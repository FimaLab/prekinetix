###### Подключение пакетов

import streamlit as st

#предварительный просмотр общего доступа
st.set_page_config(page_title="BPK", page_icon="favicon.png", layout="centered", initial_sidebar_state="auto", menu_items=None)

with open('style.css') as f:
    st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import math
import statistics  
import seaborn as sns
import statsmodels.api as sm
import streamlit.components as stc
from io import BytesIO
from pyxlsb import open_workbook as open_xlsb
from docx import Document
from docx.shared import Pt
import tempfile
import os
from cycler import cycler
from streamlit_option_menu import option_menu
import streamlit.components.v1 as components 
import codecs
from streamlit_js_eval import streamlit_js_eval

############Для запуска приложения в консоле

#cd C:\Users\Павел\OneDrive\Worktable\pyt\Bioavailability
#активировать среду my_env_name\scripts\activate
#streamlit run "C:\Users\Павел\OneDrive\Worktable\pyt\Bioavailability\validation.py"
#после введения команды в консоль, закрепляем открытую вкладку в браузере
#для просмотра измененного кода нужно обязательно сохранять файл .py и перезагружать вкладку

#C:\Users\Павел\AppData\Local\Programs\Python\Python310\Lib\site-packages путь ко всем пакетам

### сделать exe файл: 
# 1) Открыть Node.js command prompt
# 2)nativefier  --name "BPK" --icon "C:\Users\Павел\OneDrive\Worktable\icon_final_total.ico" "https://bioavailability-pk.streamlit.app" 

### создать и обновить файл требований
# pip freeze > requirements.txt

###########################################################
#область глобальных стилей

st.markdown(
    """
<style>
span[data-baseweb="tag"] {
  background-color: #3095eb !important;
}
label[data-baseweb="checkbox"] {
  background-color: #1D1D1D !important;
  border-radius: 5px;
}
label[data-baseweb="checkbox"] span {
  background-color: #2e4f4f !important;
  margin-left: 6px;
}
div[data-baseweb="notification"] {
  background-color: #2e4f4f !important;
}
</style>
""",
    unsafe_allow_html=True,
)

####[theme] синяяя тема
#primaryColor="#FFFFFF"
#backgroundColor="#04748c"
#secondaryBackgroundColor="#5d959f"
#textColor="#FFFFFF"
#font="sans serif"

#<style>
#span[data-baseweb="tag"] {
 # background-color: #0f7c9bbf !important;
#}
#label[data-baseweb="checkbox"] {
 # background-color: #355b70 !important;
 # border-radius: 5px;
#}
#label[data-baseweb="checkbox"] span {
 # background-color: #50a0af !important;
 # margin-left: 6px;
#}
#</style>

#основного меню 

#styles={
 #            "container": {"padding": "0!important", "background-color": "#24769C"},
 #            "icon": {"color": "#5DAED3", "font-size": "18px"}, 
 #            "nav-link": {"font-size": "18px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
  #           "nav-link-selected": {"background-color": "#335D70"},
  #       })

#span[data-baseweb="tag"] - стиль тега в селекторе
#label[data-baseweb="checkbox"] - стиль чекбокса всего
#label[data-baseweb="checkbox"] span - стиль чекбокса


#область глобальных функций

#сохранение загружаемых файлов 
def save_uploadedfile(uploadedfile):
    with open(os.path.join("Папка для сохранения файлов",uploadedfile.name),"wb") as f:
       f.write(uploadedfile.getbuffer())
    return st.success("Файл загружен")

#сохранение редактируемых файлов df_edit
def save_editfile(df_edit,uploadedfile_name):
    writer=pd.ExcelWriter(os.path.join("Папка для сохранения файлов",uploadedfile_name))
    df_edit.to_excel(writer,index=False)
    writer.save()

#превращает df в excel файл 
def to_excel(df_example_file):
       output = BytesIO()
       writer = pd.ExcelWriter(output, engine='xlsxwriter')
       df_example_file.to_excel(writer, index=False, sheet_name='Sheet1')
       workbook = writer.book
       worksheet = writer.sheets['Sheet1']
       format1 = workbook.add_format({'num_format': '0.00'}) 
       worksheet.set_column('A:A', None, format1)
       writer.save()  
       processed_data = output.getvalue()
       return processed_data
 
 ###возможность редактирования фрейма исходных данных
def edit_frame(df,uploadedfile_name):
       new_df = df
       list_columns_str = []
       for i in new_df.columns.tolist():
           i_new = str(i)
           list_columns_str.append(i_new)
       new_df.columns = list_columns_str

       edited_df = st.data_editor(new_df, key = ("edited_df" + uploadedfile_name))
       save_editfile(edited_df,uploadedfile_name)

       df_change = edited_df
       
       list_change_values = df_change.columns.tolist()
       list_change_values.remove("Номер")

       list_columns_number = []
       for i in list_change_values:
           i_new = float(i)
           list_columns_number.append(i_new)

       list_columns_number.insert(0,"Номер")

       df_change.columns = list_columns_number
       
       df = df_change
       return df

###создание Word-отчета
## функция создания отчета таблиц

def create_table(list_heading_word,list_table_word):
    ### таблицы
    zip_heading_table = zip(list_heading_word,list_table_word)

    doc = Document()

    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)
    
    for heading, df in zip_heading_table:
        doc.add_paragraph(heading)

        name_columns = pd.DataFrame(df.columns.tolist()).T
        # add columns
        name_columns.columns = df.columns.tolist()
        df_columns = pd.concat([name_columns, df]).reset_index(drop = True)
        # add indexes
        total_name_index = df.index.name
        list_index_names = df.index.tolist()
        list_index_names.insert(0,total_name_index)
        series_index_names=pd.Series(list_index_names, name=total_name_index)
        df_series_index_names = series_index_names.to_frame()
        
        df_columns_indexes=pd.concat([df_series_index_names, df_columns], axis=1)
        
        t = doc.add_table(rows=1, cols=df_columns_indexes.shape[1])
        t.style = 'TableGrid'
        # Add the body of the data frame
        for i in range(df_columns_indexes.shape[0]):
            row = t.add_row()
            for j in range(df_columns_indexes.shape[1]):
                cell = df_columns_indexes.iat[i, j]
                row.cells[j].text = str(cell)

    bio = BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Сохранить таблицы 📃",
            data=bio.getvalue(),
            file_name="Таблицы.docx",
            mime="docx"
        )

def visualize_table(list_heading_word,list_table_word):
    zip_heading_table = zip(list_heading_word,list_table_word) ###еще раз объявляем, иначе не видит zip-объект
    #####визуализация
    for heading, df in zip_heading_table:
        st.subheader(heading)
        st.write(df)

## функция создания отчета графиков
def create_graphic(list_graphics_word,list_heading_graphics_word):
    ### документ Word
    zip_graphics_heading = zip(list_graphics_word,list_heading_graphics_word)
    doc = Document()

    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    buf = BytesIO() #костыль для того, чтобы не вылазила ошибка
    for fig, heading in zip_graphics_heading:
        buf = BytesIO()
        fig.savefig(buf, format="jpg", dpi=300, bbox_inches='tight')
        fp = tempfile.NamedTemporaryFile() 
        with open(f"{fp.name}.jpg",'wb') as ff:
             ff.write(buf.getvalue()) 
        doc.add_picture(buf)
        doc.add_paragraph(heading)
    
    doc.save(buf)
    if doc:
        st.download_button(
            label="Сохранить графики 📈",
            data=buf.getvalue(),
            file_name="Графики.docx",
            mime="docx",
            key = "graphics"
        )

## функция подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
def create_table_descriptive_statistics(df):
    col_mapping = df.columns.tolist()
    col_mapping.remove('Номер')

    list_gmean=[]
    list_cv=[] 
    for i in col_mapping:

        list_ser=df[i].tolist()
        list_ser_cv = list_ser#нужно с нулями для CV

        #убрать нули, т.к нули будут давать нулевое gmean
        count_for_range_ser=len(list_ser)
        list_range_ser=range(0,count_for_range_ser)
        
        list_ser_without_0=[]
        for i in list_range_ser:
            if list_ser[i] !=0:
               list_ser_without_0.append(list_ser[i])

        list_ser = list_ser_without_0

        def g_mean(list_ser):
            a=np.log(list_ser)
            return np.exp(a.mean())
        Gmean=g_mean(list_ser)
        list_gmean.append(Gmean)
        
        ####CV
        cv_std=lambda x: np.std(x, ddof= 1 )
        cv_mean=lambda x: np.mean(x)
        CV_std=cv_std(list_ser_cv)
        CV_mean=cv_mean(list_ser_cv)
        CV=CV_std/CV_mean * 100
        list_cv.append(CV)
        
    #для устранения None из фрейма
    list_gmean.pop(0)
    list_gmean.insert(0,0)
    list_cv.pop(0)
    list_cv.insert(0,0)
    
    df_averaged_concentrations=df.describe()
    df_averaged_concentrations_1= df_averaged_concentrations.drop(['count', '25%','75%'],axis=0)
    df_averaged_concentrations_2= df_averaged_concentrations_1.rename(index={"50%": "median"})
    df_averaged_concentrations_2.loc[len(df_averaged_concentrations_2.index )] = list_gmean
    df_averaged_3 = df_averaged_concentrations_2.rename(index={5 : "Gmean"})
    df_averaged_3.loc[len(df_averaged_3.index )] = list_cv
    df_averaged_3 = df_averaged_3.rename(index={6 : "CV, %"})

    df_index=df.set_index('Номер')
    df_concat = pd.concat([df_index,df_averaged_3],sort=False,axis=0)
    df_concat_round=df_concat.round(2)

    ###визуализация фрейма с нулями после округления
    col_mapping = df_concat_round.columns.tolist()

    list_list_series=[]
    for i in col_mapping:
        list_series = df_concat_round[i].tolist()
         
        list_series_round = []
        for i in list_series:
            value = "%.2f" % round(i,2)
            list_series_round.append(value)
             
        list_list_series.append(list_series_round)

    df_concat_round_str = pd.DataFrame(list_list_series, columns = df_concat_round.index.tolist(),index=col_mapping) 
    df_concat_round_str_transpose = df_concat_round_str.transpose()
    df_concat_round_str_transpose.index.name = 'Номер'
    
    ##изменение названий параметров описательной статистики

    df_concat_round_str_transpose1=df_concat_round_str_transpose.copy()
    df_concat_round_str_transpose1.iloc[-6,:],df_concat_round_str_transpose1.iloc[-2,:]=df_concat_round_str_transpose.iloc[-2,:],df_concat_round_str_transpose.iloc[-6,:]
    df_concat_round_str_transpose=df_concat_round_str_transpose1
    df_concat_round_str_transpose1=df_concat_round_str_transpose.copy()
    df_concat_round_str_transpose1.iloc[-4,:],df_concat_round_str_transpose1.iloc[-5,:]=df_concat_round_str_transpose.iloc[-5,:],df_concat_round_str_transpose.iloc[-4,:]
    df_concat_round_str_transpose=df_concat_round_str_transpose1
    df_concat_round_str_transpose = df_concat_round_str_transpose.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')
    
    #округление времени в качестве названий стоблцов
    list_time_round =["%.2f" % round(v,2) for v in df_concat_round_str_transpose.columns.tolist()]
    df_concat_round_str_transpose.columns = list_time_round
    
    dict_descriptive_statistics = {'df_concat_round_str_transpose': df_concat_round_str_transpose,'df_concat': df_concat}
    return dict_descriptive_statistics

#############################################################

####### Главное меню

#кнопка перезагрузки приложения
button_upload_app = st.sidebar.button('Обновить страницу', key = "Перезагрузка приложения", use_container_width = False)
if button_upload_app:
   streamlit_js_eval(js_expressions="parent.window.location.reload()")

### пустое пространство 

st.sidebar.title('ㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤㅤ')

st.sidebar.image("logo-new.png", width=110)

st.sidebar.title('Лаборатория фармакокинетики и метаболомного анализа')

st.title('Добро пожаловать в приложение по расчёту ФК параметров ')

selected = option_menu(None, ["Главная", "Исследование", 'Настройки'], 
         icons=['house-door', 'graph-up','wrench-adjustable'], 
         menu_icon="cast", default_index=0, orientation="horizontal",
         styles={
             "container": {"padding": "0!important", "background-color": "#2e4f4f"},
             "icon": {"color": "#cbe4de", "font-size": "18px"}, 
             "nav-link": {"font-size": "18px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
             "nav-link-selected": {"background-color": "#0e8388"},
         })

##########Главная
if selected == "Главная":
   
   with st.sidebar:
        file = codecs.open("часы.html","r",encoding="utf8")
        page = file.read()
        components.html(page,width=None, height=143, scrolling=False)

   st.sidebar.caption('© 2023. Павел Резванов')
   
   col1, col2 = st.columns([0.66, 0.34])

   file = codecs.open("Главный_текст.html","r",encoding="utf8")
   page = file.read()
   components.html(page,width=None, height=1000, scrolling=False)

##########Исследование

if selected == "Исследование":
   ####### Левое боковое меню
   st.sidebar.title('Меню')

   st.sidebar.subheader('Какое исследование проводится?')

   option = st.sidebar.selectbox('Выберите вид исследования',
       ('Фармакокинетика','Биодоступность', 'ФК в органах', 'Линейность дозирования','Экскреция препарата'),disabled = False, key = "Вид исследования")

   ############### файл пример

   df_example_file = pd.read_excel("server_example_file.xlsx")
   df_example_file_xlsx = to_excel(df_example_file)
   st.sidebar.download_button(label='Пример файла 💾', data=df_example_file_xlsx , file_name= 'example_file.xlsx')

   ############ памятка

   text_contents = '''1)Оглавлять колонку с номерами животных должно слово «Номер» (в верхнем регистре).
   2)Знак «№» обязательно должен присутствовать при указании номера животного, иначе приложение выдаст ошибку. 
   3) Не ставить в ячейки знак «-» в случае нулевого значения. Ставить число «0» для корректной работы приложения.
   4)Ни в каком исследовании загружаемые файлы не должны называться одинаково.
   '''
   st.sidebar.download_button('Памятка заполнения 📄', text_contents)
   
   ################################
   if option == 'Фармакокинетика':
      
       st.title('Расчет фармакокинетических параметров')

       col1, col2 = st.columns([0.66, 0.34])
      
       ####### основной экран
       with col1:
           
           panel = st.radio(
               "⚙️Панель управления",
               ("Загрузка файлов", "Таблицы","Графики"),
               horizontal=True, key= "Загрузка файлов - Расчет фармакокинетических параметров"
           )
         
           ###создание состояния
           if "measure_unit_pk" not in st.session_state:
              st.session_state["measure_unit_pk"] = ""
           if "dose_pk" not in st.session_state:
              st.session_state["dose_pk"] = ""
           
           #cписки для word-отчета
           list_heading_word=[]
           list_table_word=[]
           list_graphics_word=[]
           list_heading_graphics_word=[]

           if panel == "Загрузка файлов":
              
              ######### боковое меню справа
              with col2:
                   selected = option_menu(None, ["Включение параметров в исследование"], 
                   icons=['menu-button'], 
                   menu_icon="cast", default_index=0, orientation="vertical",
                   styles={
                       "container": {"padding": "0!important", "background-color": "#2e4f4f"},
                       "icon": {"color": "#cbe4de", "font-size": "13px"}, 
                       "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
                       "nav-link-selected": {"background-color": "#0e8388"},
                   })

                   if selected == "Включение параметров в исследование":
                      type_parameter = st.selectbox('Выберите параметр',
                   ('Второй максимум',"Вид введения"),disabled = False, key = "Вид параметра - фк")
                      
                   if "agree_cmax2 - фк" not in st.session_state:
                         st.session_state["agree_cmax2 - фк"] = False

                   if type_parameter == 'Второй максимум':

                      st.session_state["agree_cmax2 - фк"] = st.checkbox('Добавить возможность выбора второго максимума на ФК кривой', key = "Возможность добавления Cmax2 - фк", value = st.session_state["agree_cmax2 - фк"])
                      
                      if st.session_state["agree_cmax2 - фк"] == True:
                         st.write('✔️Параметр добавлен!')

                   if "agree_injection - фк" not in st.session_state:
                         st.session_state["agree_injection - фк"] = False

                   if type_parameter == "Вид введения":

                      st.session_state["agree_injection - фк"] = st.checkbox('Внутривенное введение', key = "Возможность добавления injection - фк", value = st.session_state["agree_injection - фк"])
                      
                      if st.session_state["agree_injection - фк"] == True:
                         st.write('💉Параметр добавлен!')
              
              measure_unit_pk = st.text_input("Введите единицы измерения концентрации", key='Единицы измерения при расчете фармакокинетических параметров', value = st.session_state["measure_unit_pk"])
              
              st.session_state["measure_unit_pk"] = measure_unit_pk

              #cостояние радио-кнопки "method_auc"
              if "index_method_auc - фк" not in st.session_state:
                  st.session_state["index_method_auc - фк"] = 0

              method_auc = st.radio("📌Метод подсчёта AUC0-t",('linear',"linear-up/log-down"),key = "Метод подсчёта AUC0-t - фк", index = st.session_state["index_method_auc - фк"])
              
              if st.session_state["Метод подсчёта AUC0-t - фк"] == 'linear':
                 st.session_state["index_method_auc - фк"] = 0
              if st.session_state["Метод подсчёта AUC0-t - фк"] == "linear-up/log-down":
                 st.session_state["index_method_auc - фк"] = 1
                           
              uploaded_file_pk = st.file_uploader("Выбрать файл концентраций ЛС (формат XLSX)", key='Файл введения ЛС при расчете фк')
              
              #сохранение файла
              if uploaded_file_pk is not None:
                 save_uploadedfile(uploaded_file_pk)
                 st.session_state["uploaded_file_pk"] = uploaded_file_pk.name

              dose_pk = st.text_input("Доза при введении ЛС", key='Доза при введении ЛС при при расчете фк', value = st.session_state["dose_pk"])
              
              st.session_state["dose_pk"] = dose_pk
              
              if "uploaded_file_pk" in st.session_state and dose_pk and measure_unit_pk:

                 df = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_pk"]))
                 st.subheader('Индивидуальные значения концентраций в крови после введения ЛС')
                 
                 ###интерактивная таблица
                 df = edit_frame(df,st.session_state["uploaded_file_pk"])

                 ###количество животных 
                 count_rows_number_pk= len(df.axes[0])
           
                 table_heading='Индивидуальные и усредненные значения концентраций в крови после введения ЛС'
                 list_heading_word.append(table_heading)

                 ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                 df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']
                 
                 list_table_word.append(df_concat_round_str_transpose)
              
              ########### графики    

              ######индивидуальные    

                 # в линейных координатах
                 col_mapping = df.columns.tolist()
                 col_mapping.remove('Номер')

                 count_row_df = len(df.axes[0])

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)
                 
                 if st.session_state["agree_injection - фк"] == True: 
                    list_time.remove(0)

                 for r in range(0,count_row_df):

                     list_concentration=df.iloc[r].tolist()

                     numer_animal=list_concentration[0]

                     list_concentration.pop(0) #удаление номера животного

                     list_concentration = [float(v) for v in list_concentration]

                     if st.session_state["agree_injection - фк"] == True:
                        list_concentration.remove(0)


                     fig, ax = plt.subplots()
                     plt.plot(list_time,list_concentration,marker='o',markersize=4.0, color = "black", markeredgecolor="black",markerfacecolor="black")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit_pk)
                    
                     list_graphics_word.append(fig)  

                     graphic='График индивидуального фармакокинетического профиля в крови (в линейных координатах) после введения ЛС,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

                  #в полулогарифмических координатах методом удаления точек
                     count_for_0_1=len(list_concentration)
                     list_range_for_0_1=range(0,count_for_0_1)

                     list_time_0=[]
                     list_for_log_1=[]
                     for i in list_range_for_0_1:
                         if list_concentration[i] !=0:
                            list_for_log_1.append(list_concentration[i])
                            list_time_0.append(list_time[i]) 

                     fig, ax = plt.subplots()
                     plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,color = "black",markeredgecolor="black",markerfacecolor="black")
                     ax.set_yscale("log")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit_pk)

                     list_graphics_word.append(fig) 

                     graphic='График индивидуального фармакокинетического профиля в крови (в полулогарифмических координатах) после введения ЛС,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

              # объединенные индивидуальные в линейных координатах

                 df_for_plot_conc=df.drop(['Номер'], axis=1)
                 df_for_plot_conc_1 = df_for_plot_conc.transpose()

                 if st.session_state["agree_injection - фк"] == True:
                    df_for_plot_conc_1=df_for_plot_conc_1.replace(0, None) ###т.к. внутривенное

                 list_numer_animal_for_plot=df['Номер'].tolist()
                 count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды

                 list_color = ["blue","green","red","#D6870C","violet","gold","indigo","magenta","lime","tan","teal","coral","pink","#510099","lightblue","yellowgreen","cyan","salmon","brown","black"]

                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit_pk)
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 

                 graphic="Сравнение индивидуальных фармакокинетических профилей (в линейных координатах) после введения ЛС"
                 list_heading_graphics_word.append(graphic)    
              # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                 df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)

                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit_pk)
                 ax.set_yscale("log")
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 
          
                 graphic="Сравнение индивидуальных фармакокинетических профилей (в полулогарифмических координатах) после введения ЛС"
                 list_heading_graphics_word.append(graphic) 

              ### усреденные    
              #в линейных    

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 df_averaged_concentrations=df.describe()
                 list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                 err_y_pk=df_averaged_concentrations.loc['std'].tolist()
                 
                 if st.session_state["agree_injection - фк"] == True:
                    list_time.remove(0) ###т.к. внутривенное
                    list_concentration.remove(0)
                    err_y_pk.remove(0)

                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_pk, marker='o',markersize=4.0,color = "black",markeredgecolor="black",markerfacecolor="black",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit_pk)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в линейных координатах) после введения ЛС'
                 list_heading_graphics_word.append(graphic)  

              #в полулогарифмических координатах
                 if st.session_state["agree_injection - фк"] == False:
                    list_time.remove(0) ###т.к. внутривенное
                    list_concentration.remove(0)
                    err_y_pk.remove(0) 


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_pk, marker='o',markersize=4.0,color = "black",markeredgecolor="black",markerfacecolor="black",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 ax.set_yscale("log")
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit_pk)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в полулогарифмических координатах) после введения ЛС'
                 list_heading_graphics_word.append(graphic)
                 
                 ############ Параметры ФК

                 df_without_numer=df.drop(['Номер'],axis=1)
                 count_row=df_without_numer.shape[0]

                 list_count_row=range(count_row)
       
                 ###Cmax
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                 if st.session_state["agree_cmax2 - фк"] == False:
                    list_cmax_1_pk=[]
                    for i in range(0,count_row):
                        cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                        list_cmax_1_pk.append(cmax)
                 
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                 if st.session_state["agree_cmax2 - фк"] == True:
                    ###создание состояния
                    if "selected_value_pk" not in st.session_state:
                       st.session_state["selected_value_pk"] = []
                    
                    if "feature_disable_selected_value_pk" not in st.session_state:
                        st.session_state["feature_disable_selected_value_pk"] = True

                    ###создание состояния
                    st.info('Выбери Cmax:')
                    list_columns_without_numer = df.columns.tolist()
                    list_columns_without_numer.remove('Номер')
                    selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax введения ЛС',max_selections=1)
                    st.session_state["selected_columns_pk"] = selected_columns 

                    list_keys_cmax = st.session_state["selected_value_pk"]
                    if selected_columns != [] and st.session_state["feature_disable_selected_value_pk"]:
                       selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax введения ЛС',max_selections=1)
                       list_keys_cmax.append(selected_value)

                    if list_keys_cmax != []:
                       st.session_state["selected_value_pk"] = list_keys_cmax

                    list_keys_cmax = st.session_state["selected_value_pk"]
                    list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]

                    if st.button('Очистить список Cmax', key="Очистка списка Cmax введения ЛС"):
                       del st.session_state["selected_value_pk"]
                       list_keys_cmax_sample = []
                       selected_columns = st.session_state["selected_columns_pk"]
                       st.session_state["feature_disable_selected_value_pk"] = True
                                           
                    st.write("Список Cmax:")
                    st.write(list_keys_cmax_sample)
                    
                    list_cmax_1_pk=list_keys_cmax_sample 
                    
                    list_cmax_2_pk=[]

                 if len(list_cmax_1_pk) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - фк"] == True):
                    st.session_state["feature_disable_selected_value_pk"] = False

                    ######Cmax2

                    if "feature_disable_selected_value_pk_2" not in st.session_state:
                     st.session_state["feature_disable_selected_value_pk_2"] = True

                    st.info('Выбери Cmax(2):')
                    
                    selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 введения ЛС', max_selections=1)
                    st.session_state["selected_columns_2_pk"] = selected_columns_2

                    ###создание состояния
                    if "selected_value_2_pk" not in st.session_state:
                       st.session_state["selected_value_2_pk"] = []

                    list_keys_cmax_2 = st.session_state["selected_value_2_pk"]
                    if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_pk_2"]:
                       selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 введения ЛС', max_selections=1)
                       list_keys_cmax_2.append(selected_value_2)

                    if list_keys_cmax_2 != []:
                       st.session_state["selected_value_2_pk"] = list_keys_cmax_2

                    list_keys_cmax_2 = st.session_state["selected_value_2_pk"]
                    list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                    if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) введения ЛС"):
                       del st.session_state["selected_value_2_pk"]
                       list_keys_cmax_sample_2 = []
                       selected_columns_2 = st.session_state["selected_columns_2_pk"]
                       st.session_state["feature_disable_selected_value_pk_2"] = True

                    st.write("Список Cmax(2):")
                    st.write(list_keys_cmax_sample_2)

                    list_cmax_2_pk= list_keys_cmax_sample_2

                    if len(list_cmax_2_pk) == len(df.index.tolist()):
                       st.session_state["feature_disable_selected_value_pk_2"] = False

                 if (len(list_cmax_1_pk) == len(df.index.tolist())):
                    
                    ###Tmax   
                    list_Tmax_1=[]
                    for cmax in list_cmax_1_pk:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_1.append(f"{column}")
                  
                    list_Tmax_float_1=[]           
                    for i in list_Tmax_1:
                        Tmax=float(i)
                        list_Tmax_float_1.append(Tmax)

                 if (len(list_cmax_1_pk) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - фк"] == True):
                    
                    list_Tmax_2=[]
                    for cmax in list_cmax_2_pk:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_2.append(f"{column}")
                  
                    list_Tmax_float_2=[]           
                    for i in list_Tmax_2:
                        Tmax=float(i)
                        list_Tmax_float_2.append(Tmax)  

                 if (len(list_cmax_1_pk) == len(df.index.tolist())):
                    
                    ###AUC0-t
                    list_AUC_0_T=[]
                    if method_auc == 'linear':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                           list_AUC_0_T.append(AUC_0_T)

                    if method_auc == 'linear-up/log-down':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################
                           
                           list_c = list_concentration
                           list_t = list_columns_T
                           
                           count_i = len(list_c)
                           list_range= range(0,count_i)
                           
                           list_AUC_0_T_ascending=[]
                           list_AUC_0_T_descending = []
                           AUC_0_T_ascending=0
                           AUC_0_T_descending = 0
                           a=0
                           a1=0
                           d=0
                           d1=0
                           for i in list_range:
                               if a1<count_i-1:
                                  if list_c[i+1] > list_c[i]:
                                     if a<count_i-1:
                                         AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                         a+=1
                                         list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                               if d1<count_i-1:
                                  if list_c[i+1] < list_c[i]:      
                                     if d<count_i-1:
                                         AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                         d+=1
                                         list_AUC_0_T_descending.append(AUC_0_T_descending)
                                  a1+=1
                                  d1+=1
            
                           AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                           
                           list_AUC_0_T.append(AUC_O_T)

                    ####Сmax/AUC0-t
                    list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_pk,list_AUC_0_T)
                    list_Сmax_division_AUC0_t=[]
                    for i,j in list_Сmax_division_AUC0_t_for_division:
                            list_Сmax_division_AUC0_t.append(i/j)


                    ####KEL
                    list_kel_total=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c=list_concentration

                        list_time=df_without_numer.columns.tolist()
                        list_time.remove(0) 

                        list_t=[]
                        for i in list_time:
                            i=float(i)
                            list_t.append(i)

                        #срез_без_cmax
                        max_value_c=max(list_c)
                        index_cmax=list_c.index(max_value_c)

                        list_c_without_cmax=list_c[index_cmax+1:]
                        list_t_without_cmax=list_t[index_cmax+1:]

                        #удаление всех нулей из массивов
                        count_for_0_1=len(list_c_without_cmax)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_conc_0=[]
                        for i in list_range_for_0_1:
                            if list_c_without_cmax[i] !=0:
                               list_conc_0.append(list_c_without_cmax[i])
                               list_time_0.append(list_t_without_cmax[i]) 
                        ################################

                        n_points=len(list_conc_0)
                        list_n_points = range(0,n_points)

                        #создание списков с поочередно уменьщающемся кол, точек
                        list_for_kel_c=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_c_new=list_conc_0[j:n_points]
                               list_for_kel_c.append(list_c_new)
                        list_for_kel_c.pop(-1) #удаление списка с одной точкой
                        list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                        list_for_kel_t=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_t_new=list_time_0[j:n_points]
                               list_for_kel_t.append(list_t_new)
                        list_for_kel_t.pop(-1) #удаление списка с одной точкой
                        list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                        list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                        list_kel=[]
                        list_r=[]
                        for i,j in list_ct_zip:

                            n_points_r=len(i)

                            np_c=np.asarray(i)
                            np_t_1=np.asarray(j).reshape((-1,1))

                            np_c_log=np.log(np_c)

                            model = LinearRegression().fit(np_t_1,np_c_log)

                            np_t=np.asarray(j)
                            a=np.corrcoef(np_t, np_c_log)
                            cor=((a[0])[1])
                            r_sq=cor**2

                            adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                            ########################################
                            kel=abs(model.coef_[0])
                            list_kel.append(kel)
                            list_r.append(adjusted_r_sq)

                        #делаем срезы списоков до rmax
                        max_r=max(list_r)

                        index_max_r= list_r.index(max_r)

                        list_r1=list_r
                        list_kel1=list_kel

                        number_elem_list_r1=len(list_r1)

                        list_range_kel=range(0,number_elem_list_r1) 

                        list_kel_total_1=[]
                        for i in list_range_kel:

                            if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                               list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                               break #самая ранняя удовлетовряющая условию

                        for i in list_kel_total_1:
                            list_kel_total.append(i) 


                    ####T1/2
                    list_half_live=[]
                    for i in list_kel_total:
                        half_live=math.log(2)/i
                        list_half_live.append(half_live)


                    ###AUC0-inf 

                    list_auc0_inf=[] 

                    list_of_list_c=[]
                    for i in range(0,count_row):
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c = list_concentration
                        list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                        list_of_list_c.append(list_c)

                    list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                        #AUCt-inf 
                    list_auc_t_inf=[]     
                    for i,j in list_zip_c_AUCt_inf:
                        for clast in j:
                            if clast != 0:
                               clast_true=clast
                               break
                        auc_t_inf=clast_true/i
                        list_auc_t_inf.append(auc_t_inf)

                    list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                    for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                        auc0_inf=i+j    
                        list_auc0_inf.append(auc0_inf)


                    ####CL
                    list_cl=[]

                    for i in list_auc0_inf:
                        cl = float(dose_pk)/i *1000000
                        list_cl.append(cl) 


                    ####Vd
                    list_Vd=[]

                    list_zip_kel_cl=zip(list_kel_total,list_cl)

                    for i,j in list_zip_kel_cl:
                        Vd = j/i
                        list_Vd.append(Vd)


                    ###AUMC0-t и ###AUMC0-inf
                    list_AUMCO_inf=[]

                    list_AUMC0_t=[]

                    list_C_last=[]
                    list_T_last=[]

                    if method_auc == 'linear':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           list_C_last.append(list_concentration[-1]) 
                           list_T_last.append(list_columns_T[-1]) 

                           list_len=len(list_concentration)

                           list_aumc_i=[]
                           for i in range(0,list_len):
                               AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                               list_aumc_i.append(AUMC)

                           list_aumc_i.pop(0)

                           a=0
                           list_AUMC0_t_1=[]
                           for i in list_aumc_i:
                               a+=i
                               list_AUMC0_t_1.append(a)
                           list_AUMC0_t.append(list_AUMC0_t_1[-1])
                    
                    if method_auc == 'linear-up/log-down':
                       
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           list_C_last.append(list_concentration[-1]) 
                           list_T_last.append(list_columns_T[-1])

                           list_c = list_concentration
                           list_t = list_columns_T

                           count_i = len(list_c)
                           list_range= range(0,count_i)

                           list_AUMC_0_T_ascending=[]
                           list_AUMC_0_T_descending = []
                           AUMC_0_T_ascending=0
                           AUMC_0_T_descending = 0
                           a=0
                           a1=0
                           d=0
                           d1=0
                           for i in list_range:
                               if a1<count_i-1:
                                   if list_c[i+1] > list_c[i]:
                                       if a<count_i-1:
                                           AUMC_0_T_ascending +=(list_t[i+1] - list_t[i]) *  ((list_c[i+1] * list_t[i+1] + list_c[i] * list_t[i])/2)
                                           a+=1
                                           list_AUMC_0_T_ascending.append(AUMC_0_T_ascending)
                               if d1<count_i-1:
                                   if list_c[i+1] < list_c[i]:      
                                       if d<count_i-1:
                                           coeff = (list_t[i+1] - list_t[i]) / np.log(np.asarray(list_c[i+1])/np.asarray(list_c[i]))
                                           AUMC_0_T_descending+= coeff * ((list_c[i+1] * list_t[i+1] - list_c[i] * list_t[i]) - coeff * (list_c[i+1] - list_c[i]))
                                           d+=1
                                           list_AUMC_0_T_descending.append(AUMC_0_T_descending)
                                   a1+=1
                                   d1+=1

                           AUMC_O_T = list_AUMC_0_T_ascending[-1]+list_AUMC_0_T_descending[-1]

                           list_AUMC0_t.append(AUMC_O_T)

                    ########AUMC0-inf конечный подсчет
                    list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                    list_AUMCt_inf=[]
                    for k,c,t in list_zip_for_AUMC_inf:
                        AUMCt_inf=c*t/k+c/(k*k)
                        list_AUMCt_inf.append(AUMCt_inf)


                    list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                    for i,j in list_AUMC_zip:
                        AUMCO_inf=i+j
                        list_AUMCO_inf.append(AUMCO_inf)

                    ###MRT0-inf
                    list_MRT0_inf=[]

                    list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                    for i,j in list_zip_AUMCO_inf_auc0_inf:
                        MRT0_inf=i/j
                        list_MRT0_inf.append(MRT0_inf)
                 
                 if st.session_state["agree_cmax2 - фк"] == True:
                    #####Cmax условие для дальнейшего кода
                    if len(list_cmax_1_pk) == len(df.index.tolist()) and len(list_cmax_2_pk) == len(df.index.tolist()):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_pk,list_Tmax_float_1,list_cmax_2_pk,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 

                 if len(list_cmax_1_pk) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - фк"] == False):
                    
                    ##################### Фрейм ФК параметров

                    ### пользовательский индекс
                    list_for_index=df["Номер"].tolist()
                    df_PK=pd.DataFrame(list(zip(list_cmax_1_pk,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index)
                 
                 checking_condition_cmax2 = False

                 if st.session_state["agree_cmax2 - фк"] == True:
                     
                    checking_condition_cmax2 = len(list_cmax_1_pk) == len(df.index.tolist()) and len(list_cmax_2_pk) == len(df.index.tolist()) and st.session_state["agree_cmax2 - фк"] == True

                 if checking_condition_cmax2 or (len(list_cmax_1_pk) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - фк"] == False)):
                 
                    ###описательная статистика

                    col_mapping_PK = df_PK.columns.tolist()

                    list_gmean_PK=[]

                    list_cv_PK=[] 

                    for i in col_mapping_PK:

                        list_ser_PK=df_PK[i].tolist()

                        def g_mean(list_ser_PK):
                            a=np.log(list_ser_PK)
                            return np.exp(a.mean())
                        Gmean_PK=g_mean(list_ser_PK)
                        list_gmean_PK.append(Gmean_PK)

                        cv_std_PK=lambda x: np.std(x, ddof= 1 )
                        cv_mean_PK=lambda x: np.mean(x)

                        CV_std_PK=cv_std_PK(list_ser_PK)
                        CV_mean_PK=cv_mean_PK(list_ser_PK)

                        CV_PK=(CV_std_PK/CV_mean_PK * 100)
                        list_cv_PK.append(CV_PK)


                    df_averaged_concentrations_PK=df_PK.describe()
                    df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                    df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                    df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                    df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                    df_round_without_CV_PK=df_averaged_3_PK
                    df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                    df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                    df_concat_PK_pk= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)

                    ###округление описательной статистики и ФК параметров

                    series_Cmax=df_concat_PK_pk['Cmax']
                    list_Cmax_str_f=["%.10f" % round(v,10) for v in series_Cmax.tolist()]
                    series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_pk.index.tolist(), name='Cmax ' +"("+measure_unit_pk+")")

                    if st.session_state["agree_cmax2 - фк"] == True:
                       series_Cmax_2=df_concat_PK_pk['Cmax(2)']
                       list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                       series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_pk.index.tolist(), name='Cmax(2) ' +"("+measure_unit_pk+")")

                    series_Tmax=df_concat_PK_pk['Tmax']
                    list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                    series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_pk.index.tolist(), name='Tmax ' +"("+"ч"+")")
                    
                    if st.session_state["agree_cmax2 - фк"] == True:
                       series_Tmax_2=df_concat_PK_pk['Tmax(2)']
                       list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                       series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_pk.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                    series_MRT0_inf= df_concat_PK_pk['MRT0→∞']
                    list_MRT0_inf_str_f=["%.14f" % round(v,14) for v in series_MRT0_inf.tolist()]
                    series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_pk.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                    series_half_live= df_concat_PK_pk['T1/2']
                    list_half_live_str_f=["%.15f" % round(v,15) for v in series_half_live.tolist()]
                    series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_pk.index.tolist(), name='T1/2 '+"("+"ч"+")")

                    series_AUC0_t= df_concat_PK_pk['AUC0-t']
                    list_AUC0_t_str_f=["%.11f" % round(v,11) for v in series_AUC0_t.tolist()]
                    series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_pk.index.tolist(), name='AUC0-t '+"("+measure_unit_pk+"×ч" +")")

                    series_AUC0_inf= df_concat_PK_pk['AUC0→∞']
                    list_AUC0_inf_str_f=["%.11f" % round(v,11) for v in series_AUC0_inf.tolist()]
                    series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_pk.index.tolist(), name='AUC0→∞ '+"("+measure_unit_pk+"×ч" +")")

                    series_AUMC0_inf= df_concat_PK_pk['AUMC0-∞']
                    list_AUMC0_inf_str_f=["%.11f" % round(v,11) for v in series_AUMC0_inf.tolist()]
                    series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_pk.index.tolist(), name='AUMC0-∞ '+"("+measure_unit_pk+"×ч\u00B2" +")")

                    series_Сmax_dev_AUC0_t= df_concat_PK_pk['Сmax/AUC0-t']
                    list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                    series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_pk.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                    series_Kel= df_concat_PK_pk['Kel']
                    list_Kel_str_f=["%.15f" % round(v,15) for v in series_Kel.tolist()]
                    series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_pk.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                    series_CL= df_concat_PK_pk['CL/F']
                    list_CL_str_f=["%.10f" % round(v,10) for v in series_CL.tolist()]
                    series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_pk.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                    series_Vd= df_concat_PK_pk['Vd']
                    list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                    series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_pk.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")
                    
                    if st.session_state["agree_cmax2 - фк"] == True:
                       df_total_PK_pk = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    else:
                       df_total_PK_pk = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    
                    df_total_PK_pk.index.name = 'Номер'

                    ##изменение названий параметров описательной статистики

                    df_total_PK_pk1=df_total_PK_pk.copy()
                    df_total_PK_pk1.iloc[-6,:],df_total_PK_pk1.iloc[-2,:]=df_total_PK_pk.iloc[-2,:],df_total_PK_pk.iloc[-6,:]

                    df_total_PK_pk=df_total_PK_pk1

                    df_total_PK_pk1=df_total_PK_pk.copy()
                    df_total_PK_pk1.iloc[-4,:],df_total_PK_pk1.iloc[-5,:]=df_total_PK_pk.iloc[-5,:],df_total_PK_pk.iloc[-4,:]

                    df_total_PK_pk=df_total_PK_pk1

                    df_total_PK_pk = df_total_PK_pk.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')

                    table_heading='Фармакокинетические показатели в крови после введения лс'
                    list_heading_word.append(table_heading)
                    
                    list_table_word.append(df_total_PK_pk)
                    writer = pd.ExcelWriter("C:/Users/Павел/OneDrive/Рабочий стол/Статьи для валидации приложения/проекты_феникс/проекты_феникс/my_app_res_ph_07_BE_2024.xlsx", engine='xlsxwriter')
                    
                    # Записать ваш DataFrame в файл     
                    df_total_PK_pk.to_excel(writer, 'Sheet1')

                    # Сохраним результат 
                    writer.save()
              else:
                  st.write("")
              ###сохранение состояния 
              st.session_state["list_heading_word"] = list_heading_word
              st.session_state["list_table_word"] = list_table_word
              st.session_state["list_graphics_word"] = list_graphics_word
              st.session_state["list_heading_graphics_word"] = list_heading_graphics_word
             
       #отдельная панель, чтобы уменьшить размер вывода результатов

       col1, col2 = st.columns([0.66,0.34])
       
       with col1:
        
          #####Создание word отчета
          if panel == "Таблицы":

                list_heading_word = st.session_state["list_heading_word"]
                list_table_word = st.session_state["list_table_word"]
                
                ###вызов функции визуализации таблиц
                visualize_table(list_heading_word,list_table_word)

                with col2:
                     
                     selected = option_menu(None, ["Cформированный отчeт"], 
                     icons=['file-earmark-arrow-down-fill'], 
                     menu_icon="cast", default_index=0, orientation="vertical",
                     styles={
                        "container": {"padding": "0!important", "background-color": "#2e4f4f"},
                        "icon": {"color": "#cbe4de", "font-size": "16px"}, 
                        "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
                        "nav-link-selected": {"background-color": "#0e8388"},
                     })

                     if selected == "Cформированный отчeт":

                        ###вызов функции создания Word-отчета таблиц
                        create_table(list_heading_word,list_table_word)

          if panel == "Графики":
                
                list_graphics_word = st.session_state["list_graphics_word"]
                list_heading_graphics_word = st.session_state["list_heading_graphics_word"]
                
                #######визуализация

                #классификация графиков по кнопкам
                type_graphics = st.selectbox('Выберите вид графиков',
          ('Индивидуальные фармакокинетические профили', 'Сравнение индивидуальных фармакокинетических профилей', 'Графики усредненного фармакокинетического профиля'),disabled = False, key = "Вид графика - фк" )

                count_graphics_for_visual = len(list_heading_graphics_word)
                list_range_count_graphics_for_visual = range(0,count_graphics_for_visual)
                
                for i in list_range_count_graphics_for_visual:
                    if list_heading_graphics_word[i].__contains__("индивидуального"): 
                       if type_graphics == 'Индивидуальные фармакокинетические профили':
                          st.pyplot(list_graphics_word[i])
                          st.subheader(list_heading_graphics_word[i])
                    if list_heading_graphics_word[i].__contains__("Сравнение индивидуальных"):   
                       if type_graphics == 'Сравнение индивидуальных фармакокинетических профилей':
                          st.pyplot(list_graphics_word[i])
                          st.subheader(list_heading_graphics_word[i])
                    if list_heading_graphics_word[i].__contains__("усредненного"):
                       if type_graphics == 'Графики усредненного фармакокинетического профиля':
                          st.pyplot(list_graphics_word[i])
                          st.subheader(list_heading_graphics_word[i])

                with col2:
                     
                     selected = option_menu(None, ["Cформированный отчeт"], 
                     icons=['file-earmark-arrow-down-fill'], 
                     menu_icon="cast", default_index=0, orientation="vertical",
                     styles={
                        "container": {"padding": "0!important", "background-color": "#2e4f4f"},
                        "icon": {"color": "#cbe4de", "font-size": "16px"}, 
                        "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
                        "nav-link-selected": {"background-color": "#0e8388"},
                     })
                      
                     if selected == "Cформированный отчeт":
                        ###вызов функции создания Word-отчета графиков
                        create_graphic(list_graphics_word,list_heading_graphics_word)