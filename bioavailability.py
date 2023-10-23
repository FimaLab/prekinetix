###### Подключение пакетов

import streamlit as st

with open('style.css') as f:
    st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import math
import statistics  
import seaborn as sns
import statsmodels.api as sm
import streamlit.components as stc
from io import BytesIO
from pyxlsb import open_workbook as open_xlsb
from docx import Document
from docx.shared import Pt
import tempfile
import os
import random
from cycler import cycler
from streamlit_option_menu import option_menu
import streamlit.components.v1 as components 
import codecs
from streamlit_js_eval import streamlit_js_eval

############Для запуска приложения в консоле

#cd C:\Users\Павел\OneDrive\Worktable\pyt\Bioavailability
#активировать среду my_env_name\scripts\activate
#streamlit run "C:\Users\Павел\OneDrive\Worktable\pyt\Bioavailability\bioavailability.py"
#после введения команды в консоль, закрепляем открытую вкладку в браузере
#для просмотра измененного кода нужно обязательно сохранять файл .py и перезагружать вкладку

#C:\Users\Павел\AppData\Local\Programs\Python\Python310\Lib\site-packages путь ко всем пакетам

### сделать exe файл: 
# 1) Открыть Node.js command prompt
# 2)nativefier  --name "BPK" --icon "C:\Users\Павел\OneDrive\Worktable\icon_final_total.ico" "https://bioavailability-pk.streamlit.app" 

### создать и обновить файл требований
# pip freeze > requirements.txt

###########################################################
#область глобальных стилей

st.markdown(
    """
<style>
span[data-baseweb="tag"] {
  background-color: #0f7c9bbf !important;
}
label[data-baseweb="checkbox"] {
  background-color: #355b70 !important;
  border-radius: 5px;
}
label[data-baseweb="checkbox"] span {
  background-color: #50a0af !important;
  margin-left: 6px;
}
</style>
""",
    unsafe_allow_html=True,
)

#span[data-baseweb="tag"] - стиль тега в селекторе
#label[data-baseweb="checkbox"] - стиль чекбокса всего
#label[data-baseweb="checkbox"] span - стиль чекбокса


#область глобальных функций

#сохранение загружаемых файлов 
def save_uploadedfile(uploadedfile):
    with open(os.path.join("Папка для сохранения файлов",uploadedfile.name),"wb") as f:
       f.write(uploadedfile.getbuffer())
    return st.success("Файл загружен")

#сохранение редактируемых файлов df_edit
def save_editfile(df_edit,uploadedfile_name):
    writer=pd.ExcelWriter(os.path.join("Папка для сохранения файлов",uploadedfile_name))
    df_edit.to_excel(writer,index=False)
    writer.save()

#превращает df в excel файл 
def to_excel(df_example_file):
       output = BytesIO()
       writer = pd.ExcelWriter(output, engine='xlsxwriter')
       df_example_file.to_excel(writer, index=False, sheet_name='Sheet1')
       workbook = writer.book
       worksheet = writer.sheets['Sheet1']
       format1 = workbook.add_format({'num_format': '0.00'}) 
       worksheet.set_column('A:A', None, format1)
       writer.save()  
       processed_data = output.getvalue()
       return processed_data
 
 ###возможность редактирования фрейма исходных данных
def edit_frame(df,uploadedfile_name):
       new_df = df
       list_columns_str = []
       for i in new_df.columns.tolist():
           i_new = str(i)
           list_columns_str.append(i_new)
       new_df.columns = list_columns_str

       edited_df = st.data_editor(new_df, key = ("edited_df" + uploadedfile_name))
       save_editfile(edited_df,uploadedfile_name)

       df_change = edited_df
       
       list_change_values = df_change.columns.tolist()
       list_change_values.remove("Номер")

       list_columns_number = []
       for i in list_change_values:
           i_new = float(i)
           list_columns_number.append(i_new)

       list_columns_number.insert(0,"Номер")

       df_change.columns = list_columns_number
       
       df = df_change
       return df

###создание Word-отчета
## функция создания отчета таблиц

def create_table(list_heading_word,list_table_word):
    ### таблицы
    zip_heading_table = zip(list_heading_word,list_table_word)

    doc = Document()

    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)
    
    for heading, df in zip_heading_table:
        doc.add_paragraph(heading)

        name_columns = pd.DataFrame(df.columns.tolist()).T
        # add columns
        name_columns.columns = df.columns.tolist()
        df_columns = pd.concat([name_columns, df]).reset_index(drop = True)
        # add indexes
        total_name_index = df.index.name
        list_index_names = df.index.tolist()
        list_index_names.insert(0,total_name_index)
        series_index_names=pd.Series(list_index_names, name=total_name_index)
        df_series_index_names = series_index_names.to_frame()
        
        df_columns_indexes=pd.concat([df_series_index_names, df_columns], axis=1)
        
        t = doc.add_table(rows=1, cols=df_columns_indexes.shape[1])
        t.style = 'TableGrid'
        # Add the body of the data frame
        for i in range(df_columns_indexes.shape[0]):
            row = t.add_row()
            for j in range(df_columns_indexes.shape[1]):
                cell = df_columns_indexes.iat[i, j]
                row.cells[j].text = str(cell)

    bio = BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Сохранить таблицы 📃",
            data=bio.getvalue(),
            file_name="Таблицы.docx",
            mime="docx"
        )

    zip_heading_table = zip(list_heading_word,list_table_word) ###еще раз объявляем, иначе не видит zip-объект
    #####визуализация
    for heading, df in zip_heading_table:
        st.subheader(heading)
        st.write(df)

## функция создания отчета графиков
def create_graphic(list_graphics_word,list_heading_graphics_word):
    ### документ Word
    zip_graphics_heading = zip(list_graphics_word,list_heading_graphics_word)
    doc = Document()

    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    buf = BytesIO() #костыль для того, чтобы не вылазила ошибка
    for fig, heading in zip_graphics_heading:
        buf = BytesIO()
        fig.savefig(buf, format="jpg", dpi=300, bbox_inches='tight')
        fp = tempfile.NamedTemporaryFile() 
        with open(f"{fp.name}.jpg",'wb') as ff:
             ff.write(buf.getvalue()) 
        doc.add_picture(buf)
        doc.add_paragraph(heading)
    
    doc.save(buf)
    if doc:
        st.download_button(
            label="Сохранить графики 📈",
            data=buf.getvalue(),
            file_name="Графики.docx",
            mime="docx",
            key = "graphics"
        )

## функция подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
def create_table_descriptive_statistics(df):
    col_mapping = df.columns.tolist()
    col_mapping.remove('Номер')

    list_gmean=[]
    list_cv=[] 
    for i in col_mapping:

        list_ser=df[i].tolist()
        list_ser_cv = list_ser#нужно с нулями для CV

        #убрать нули, т.к нули будут давать нулевое gmean
        count_for_range_ser=len(list_ser)
        list_range_ser=range(0,count_for_range_ser)
        
        list_ser_without_0=[]
        for i in list_range_ser:
            if list_ser[i] !=0:
               list_ser_without_0.append(list_ser[i])

        list_ser = list_ser_without_0

        def g_mean(list_ser):
            a=np.log(list_ser)
            return np.exp(a.mean())
        Gmean=g_mean(list_ser)
        list_gmean.append(Gmean)
        
        ####CV
        cv_std=lambda x: np.std(x, ddof= 1 )
        cv_mean=lambda x: np.mean(x)
        CV_std=cv_std(list_ser_cv)
        CV_mean=cv_mean(list_ser_cv)
        CV=CV_std/CV_mean * 100
        list_cv.append(CV)
        
    #для устранения None из фрейма
    list_gmean.pop(0)
    list_gmean.insert(0,0)
    list_cv.pop(0)
    list_cv.insert(0,0)
    
    df_averaged_concentrations=df.describe()
    df_averaged_concentrations_1= df_averaged_concentrations.drop(['count', '25%','75%'],axis=0)
    df_averaged_concentrations_2= df_averaged_concentrations_1.rename(index={"50%": "median"})
    df_averaged_concentrations_2.loc[len(df_averaged_concentrations_2.index )] = list_gmean
    df_averaged_3 = df_averaged_concentrations_2.rename(index={5 : "Gmean"})
    df_averaged_3.loc[len(df_averaged_3.index )] = list_cv
    df_averaged_3 = df_averaged_3.rename(index={6 : "CV, %"})

    df_index=df.set_index('Номер')
    df_concat = pd.concat([df_index,df_averaged_3],sort=False,axis=0)
    df_concat_round=df_concat.round(2)

    ###визуализация фрейма с нулями после округления
    col_mapping = df_concat_round.columns.tolist()

    list_list_series=[]
    for i in col_mapping:
        list_series = df_concat_round[i].tolist()
         
        list_series_round = []
        for i in list_series:
            value = "%.2f" % round(i,2)
            list_series_round.append(value)
             
        list_list_series.append(list_series_round)

    df_concat_round_str = pd.DataFrame(list_list_series, columns = df_concat_round.index.tolist(),index=col_mapping) 
    df_concat_round_str_transpose = df_concat_round_str.transpose()
    df_concat_round_str_transpose.index.name = 'Номер'
    
    ##изменение названий параметров описательной статистики

    df_concat_round_str_transpose1=df_concat_round_str_transpose.copy()
    df_concat_round_str_transpose1.iloc[-6,:],df_concat_round_str_transpose1.iloc[-2,:]=df_concat_round_str_transpose.iloc[-2,:],df_concat_round_str_transpose.iloc[-6,:]
    df_concat_round_str_transpose=df_concat_round_str_transpose1
    df_concat_round_str_transpose1=df_concat_round_str_transpose.copy()
    df_concat_round_str_transpose1.iloc[-4,:],df_concat_round_str_transpose1.iloc[-5,:]=df_concat_round_str_transpose.iloc[-5,:],df_concat_round_str_transpose.iloc[-4,:]
    df_concat_round_str_transpose=df_concat_round_str_transpose1
    df_concat_round_str_transpose = df_concat_round_str_transpose.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')
    
    #округление времени в качестве названий стоблцов
    list_time_round =["%.2f" % round(v,2) for v in df_concat_round_str_transpose.columns.tolist()]
    df_concat_round_str_transpose.columns = list_time_round
    
    dict_descriptive_statistics = {'df_concat_round_str_transpose': df_concat_round_str_transpose,'df_concat': df_concat}
    return dict_descriptive_statistics

#############################################################

####### Главное меню

#кнопка перезагрузки приложения
button_upload_app = st.sidebar.button('🔄', key = "Перезагрузка приложения", use_container_width = False)
if button_upload_app:
   streamlit_js_eval(js_expressions="parent.window.location.reload()")

### пустое пространство 

st.sidebar.title('ㅤㅤㅤㅤㅤㅤㅤ')

st.sidebar.image("logo-new.png", width=110)

st.sidebar.title('Лаборатория фармакокинетики и метаболомного анализа')

st.title('Добро пожаловать в приложение по расчёту ФК параметров ')

selected = option_menu(None, ["Главная", "Исследование", 'Настройки'], 
         icons=['house-door', 'graph-up','wrench-adjustable'], 
         menu_icon="cast", default_index=0, orientation="horizontal",
         styles={
             "container": {"padding": "0!important", "background-color": "#24769C"},
             "icon": {"color": "#5DAED3", "font-size": "18px"}, 
             "nav-link": {"font-size": "18px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
             "nav-link-selected": {"background-color": "#335D70"},
         })

##########Главная
if selected == "Главная":
   
   col1, col2 = st.columns([0.66, 0.34])

   with col1:
        file_land = codecs.open("заставка.html","r",encoding="utf8")
        page_land = file_land.read()
        components.html(page_land,width=None, height=550, scrolling=False)

   file = codecs.open("Главный_текст.html","r",encoding="utf8")
   page = file.read()
   components.html(page,width=None, height=600, scrolling=False)

##########Исследование

if selected == "Исследование":
   ####### Левое боковое меню
   st.sidebar.title('Меню')

   st.sidebar.subheader('Какое исследование проводится?')

   option = st.sidebar.selectbox('Выберите вид исследования',
       ('Изучение абсолютной и относительной биодоступности препарата', 'Изучение фармакокинетики в органах животных', 'Линейность дозирования','Изучение экскреции препарата'),disabled = False, key = "Вид исследования")

   ############### файл пример

   df_example_file = pd.read_excel("server_example_file.xlsx")
   df_example_file_xlsx = to_excel(df_example_file)
   st.sidebar.download_button(label='Пример файла 💾', data=df_example_file_xlsx , file_name= 'example_file.xlsx')

   ############ памятка

   text_contents = '''1)Оглавлять колонку с номерами животных должно слово «Номер» (в верхнем регистре).
   2)Знак «№» обязательно должен присутствовать при указании номера животного, иначе приложение выдаст ошибку. 
   3) Не ставить в ячейки знак «-» в случае нулевого значения. Ставить число «0» для корректной работы приложения.
   4)Ни в каком исследовании загружаемые файлы не должны называться одинаково.
   '''
   st.sidebar.download_button('Памятка заполнения 📄', text_contents)
   
   ################################

   if option == 'Изучение абсолютной и относительной биодоступности препарата':
       
       st.title('Изучение абсолютной и относительной биодоступности препарата')

       col1, col2 = st.columns([0.66, 0.34])
       
       ######### боковое меню справа
       with col2:
            selected = option_menu(None, ["Включение параметров в исследование"], 
            icons=['menu-button'], 
            menu_icon="cast", default_index=0, orientation="vertical",
            styles={
                "container": {"padding": "0!important", "background-color": "#24769C"},
                "icon": {"color": "#5DAED3", "font-size": "13px"}, 
                "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
                "nav-link-selected": {"background-color": "#335D70"},
            })

            if selected == "Включение параметров в исследование":
               type_parameter = st.selectbox('Выберите параметр',
            ('Cmax(2)',"Вид введения"),disabled = False, key = "Вид параметра - ИБ")
               

            if type_parameter == 'Cmax(2)':
               
               if "agree_cmax2 - ИБ" not in st.session_state:
                  st.session_state["agree_cmax2 - ИБ"] = False

               st.session_state["agree_cmax2 - ИБ"] = st.checkbox('Добавить возможность выбора Cmax(2)', key = "Возможность добавления Cmax2 - ИБ", value = st.session_state["agree_cmax2 - ИБ"])
               
               if st.session_state["agree_cmax2 - ИБ"] == True:
                  st.write('🧠Параметр добавлен!')

       ####### основной экран
       with col1:
           
           panel = st.radio(
               "⚙️Панель управления",
               ("Загрузка файлов", "Таблицы","Графики"),
               horizontal=True, key= "Загрузка файлов - Изучение абсолютной и относительной биодоступности препарата"
           )

           ###создание состояния
           if "measure_unit" not in st.session_state:
              st.session_state["measure_unit"] = ""
           if "dose_iv" not in st.session_state:
              st.session_state["dose_iv"] = ""
           if "dose_po_sub" not in st.session_state:   
              st.session_state["dose_po_sub"] = ""
           if "dose_po_rdf" not in st.session_state:   
              st.session_state["dose_po_rdf"] = ""
              
           #cписки для word-отчета
           list_heading_word=[]
           list_table_word=[]
           list_graphics_word=[]
           list_heading_graphics_word=[]

           if panel == "Загрузка файлов":
          
              measure_unit = st.text_input("Введите единицы измерения концентрации", key='Единицы измерения при изучении абсолютной и относительной биодоступности препарата', value = st.session_state["measure_unit"])
              
              st.session_state["measure_unit"] = measure_unit

              #cостояние радио-кнопки "method_auc"
              if "index_method_auc - ИБ" not in st.session_state:
                  st.session_state["index_method_auc - ИБ"] = 0

              method_auc = st.radio("📌Метод подсчёта AUC0-t",('linear',"linear-up/log-down"),key = "Метод подсчёта AUC0-t - ИБ", index = st.session_state["index_method_auc - ИБ"])
              
              if st.session_state["Метод подсчёта AUC0-t - ИБ"] == 'linear':
                 st.session_state["index_method_auc - ИБ"] = 0
              if st.session_state["Метод подсчёта AUC0-t - ИБ"] == "linear-up/log-down":
                 st.session_state["index_method_auc - ИБ"] = 1

              st.title('Внутривенное введение субстанции')
              
              uploaded_file_1 = st.file_uploader("Выбрать файл внутривенного введения (формат XLSX)", key='Файл внутривенного введения при изучении абсолютной и относительной биодоступности препарата')
              
              #сохранение файла
              if uploaded_file_1 is not None:
                 save_uploadedfile(uploaded_file_1)
                 st.session_state["uploaded_file_1"] = uploaded_file_1.name
                 
              dose_iv = st.text_input("Доза при внутривенном введении", key='Доза при внутривенном введении при изучении абсолютной и относительной биодоступности препарата', value = st.session_state["dose_iv"])
              
              st.session_state["dose_iv"] = dose_iv

              if "uploaded_file_1" in st.session_state and dose_iv and measure_unit:
                 df = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_1"]))
                 st.subheader('Индивидуальные значения концентраций в крови после внутривенного введения субстанции')
                 
                 ###интерактивная таблица
                 df = edit_frame(df,st.session_state["uploaded_file_1"])

                 ###количество животных 
                 count_rows_number_iv= len(df.axes[0])
                
                 ################

                 table_heading='Индивидуальные и усредненные значения концентраций в крови после внутривенного введения субстанции'
                 list_heading_word.append(table_heading)

                 ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                 df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']

                 list_table_word.append(df_concat_round_str_transpose)
              ########### графики    

              ######индивидуальные    

                 # в линейных координатах
                 col_mapping = df.columns.tolist()
                 col_mapping.remove('Номер')

                 count_row_df = len(df.axes[0])

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 list_time.remove(0) ###т.к. внутривенное

                 for r in range(0,count_row_df):

                     list_concentration=df.iloc[r].tolist()

                     numer_animal=list_concentration[0]

                     list_concentration.pop(0) #удаление номера животного

                     list_concentration = [float(v) for v in list_concentration]

                     list_concentration.remove(0) ###т.к. внутривенное

                     fig, ax = plt.subplots()
                     plt.plot(list_time,list_concentration,marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)
                     
                     list_graphics_word.append(fig) 
                     
                     #переобъявляем переменную названия графика
                     graphic='График индивидуального фармакокинетического профиля в крови (в линейных координатах) после внутривенного введения субстанции,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

                  #в полулогарифмических координатах методом удаления точек
                     count_for_0_1=len(list_concentration)
                     list_range_for_0_1=range(0,count_for_0_1)

                     list_time_0=[]
                     list_for_log_1=[]
                     for i in list_range_for_0_1:
                         if list_concentration[i] !=0:
                            list_for_log_1.append(list_concentration[i])
                            list_time_0.append(list_time[i]) 

                     fig, ax = plt.subplots()
                     plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     ax.set_yscale("log")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)
                     
                     
                     list_graphics_word.append(fig)
                     
                     graphic='График индивидуального фармакокинетического профиля в крови (в полулогарифмических координатах) после внутривенного введения субстанции,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

              # объединенные индивидуальные в линейных координатах

                 df_for_plot_conc=df.drop(['Номер'], axis=1)
                 df_for_plot_conc_1 = df_for_plot_conc.transpose()

                 df_for_plot_conc_1=df_for_plot_conc_1.replace(0, None) ###т.к. внутривенное

                 list_numer_animal_for_plot=df['Номер'].tolist()
                 count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды

                 list_color = [] ## генерация 500 цветов
                 for i in range(0,500):
                     hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                     list_color.append(hexadecimal)
                 
                 fig, ax = plt.subplots()
                
                 ax.set_prop_cycle(cycler(color=list_color))
                
                 plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)
                 
                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))
                    
                 list_graphics_word.append(fig)

                 graphic="Сравнение индивидуальных фармакокинетических профилей (в линейных координатах) после внутривенного введения субстанции"
                 list_heading_graphics_word.append(graphic)    
              # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                 df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)
                 

                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 ax.set_yscale("log")
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))
                 
                 list_graphics_word.append(fig)

                 graphic="Сравнение индивидуальных фармакокинетических профилей (в полулогарифмических координатах) после внутривенного введения субстанции"
                 list_heading_graphics_word.append(graphic)
                  ###усредненные    
              # в линейных координатах
                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 df_averaged_concentrations=df.describe()
                 list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                 err_y_1=df_averaged_concentrations.loc['std'].tolist()
                 
                 list_time.remove(0) ###т.к. внутривенное
                 list_concentration.remove(0)
                 err_y_1.remove(0) 
                 
                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)
                 
                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в линейных координатах) после внутривенного введения субстанции'
                 list_heading_graphics_word.append(graphic)



              #в полулогарифмических координатах
                 #для полулогарифм. построим без нуля (ноль уже удален)


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 ax.set_yscale("log")
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)

                 
                 list_graphics_word.append(fig)

                 graphic='График усредненного фармакокинетического профиля в крови (в полулогарифмических координатах) после внутривенного введения субстанции'
                 list_heading_graphics_word.append(graphic)


                 ############ Параметры ФК

                 df_without_numer=df.drop(['Номер'],axis=1)
                 count_row=df_without_numer.shape[0]

                 list_count_row=range(count_row)
       
                 ###Cmax
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                 if st.session_state["agree_cmax2 - ИБ"] == False:
                    list_cmax_1_iv=[]
                    for i in range(0,count_row):
                        cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                        list_cmax_1_iv.append(cmax)
                 
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    ###создание состояния
                    if "selected_value_iv" not in st.session_state:
                       st.session_state["selected_value_iv"] = []
                    
                    if "feature_disable_selected_value_iv" not in st.session_state:
                        st.session_state["feature_disable_selected_value_iv"] = True

                    ###создание состояния
                    st.info('Выбери Cmax:')
                    list_columns_without_numer = df.columns.tolist()
                    list_columns_without_numer.remove('Номер')
                    selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax внутривенного введения субстанции',max_selections=1)
                    st.session_state["selected_columns_iv"] = selected_columns 

                    list_keys_cmax = st.session_state["selected_value_iv"]
                    if selected_columns != [] and st.session_state["feature_disable_selected_value_iv"]:
                       selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax внутривенного введения субстанции',max_selections=1)
                       list_keys_cmax.append(selected_value)

                    if list_keys_cmax != []:
                       st.session_state["selected_value_iv"] = list_keys_cmax

                    list_keys_cmax = st.session_state["selected_value_iv"]
                    list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]

                    if st.button('Очистить список Cmax', key="Очистка списка Cmax внутривенного введения субстанции"):
                       del st.session_state["selected_value_iv"]
                       list_keys_cmax_sample = []
                       selected_columns = st.session_state["selected_columns_iv"]
                       st.session_state["feature_disable_selected_value_iv"] = True
                                           
                    st.write("Список Cmax:")
                    st.write(list_keys_cmax_sample)
                    

                    list_cmax_1_iv=list_keys_cmax_sample 
                    
                    list_cmax_2_iv=[]

                 if (len(list_cmax_1_iv) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    st.session_state["feature_disable_selected_value_iv"] = False

                    ######Cmax2

                    if "feature_disable_selected_value_iv_2" not in st.session_state:
                     st.session_state["feature_disable_selected_value_iv_2"] = True

                    st.info('Выбери Cmax(2):')
                    
                    selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 внутривенного введения субстанции', max_selections=1)
                    st.session_state["selected_columns_2_iv"] = selected_columns_2

                    ###создание состояния
                    if "selected_value_2_iv" not in st.session_state:
                       st.session_state["selected_value_2_iv"] = []

                    list_keys_cmax_2 = st.session_state["selected_value_2_iv"]
                    if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_iv_2"]:
                       selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 внутривенного введения субстанции', max_selections=1)
                       list_keys_cmax_2.append(selected_value_2)

                    if list_keys_cmax_2 != []:
                       st.session_state["selected_value_2_iv"] = list_keys_cmax_2

                    list_keys_cmax_2 = st.session_state["selected_value_2_iv"]
                    list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                    if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) внутривенного введения субстанции"):
                       del st.session_state["selected_value_2_iv"]
                       list_keys_cmax_sample_2 = []
                       selected_columns_2 = st.session_state["selected_columns_2_iv"]
                       st.session_state["feature_disable_selected_value_iv_2"] = True

                    st.write("Список Cmax(2):")
                    st.write(list_keys_cmax_sample_2)

                    list_cmax_2_iv= list_keys_cmax_sample_2

                    if len(list_cmax_2_iv) == len(df.index.tolist()):
                       st.session_state["feature_disable_selected_value_iv_2"] = False

                 if (len(list_cmax_1_iv) == len(df.index.tolist())):
                    
                    ###Tmax   
                    list_Tmax_1=[]
                    for cmax in list_cmax_1_iv:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_1.append(f"{column}")
                  
                    list_Tmax_float_1=[]           
                    for i in list_Tmax_1:
                        Tmax=float(i)
                        list_Tmax_float_1.append(Tmax)

                 if (len(list_cmax_1_iv) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    
                    list_Tmax_2=[]
                    for cmax in list_cmax_2_iv:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_2.append(f"{column}")
                  
                    list_Tmax_float_2=[]           
                    for i in list_Tmax_2:
                        Tmax=float(i)
                        list_Tmax_float_2.append(Tmax)  

                 if (len(list_cmax_1_iv) == len(df.index.tolist())):
                    
                    ###AUC0-t
                    list_AUC_0_T=[]
                    if method_auc == 'linear':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                           list_AUC_0_T.append(AUC_0_T)

                    if method_auc == 'linear-up/log-down':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################
                           
                           list_c = list_concentration
                           list_t = list_columns_T
                           
                           count_i = len(list_c)
                           list_range= range(0,count_i)
                           
                           list_AUC_0_T_ascending=[]
                           list_AUC_0_T_descending = []
                           AUC_0_T_ascending=0
                           AUC_0_T_descending = 0
                           a=0
                           a1=0
                           d=0
                           d1=0
                           for i in list_range:
                               if a1<count_i-1:
                                  if list_c[i+1] > list_c[i]:
                                     if a<count_i-1:
                                         AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                         a+=1
                                         list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                               if d1<count_i-1:
                                  if list_c[i+1] < list_c[i]:      
                                     if d<count_i-1:
                                         AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                         d+=1
                                         list_AUC_0_T_descending.append(AUC_0_T_descending)
                                  a1+=1
                                  d1+=1
            
                           AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                           
                           list_AUC_0_T.append(AUC_O_T)

                    ####Сmax/AUC0-t
                    list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_iv,list_AUC_0_T)
                    list_Сmax_division_AUC0_t=[]
                    for i,j in list_Сmax_division_AUC0_t_for_division:
                            list_Сmax_division_AUC0_t.append(i/j)


                    ####KEL
                    list_kel_total=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c=list_concentration

                        list_time=df_without_numer.columns.tolist()
                        list_time.remove(0) 

                        list_t=[]
                        for i in list_time:
                            i=float(i)
                            list_t.append(i)

                        #срез_без_cmax
                        max_value_c=max(list_c)
                        index_cmax=list_c.index(max_value_c)

                        list_c_without_cmax=list_c[index_cmax+1:]
                        list_t_without_cmax=list_t[index_cmax+1:]

                        #удаление всех нулей из массивов
                        count_for_0_1=len(list_c_without_cmax)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_conc_0=[]
                        for i in list_range_for_0_1:
                            if list_c_without_cmax[i] !=0:
                               list_conc_0.append(list_c_without_cmax[i])
                               list_time_0.append(list_t_without_cmax[i]) 
                        ################################

                        n_points=len(list_conc_0)
                        list_n_points = range(0,n_points)

                        #создание списков с поочередно уменьщающемся кол, точек
                        list_for_kel_c=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_c_new=list_conc_0[j:n_points]
                               list_for_kel_c.append(list_c_new)
                        list_for_kel_c.pop(-1) #удаление списка с одной точкой
                        list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                        list_for_kel_t=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_t_new=list_time_0[j:n_points]
                               list_for_kel_t.append(list_t_new)
                        list_for_kel_t.pop(-1) #удаление списка с одной точкой
                        list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                        list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                        list_kel=[]
                        list_r=[]
                        for i,j in list_ct_zip:

                            n_points_r=len(i)

                            np_c=np.asarray(i)
                            np_t_1=np.asarray(j).reshape((-1,1))

                            np_c_log=np.log(np_c)

                            model = LinearRegression().fit(np_t_1,np_c_log)

                            np_t=np.asarray(j)
                            a=np.corrcoef(np_t, np_c_log)
                            cor=((a[0])[1])
                            r_sq=cor**2

                            adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                            ########################################
                            kel=abs(model.coef_[0])
                            list_kel.append(kel)
                            list_r.append(adjusted_r_sq)

                        #делаем срезы списоков до rmax
                        max_r=max(list_r)

                        index_max_r= list_r.index(max_r)

                        list_r1=list_r
                        list_kel1=list_kel

                        number_elem_list_r1=len(list_r1)

                        list_range_kel=range(0,number_elem_list_r1) 

                        list_kel_total_1=[]
                        for i in list_range_kel:

                            if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                               list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                               break #самая ранняя удовлетовряющая условию

                        for i in list_kel_total_1:
                            list_kel_total.append(i) 


                    ####T1/2
                    list_half_live=[]
                    for i in list_kel_total:
                        half_live=math.log(2)/i
                        list_half_live.append(half_live)


                    ###AUC0-inf 

                    list_auc0_inf=[] 

                    list_of_list_c=[]
                    for i in range(0,count_row):
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c = list_concentration
                        list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                        list_of_list_c.append(list_c)

                    list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                        #AUCt-inf 
                    list_auc_t_inf=[]     
                    for i,j in list_zip_c_AUCt_inf:
                        for clast in j:
                            if clast != 0:
                               clast_true=clast
                               break
                        auc_t_inf=clast_true/i
                        list_auc_t_inf.append(auc_t_inf)

                    list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                    for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                        auc0_inf=i+j    
                        list_auc0_inf.append(auc0_inf)


                    ####CL
                    list_cl=[]

                    for i in list_auc0_inf:
                        cl = float(dose_iv)/i * 1000
                        list_cl.append(cl)


                    ####Vd
                    list_Vd=[]

                    list_zip_kel_cl=zip(list_kel_total,list_cl)

                    for i,j in list_zip_kel_cl:
                        Vd = j/i
                        list_Vd.append(Vd)


                    ###AUMC
                    list_AUMCO_inf=[]

                    list_AUMC0_t=[]

                    list_C_last=[]
                    list_T_last=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                        ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                        cmax = max(list_concentration)
                        index_cmax = list_concentration.index(cmax)
                        list_before_cmax = list_concentration[0:index_cmax]
                        list_after_cmax = list_concentration[index_cmax:]
                        list_before_cmax_t = list_columns_T[0:index_cmax]
                        list_after_cmax_t = list_columns_T[index_cmax:]

                        count_list_concentration = len(list_after_cmax)
                        list_range_for_remove_0 = range(0,count_list_concentration)

                        list_conc_without_0=[]
                        list_t_without_0=[]
                        for i in list_range_for_remove_0:
                            if list_after_cmax[i] !=0:
                               list_conc_without_0.append(list_after_cmax[i])
                               list_t_without_0.append(list_after_cmax_t[i])

                        list_concentration = list_before_cmax + list_conc_without_0
                        list_columns_T = list_before_cmax_t + list_t_without_0
                        ######################

                        list_C_last.append(list_concentration[-1]) 
                        list_T_last.append(list_columns_T[-1]) 

                        list_len=len(list_concentration)

                        list_aumc_i=[]
                        for i in range(0,list_len):
                            AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                            list_aumc_i.append(AUMC)

                        list_aumc_i.pop(0)

                        a=0
                        list_AUMC0_t_1=[]
                        for i in list_aumc_i:
                            a+=i
                            list_AUMC0_t_1.append(a)
                        list_AUMC0_t.append(list_AUMC0_t_1[-1])

                    list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                    list_AUMCt_inf=[]
                    for k,c,t in list_zip_for_AUMC_inf:
                        AUMCt_inf=c*t/k+c/(k*k)
                        list_AUMCt_inf.append(AUMCt_inf)


                    list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                    for i,j in list_AUMC_zip:
                        AUMCO_inf=i+j
                        list_AUMCO_inf.append(AUMCO_inf)

                    ###MRT0-inf
                    list_MRT0_inf=[]

                    list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                    for i,j in list_zip_AUMCO_inf_auc0_inf:
                        MRT0_inf=i/j
                        list_MRT0_inf.append(MRT0_inf)
                 
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    #####Cmax условие для дальнейшего кода
                    if len(list_cmax_1_iv) == len(df.index.tolist()) and len(list_cmax_2_iv) == len(df.index.tolist()):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_iv,list_Tmax_float_1,list_cmax_2_iv,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 
                 
                 if len(list_cmax_1_iv) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False):
                    
                    ##################### Фрейм ФК параметров

                    ### пользовательский индекс
                    list_for_index=df["Номер"].tolist()
                    df_PK=pd.DataFrame(list(zip(list_cmax_1_iv,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index)
                 
                 checking_condition_cmax2 = False

                 if st.session_state["agree_cmax2 - ИБ"] == True:
                     
                    checking_condition_cmax2 = len(list_cmax_1_iv) == len(df.index.tolist()) and len(list_cmax_2_iv) == len(df.index.tolist()) and st.session_state["agree_cmax2 - ИБ"] == True


                 if checking_condition_cmax2 or (len(list_cmax_1_iv) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False)):
                    
                    ###описательная статистика

                    col_mapping_PK = df_PK.columns.tolist()

                    list_gmean_PK=[]

                    list_cv_PK=[] 

                    for i in col_mapping_PK:

                        list_ser_PK=df_PK[i].tolist()

                        def g_mean(list_ser_PK):
                            a=np.log(list_ser_PK)
                            return np.exp(a.mean())
                        Gmean_PK=g_mean(list_ser_PK)
                        list_gmean_PK.append(Gmean_PK)

                        cv_std_PK=lambda x: np.std(x, ddof= 1 )
                        cv_mean_PK=lambda x: np.mean(x)

                        CV_std_PK=cv_std_PK(list_ser_PK)
                        CV_mean_PK=cv_mean_PK(list_ser_PK)

                        CV_PK=(CV_std_PK/CV_mean_PK * 100)
                        list_cv_PK.append(CV_PK)


                    df_averaged_concentrations_PK=df_PK.describe()
                    df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                    df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                    df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                    df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                    df_round_without_CV_PK=df_averaged_3_PK
                    df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                    df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                    df_concat_PK_iv= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)

                    ###округление описательной статистики и ФК параметров
                    
                    series_Cmax=df_concat_PK_iv['Cmax']
                    list_Cmax_str_f=["%.2f" % round(v,2) for v in series_Cmax.tolist()]
                    series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_iv.index.tolist(), name='Cmax ' +"("+measure_unit+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Cmax_2=df_concat_PK_iv['Cmax(2)']
                       list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                       series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_iv.index.tolist(), name='Cmax(2) ' +"("+measure_unit+")")

                    series_Tmax=df_concat_PK_iv['Tmax']
                    list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                    series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_iv.index.tolist(), name='Tmax ' +"("+"ч"+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Tmax_2=df_concat_PK_iv['Tmax(2)']
                       list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                       series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_iv.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                    series_MRT0_inf= df_concat_PK_iv['MRT0→∞']
                    list_MRT0_inf_str_f=["%.3f" % round(v,3) for v in series_MRT0_inf.tolist()]
                    series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_iv.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                    series_half_live= df_concat_PK_iv['T1/2']
                    list_half_live_str_f=["%.2f" % round(v,2) for v in series_half_live.tolist()]
                    series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_iv.index.tolist(), name='T1/2 '+"("+"ч"+")")

                    series_AUC0_t= df_concat_PK_iv['AUC0-t']
                    list_AUC0_t_str_f=["%.2f" % round(v,2) for v in series_AUC0_t.tolist()]
                    series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_iv.index.tolist(), name='AUC0-t '+"("+measure_unit+"×ч" +")")

                    series_AUC0_inf= df_concat_PK_iv['AUC0→∞']
                    list_AUC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUC0_inf.tolist()]
                    series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_iv.index.tolist(), name='AUC0→∞ '+"("+measure_unit+"×ч" +")")

                    series_AUMC0_inf= df_concat_PK_iv['AUMC0-∞']
                    list_AUMC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUMC0_inf.tolist()]
                    series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_iv.index.tolist(), name='AUMC0-∞ '+"("+measure_unit+"×ч\u00B2" +")")

                    series_Сmax_dev_AUC0_t= df_concat_PK_iv['Сmax/AUC0-t']
                    list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                    series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_iv.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                    series_Kel= df_concat_PK_iv['Kel']
                    list_Kel_str_f=["%.4f" % round(v,4) for v in series_Kel.tolist()]
                    series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_iv.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                    series_CL= df_concat_PK_iv['CL/F']
                    list_CL_str_f=["%.2f" % round(v,2) for v in series_CL.tolist()]
                    series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_iv.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                    series_Vd= df_concat_PK_iv['Vd']
                    list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                    series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_iv.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")

                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       df_total_PK_iv = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    else:
                       df_total_PK_iv = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1)

                    df_total_PK_iv.index.name = 'Номер'

                    ##изменение названий параметров описательной статистики

                    df_total_PK_iv1=df_total_PK_iv.copy()
                    df_total_PK_iv1.iloc[-6,:],df_total_PK_iv1.iloc[-2,:]=df_total_PK_iv.iloc[-2,:],df_total_PK_iv.iloc[-6,:]

                    df_total_PK_iv=df_total_PK_iv1

                    df_total_PK_iv1=df_total_PK_iv.copy()
                    df_total_PK_iv1.iloc[-4,:],df_total_PK_iv1.iloc[-5,:]=df_total_PK_iv.iloc[-5,:],df_total_PK_iv.iloc[-4,:]

                    df_total_PK_iv=df_total_PK_iv1

                    df_total_PK_iv = df_total_PK_iv.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')

                    table_heading='Фармакокинетические показатели в крови после внутривенного введения субстанции'
                    list_heading_word.append(table_heading)
                    
                    list_table_word.append(df_total_PK_iv)

                    ####получение интервала для средних ФК параметров
                    list_PK_Cmax_1_not_round = df_PK['Cmax'].tolist()
                    list_PK_Tmax_1_not_round = df_PK['Tmax'].tolist() 
                    list_PK_MRT0_inf_not_round = df_PK['MRT0→∞'].tolist() 
                    list_PK_half_live_not_round = df_PK['T1/2'].tolist() 
                    list_PK_AUC0_t_not_round = df_PK['AUC0-t'].tolist()
                    list_PK_AUC0_inf_not_round = df_PK['AUC0→∞'].tolist()
                    list_PK_AUMC0_inf_not_round = df_PK['AUMC0-∞'].tolist()
                    list_PK_Сmax_dev_AUC0_t_not_round = df_PK['Сmax/AUC0-t'].tolist()
                    list_PK_Kel_not_round = df_PK['Kel'].tolist()

                    list_list_PK_parametr_iv=[list_PK_Cmax_1_not_round,list_PK_AUC0_t_not_round,list_PK_Kel_not_round,list_PK_AUC0_inf_not_round,list_PK_half_live_not_round,list_PK_AUMC0_inf_not_round,list_PK_MRT0_inf_not_round,list_PK_Сmax_dev_AUC0_t_not_round]
                    list_parametr_mean_h_iv=[]
                    for i in list_list_PK_parametr_iv:
                         n=len(i)

                         def confidential_interval(i):
                             if n < 30:
                                h = statistics.stdev(i)
                                mean = np.mean(i)
                             else:
                                h = statistics.stdev(i)  ### прояснить момент с n-1
                                mean = np.mean(i)
                             return ([mean,h]) 
                         func_mean_h = confidential_interval(i)

                         list_parametr_mean_h_iv.append(func_mean_h)

                    list_mean_h_iv_Cmax_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_iv[0]]
                    parametr_round_mean_h_Cmax=str(list_mean_h_iv_Cmax_round[0]) +"±"+str(list_mean_h_iv_Cmax_round[1])

                    list_mean_h_iv_AUC0_t_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_iv[1]] 
                    parametr_round_mean_h_AUC0_t=str(list_mean_h_iv_AUC0_t_round[0]) +"±"+str(list_mean_h_iv_AUC0_t_round[1]) 

                    list_mean_h_iv_Kel_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_iv[2]]
                    parametr_round_mean_h_Kel=str(list_mean_h_iv_Kel_round[0]) +"±"+str(list_mean_h_iv_Kel_round[1])

                    list_mean_h_iv_AUC0_inf_round= ["%.2f" % round(v,2) for v in list_parametr_mean_h_iv[3]]
                    parametr_round_mean_h_AUC0_inf=str(list_mean_h_iv_AUC0_inf_round[0]) +"±"+str(list_mean_h_iv_AUC0_inf_round[1]) 

                    list_mean_h_iv_half_live_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_iv[4]]
                    parametr_round_mean_h_half_live=str(list_mean_h_iv_half_live_round[0]) +"±"+str(list_mean_h_iv_half_live_round[1])

                    list_mean_h_iv_AUMC0_inf_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_iv[5]] 
                    parametr_round_mean_h_AUMC0_inf=str(list_mean_h_iv_AUMC0_inf_round[0]) +"±"+str(list_mean_h_iv_AUMC0_inf_round[1]) 

                    list_mean_h_iv_MRT0_inf_round=["%.3f" % round(v,3) for v in list_parametr_mean_h_iv[6]]
                    parametr_round_mean_h_MRT0_inf=str(list_mean_h_iv_MRT0_inf_round[0]) +"±"+str(list_mean_h_iv_MRT0_inf_round[1])

                    list_mean_h_iv_Сmax_dev_AUC0_t_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_iv[7]]
                    parametr_round_mean_h_Сmax_dev_AUC0_t=str(list_mean_h_iv_Сmax_dev_AUC0_t_round[0]) +"±"+str(list_mean_h_iv_Сmax_dev_AUC0_t_round[1])

                    list_parametr_round_mean_h_iv= [parametr_round_mean_h_Cmax,parametr_round_mean_h_AUC0_t,parametr_round_mean_h_Kel,parametr_round_mean_h_AUC0_inf,parametr_round_mean_h_half_live,parametr_round_mean_h_AUMC0_inf,parametr_round_mean_h_MRT0_inf,parametr_round_mean_h_Сmax_dev_AUC0_t]

                    t_mean_iv = str("%.2f" % round(np.mean(list_PK_Tmax_1_not_round),2))     
                    list_parametr_round_mean_h_iv.insert(1,t_mean_iv)

              ############################################################################################################### 
              st.title('Пероральное введение субстанции')
              
              uploaded_file_2 = st.file_uploader("Выбрать файл перорального введения субстанции (формат XLSX)", key='Файл перорального введения субстанции при изучении абсолютной и относительной биодоступности препарата')
              
              #сохранение файла
              if uploaded_file_2 is not None:
                 save_uploadedfile(uploaded_file_2)
                 st.session_state["uploaded_file_2"] = uploaded_file_2.name

              dose_po_sub = st.text_input("Доза при пероральном введении субстанции", key='Доза при пероральном введении субстанции при изучении абсолютной и относительной биодоступности препарата', value = st.session_state["dose_po_sub"])
              
              st.session_state["dose_po_sub"] = dose_po_sub

              if "uploaded_file_2" in st.session_state and dose_po_sub and measure_unit:

                 df = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_2"]))
                 st.subheader('Индивидуальные значения концентраций в крови после перорального введения субстанции')
                 
                 ###интерактивная таблица
                 df = edit_frame(df,st.session_state["uploaded_file_2"])

                 ###количество животных 
                 count_rows_number_sub= len(df.axes[0])
           
                 table_heading='Индивидуальные и усредненные значения концентраций в крови после перорального введения субстанции'
                 list_heading_word.append(table_heading)

                 ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                 df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']
                 
                 list_table_word.append(df_concat_round_str_transpose)

              ########### графики    

              ######индивидуальные    

                 # в линейных координатах
                 col_mapping = df.columns.tolist()
                 col_mapping.remove('Номер')

                 count_row_df = len(df.axes[0])

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 for r in range(0,count_row_df):

                     list_concentration=df.iloc[r].tolist()

                     numer_animal=list_concentration[0]

                     list_concentration.pop(0) #удаление номера животного

                     list_concentration = [float(v) for v in list_concentration]


                     fig, ax = plt.subplots()
                     plt.plot(list_time,list_concentration,marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)
                    
                     list_graphics_word.append(fig)  

                     graphic='График индивидуального фармакокинетического профиля в крови (в линейных координатах) после перорального введения субстанции,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

                  #в полулогарифмических координатах методом удаления точек
                     count_for_0_1=len(list_concentration)
                     list_range_for_0_1=range(0,count_for_0_1)

                     list_time_0=[]
                     list_for_log_1=[]
                     for i in list_range_for_0_1:
                         if list_concentration[i] !=0:
                            list_for_log_1.append(list_concentration[i])
                            list_time_0.append(list_time[i]) 

                     fig, ax = plt.subplots()
                     plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     ax.set_yscale("log")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)

                     list_graphics_word.append(fig) 

                     graphic='График индивидуального фармакокинетического профиля в крови (в полулогарифмических координатах) после перорального введения субстанции,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

              # объединенные индивидуальные в линейных координатах

                 df_for_plot_conc=df.drop(['Номер'], axis=1)
                 df_for_plot_conc_1 = df_for_plot_conc.transpose()
                 list_numer_animal_for_plot=df['Номер'].tolist()
                 count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды
                 list_color = [] ## генерация 500 цветов
                 for i in range(0,500):
                     hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                     list_color.append(hexadecimal)

                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 

                 graphic="Сравнение индивидуальных фармакокинетических профилей (в линейных координатах) после перорального введения субстанции"
                 list_heading_graphics_word.append(graphic)    
              # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                 df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)


                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 ax.set_yscale("log")
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 
          
                 graphic="Сравнение индивидуальных фармакокинетических профилей (в полулогарифмических координатах) после перорального введения субстанции"
                 list_heading_graphics_word.append(graphic) 

              ### усреденные    
              #в линейных    

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 df_averaged_concentrations=df.describe()
                 list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                 err_y_2=df_averaged_concentrations.loc['std'].tolist()


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_2, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в линейных координатах) после перорального введения субстанции'
                 list_heading_graphics_word.append(graphic)  

              #в полулогарифмических координатах
                 list_time.remove(0)
                 list_concentration.remove(0)
                 err_y_2.remove(0) 


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_2, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 ax.set_yscale("log")
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в полулогарифмических координатах) после перорального введения субстанции'
                 list_heading_graphics_word.append(graphic) 

                 ############ Параметры ФК

                 df_without_numer=df.drop(['Номер'],axis=1)
                 count_row=df_without_numer.shape[0]

                 list_count_row=range(count_row)
       
                 ###Cmax
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                 if st.session_state["agree_cmax2 - ИБ"] == False:
                    list_cmax_1_sub=[]
                    for i in range(0,count_row):
                        cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                        list_cmax_1_sub.append(cmax)
                 
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    ###создание состояния
                    if "selected_value_sub" not in st.session_state:
                       st.session_state["selected_value_sub"] = []
                    
                    if "feature_disable_selected_value_sub" not in st.session_state:
                        st.session_state["feature_disable_selected_value_sub"] = True

                    ###создание состояния
                    st.info('Выбери Cmax:')
                    list_columns_without_numer = df.columns.tolist()
                    list_columns_without_numer.remove('Номер')
                    selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax перорального введения субстанции',max_selections=1)
                    st.session_state["selected_columns_sub"] = selected_columns 

                    list_keys_cmax = st.session_state["selected_value_sub"]
                    if selected_columns != [] and st.session_state["feature_disable_selected_value_sub"]:
                       selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax перорального введения субстанции',max_selections=1)
                       list_keys_cmax.append(selected_value)

                    if list_keys_cmax != []:
                       st.session_state["selected_value_sub"] = list_keys_cmax

                    list_keys_cmax = st.session_state["selected_value_sub"]
                    list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]

                    if st.button('Очистить список Cmax', key="Очистка списка Cmax перорального введения субстанции"):
                       del st.session_state["selected_value_sub"]
                       list_keys_cmax_sample = []
                       selected_columns = st.session_state["selected_columns_sub"]
                       st.session_state["feature_disable_selected_value_sub"] = True
                                           
                    st.write("Список Cmax:")
                    st.write(list_keys_cmax_sample)
                    

                    list_cmax_1_sub=list_keys_cmax_sample 
                    
                    list_cmax_2_sub=[]

                 if len(list_cmax_1_sub) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    st.session_state["feature_disable_selected_value_sub"] = False

                    ######Cmax2

                    if "feature_disable_selected_value_sub_2" not in st.session_state:
                     st.session_state["feature_disable_selected_value_sub_2"] = True

                    st.info('Выбери Cmax(2):')
                    
                    selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 перорального введения субстанции', max_selections=1)
                    st.session_state["selected_columns_2_sub"] = selected_columns_2

                    ###создание состояния
                    if "selected_value_2_sub" not in st.session_state:
                       st.session_state["selected_value_2_sub"] = []

                    list_keys_cmax_2 = st.session_state["selected_value_2_sub"]
                    if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_sub_2"]:
                       selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 перорального введения субстанции', max_selections=1)
                       list_keys_cmax_2.append(selected_value_2)

                    if list_keys_cmax_2 != []:
                       st.session_state["selected_value_2_sub"] = list_keys_cmax_2

                    list_keys_cmax_2 = st.session_state["selected_value_2_sub"]
                    list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                    if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) перорального введения субстанции"):
                       del st.session_state["selected_value_2_sub"]
                       list_keys_cmax_sample_2 = []
                       selected_columns_2 = st.session_state["selected_columns_2_sub"]
                       st.session_state["feature_disable_selected_value_sub_2"] = True

                    st.write("Список Cmax(2):")
                    st.write(list_keys_cmax_sample_2)

                    list_cmax_2_sub= list_keys_cmax_sample_2

                    if len(list_cmax_2_sub) == len(df.index.tolist()):
                       st.session_state["feature_disable_selected_value_sub_2"] = False

                 if (len(list_cmax_1_sub) == len(df.index.tolist())):
                    
                    ###Tmax   
                    list_Tmax_1=[]
                    for cmax in list_cmax_1_sub:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_1.append(f"{column}")
                  
                    list_Tmax_float_1=[]           
                    for i in list_Tmax_1:
                        Tmax=float(i)
                        list_Tmax_float_1.append(Tmax)

                 if (len(list_cmax_1_sub) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    
                    list_Tmax_2=[]
                    for cmax in list_cmax_2_sub:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_2.append(f"{column}")
                  
                    list_Tmax_float_2=[]           
                    for i in list_Tmax_2:
                        Tmax=float(i)
                        list_Tmax_float_2.append(Tmax)  

                 if (len(list_cmax_1_sub) == len(df.index.tolist())):
                    
                    ###AUC0-t
                    list_AUC_0_T=[]
                    if method_auc == 'linear':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                           list_AUC_0_T.append(AUC_0_T)

                    if method_auc == 'linear-up/log-down':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################
                           
                           list_c = list_concentration
                           list_t = list_columns_T
                           
                           count_i = len(list_c)
                           list_range= range(0,count_i)
                           
                           list_AUC_0_T_ascending=[]
                           list_AUC_0_T_descending = []
                           AUC_0_T_ascending=0
                           AUC_0_T_descending = 0
                           a=0
                           a1=0
                           d=0
                           d1=0
                           for i in list_range:
                               if a1<count_i-1:
                                  if list_c[i+1] > list_c[i]:
                                     if a<count_i-1:
                                         AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                         a+=1
                                         list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                               if d1<count_i-1:
                                  if list_c[i+1] < list_c[i]:      
                                     if d<count_i-1:
                                         AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                         d+=1
                                         list_AUC_0_T_descending.append(AUC_0_T_descending)
                                  a1+=1
                                  d1+=1
            
                           AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                           
                           list_AUC_0_T.append(AUC_O_T)

                    ####Сmax/AUC0-t
                    list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_sub,list_AUC_0_T)
                    list_Сmax_division_AUC0_t=[]
                    for i,j in list_Сmax_division_AUC0_t_for_division:
                            list_Сmax_division_AUC0_t.append(i/j)


                    ####KEL
                    list_kel_total=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c=list_concentration

                        list_time=df_without_numer.columns.tolist()
                        list_time.remove(0) 

                        list_t=[]
                        for i in list_time:
                            i=float(i)
                            list_t.append(i)

                        #срез_без_cmax
                        max_value_c=max(list_c)
                        index_cmax=list_c.index(max_value_c)

                        list_c_without_cmax=list_c[index_cmax+1:]
                        list_t_without_cmax=list_t[index_cmax+1:]

                        #удаление всех нулей из массивов
                        count_for_0_1=len(list_c_without_cmax)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_conc_0=[]
                        for i in list_range_for_0_1:
                            if list_c_without_cmax[i] !=0:
                               list_conc_0.append(list_c_without_cmax[i])
                               list_time_0.append(list_t_without_cmax[i]) 
                        ################################

                        n_points=len(list_conc_0)
                        list_n_points = range(0,n_points)

                        #создание списков с поочередно уменьщающемся кол, точек
                        list_for_kel_c=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_c_new=list_conc_0[j:n_points]
                               list_for_kel_c.append(list_c_new)
                        list_for_kel_c.pop(-1) #удаление списка с одной точкой
                        list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                        list_for_kel_t=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_t_new=list_time_0[j:n_points]
                               list_for_kel_t.append(list_t_new)
                        list_for_kel_t.pop(-1) #удаление списка с одной точкой
                        list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                        list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                        list_kel=[]
                        list_r=[]
                        for i,j in list_ct_zip:

                            n_points_r=len(i)

                            np_c=np.asarray(i)
                            np_t_1=np.asarray(j).reshape((-1,1))

                            np_c_log=np.log(np_c)

                            model = LinearRegression().fit(np_t_1,np_c_log)

                            np_t=np.asarray(j)
                            a=np.corrcoef(np_t, np_c_log)
                            cor=((a[0])[1])
                            r_sq=cor**2

                            adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                            ########################################
                            kel=abs(model.coef_[0])
                            list_kel.append(kel)
                            list_r.append(adjusted_r_sq)

                        #делаем срезы списоков до rmax
                        max_r=max(list_r)

                        index_max_r= list_r.index(max_r)

                        list_r1=list_r
                        list_kel1=list_kel

                        number_elem_list_r1=len(list_r1)

                        list_range_kel=range(0,number_elem_list_r1) 

                        list_kel_total_1=[]
                        for i in list_range_kel:

                            if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                               list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                               break #самая ранняя удовлетовряющая условию

                        for i in list_kel_total_1:
                            list_kel_total.append(i) 


                    ####T1/2
                    list_half_live=[]
                    for i in list_kel_total:
                        half_live=math.log(2)/i
                        list_half_live.append(half_live)


                    ###AUC0-inf 

                    list_auc0_inf=[] 

                    list_of_list_c=[]
                    for i in range(0,count_row):
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c = list_concentration
                        list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                        list_of_list_c.append(list_c)

                    list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                        #AUCt-inf 
                    list_auc_t_inf=[]     
                    for i,j in list_zip_c_AUCt_inf:
                        for clast in j:
                            if clast != 0:
                               clast_true=clast
                               break
                        auc_t_inf=clast_true/i
                        list_auc_t_inf.append(auc_t_inf)

                    list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                    for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                        auc0_inf=i+j    
                        list_auc0_inf.append(auc0_inf)


                    ####CL
                    list_cl=[]

                    for i in list_auc0_inf:
                        cl = float(dose_po_sub)/i * 1000
                        list_cl.append(cl)


                    ####Vd
                    list_Vd=[]

                    list_zip_kel_cl=zip(list_kel_total,list_cl)

                    for i,j in list_zip_kel_cl:
                        Vd = j/i
                        list_Vd.append(Vd)


                    ###AUMC
                    list_AUMCO_inf=[]

                    list_AUMC0_t=[]

                    list_C_last=[]
                    list_T_last=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                        ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                        cmax = max(list_concentration)
                        index_cmax = list_concentration.index(cmax)
                        list_before_cmax = list_concentration[0:index_cmax]
                        list_after_cmax = list_concentration[index_cmax:]
                        list_before_cmax_t = list_columns_T[0:index_cmax]
                        list_after_cmax_t = list_columns_T[index_cmax:]

                        count_list_concentration = len(list_after_cmax)
                        list_range_for_remove_0 = range(0,count_list_concentration)

                        list_conc_without_0=[]
                        list_t_without_0=[]
                        for i in list_range_for_remove_0:
                            if list_after_cmax[i] !=0:
                               list_conc_without_0.append(list_after_cmax[i])
                               list_t_without_0.append(list_after_cmax_t[i])

                        list_concentration = list_before_cmax + list_conc_without_0
                        list_columns_T = list_before_cmax_t + list_t_without_0
                        ######################

                        list_C_last.append(list_concentration[-1]) 
                        list_T_last.append(list_columns_T[-1]) 

                        list_len=len(list_concentration)

                        list_aumc_i=[]
                        for i in range(0,list_len):
                            AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                            list_aumc_i.append(AUMC)

                        list_aumc_i.pop(0)

                        a=0
                        list_AUMC0_t_1=[]
                        for i in list_aumc_i:
                            a+=i
                            list_AUMC0_t_1.append(a)
                        list_AUMC0_t.append(list_AUMC0_t_1[-1])

                    list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                    list_AUMCt_inf=[]
                    for k,c,t in list_zip_for_AUMC_inf:
                        AUMCt_inf=c*t/k+c/(k*k)
                        list_AUMCt_inf.append(AUMCt_inf)


                    list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                    for i,j in list_AUMC_zip:
                        AUMCO_inf=i+j
                        list_AUMCO_inf.append(AUMCO_inf)

                    ###MRT0-inf
                    list_MRT0_inf=[]

                    list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                    for i,j in list_zip_AUMCO_inf_auc0_inf:
                        MRT0_inf=i/j
                        list_MRT0_inf.append(MRT0_inf)
                 
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    #####Cmax условие для дальнейшего кода
                    if len(list_cmax_1_sub) == len(df.index.tolist()) and len(list_cmax_2_sub) == len(df.index.tolist()):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_sub,list_Tmax_float_1,list_cmax_2_sub,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 

                 if len(list_cmax_1_sub) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False):
                    
                    ##################### Фрейм ФК параметров

                    ### пользовательский индекс
                    list_for_index=df["Номер"].tolist()
                    df_PK=pd.DataFrame(list(zip(list_cmax_1_sub,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index)
                 
                 checking_condition_cmax2 = False

                 if st.session_state["agree_cmax2 - ИБ"] == True:
                     
                    checking_condition_cmax2 = len(list_cmax_1_sub) == len(df.index.tolist()) and len(list_cmax_2_sub) == len(df.index.tolist()) and st.session_state["agree_cmax2 - ИБ"] == True

                 if checking_condition_cmax2 or (len(list_cmax_1_sub) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False)):
                 
                    ###описательная статистика

                    col_mapping_PK = df_PK.columns.tolist()

                    list_gmean_PK=[]

                    list_cv_PK=[] 

                    for i in col_mapping_PK:

                        list_ser_PK=df_PK[i].tolist()

                        def g_mean(list_ser_PK):
                            a=np.log(list_ser_PK)
                            return np.exp(a.mean())
                        Gmean_PK=g_mean(list_ser_PK)
                        list_gmean_PK.append(Gmean_PK)

                        cv_std_PK=lambda x: np.std(x, ddof= 1 )
                        cv_mean_PK=lambda x: np.mean(x)

                        CV_std_PK=cv_std_PK(list_ser_PK)
                        CV_mean_PK=cv_mean_PK(list_ser_PK)

                        CV_PK=(CV_std_PK/CV_mean_PK * 100)
                        list_cv_PK.append(CV_PK)


                    df_averaged_concentrations_PK=df_PK.describe()
                    df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                    df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                    df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                    df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                    df_round_without_CV_PK=df_averaged_3_PK
                    df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                    df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                    df_concat_PK_po_sub= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)

                    ###округление описательной статистики и ФК параметров

                    series_Cmax=df_concat_PK_po_sub['Cmax']
                    list_Cmax_str_f=["%.2f" % round(v,2) for v in series_Cmax.tolist()]
                    series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_po_sub.index.tolist(), name='Cmax ' +"("+measure_unit+")")

                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Cmax_2=df_concat_PK_po_sub['Cmax(2)']
                       list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                       series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_po_sub.index.tolist(), name='Cmax(2) ' +"("+measure_unit+")")

                    series_Tmax=df_concat_PK_po_sub['Tmax']
                    list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                    series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_po_sub.index.tolist(), name='Tmax ' +"("+"ч"+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Tmax_2=df_concat_PK_po_sub['Tmax(2)']
                       list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                       series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_po_sub.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                    series_MRT0_inf= df_concat_PK_po_sub['MRT0→∞']
                    list_MRT0_inf_str_f=["%.3f" % round(v,3) for v in series_MRT0_inf.tolist()]
                    series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_po_sub.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                    series_half_live= df_concat_PK_po_sub['T1/2']
                    list_half_live_str_f=["%.2f" % round(v,2) for v in series_half_live.tolist()]
                    series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_po_sub.index.tolist(), name='T1/2 '+"("+"ч"+")")

                    series_AUC0_t= df_concat_PK_po_sub['AUC0-t']
                    list_AUC0_t_str_f=["%.2f" % round(v,2) for v in series_AUC0_t.tolist()]
                    series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_po_sub.index.tolist(), name='AUC0-t '+"("+measure_unit+"×ч" +")")

                    series_AUC0_inf= df_concat_PK_po_sub['AUC0→∞']
                    list_AUC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUC0_inf.tolist()]
                    series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_po_sub.index.tolist(), name='AUC0→∞ '+"("+measure_unit+"×ч" +")")

                    series_AUMC0_inf= df_concat_PK_po_sub['AUMC0-∞']
                    list_AUMC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUMC0_inf.tolist()]
                    series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_po_sub.index.tolist(), name='AUMC0-∞ '+"("+measure_unit+"×ч\u00B2" +")")

                    series_Сmax_dev_AUC0_t= df_concat_PK_po_sub['Сmax/AUC0-t']
                    list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                    series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_po_sub.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                    series_Kel= df_concat_PK_po_sub['Kel']
                    list_Kel_str_f=["%.4f" % round(v,4) for v in series_Kel.tolist()]
                    series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_po_sub.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                    series_CL= df_concat_PK_po_sub['CL/F']
                    list_CL_str_f=["%.2f" % round(v,2) for v in series_CL.tolist()]
                    series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_po_sub.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                    series_Vd= df_concat_PK_po_sub['Vd']
                    list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                    series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_po_sub.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       df_total_PK_po_sub = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    else:
                       df_total_PK_po_sub = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    
                    df_total_PK_po_sub.index.name = 'Номер'

                    ##изменение названий параметров описательной статистики

                    df_total_PK_po_sub1=df_total_PK_po_sub.copy()
                    df_total_PK_po_sub1.iloc[-6,:],df_total_PK_po_sub1.iloc[-2,:]=df_total_PK_po_sub.iloc[-2,:],df_total_PK_po_sub.iloc[-6,:]

                    df_total_PK_po_sub=df_total_PK_po_sub1

                    df_total_PK_po_sub1=df_total_PK_po_sub.copy()
                    df_total_PK_po_sub1.iloc[-4,:],df_total_PK_po_sub1.iloc[-5,:]=df_total_PK_po_sub.iloc[-5,:],df_total_PK_po_sub.iloc[-4,:]

                    df_total_PK_po_sub=df_total_PK_po_sub1

                    df_total_PK_po_sub = df_total_PK_po_sub.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')

                    table_heading='Фармакокинетические показатели в крови после перорального введения субстанции'
                    list_heading_word.append(table_heading)
                    
                    list_table_word.append(df_total_PK_po_sub)

                    ####получение интервала для средних ФК параметров
                    list_PK_Cmax_1_not_round = df_PK['Cmax'].tolist()
                    list_PK_Tmax_1_not_round = df_PK['Tmax'].tolist() 
                    list_PK_MRT0_inf_not_round = df_PK['MRT0→∞'].tolist() 
                    list_PK_half_live_not_round = df_PK['T1/2'].tolist() 
                    list_PK_AUC0_t_not_round = df_PK['AUC0-t'].tolist()
                    list_PK_AUC0_inf_not_round = df_PK['AUC0→∞'].tolist()
                    list_PK_AUMC0_inf_not_round = df_PK['AUMC0-∞'].tolist()
                    list_PK_Сmax_dev_AUC0_t_not_round = df_PK['Сmax/AUC0-t'].tolist()
                    list_PK_Kel_not_round = df_PK['Kel'].tolist()

                    list_list_PK_parametr_po_sub=[list_PK_Cmax_1_not_round,list_PK_AUC0_t_not_round,list_PK_Kel_not_round,list_PK_AUC0_inf_not_round,list_PK_half_live_not_round,list_PK_AUMC0_inf_not_round,list_PK_MRT0_inf_not_round,list_PK_Сmax_dev_AUC0_t_not_round]
                    list_parametr_mean_h_po_sub=[]
                    for i in list_list_PK_parametr_po_sub:
                         n=len(i)

                         def confidential_interval(i):
                             if n < 30:
                                h = statistics.stdev(i)
                                mean = np.mean(i)
                             else:
                                h = statistics.stdev(i)  ### прояснить момент с n-1
                                mean = np.mean(i)
                             return ([mean,h]) 
                         func_mean_h = confidential_interval(i)

                         list_parametr_mean_h_po_sub.append(func_mean_h)

                    list_mean_h_po_sub_Cmax_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_sub[0]]
                    parametr_round_mean_h_Cmax=str(list_mean_h_po_sub_Cmax_round[0]) +"±"+str(list_mean_h_po_sub_Cmax_round[1])

                    list_mean_h_po_sub_AUC0_t_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_sub[1]] 
                    parametr_round_mean_h_AUC0_t=str(list_mean_h_po_sub_AUC0_t_round[0]) +"±"+str(list_mean_h_po_sub_AUC0_t_round[1]) 

                    list_mean_h_po_sub_Kel_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_po_sub[2]]
                    parametr_round_mean_h_Kel=str(list_mean_h_po_sub_Kel_round[0]) +"±"+str(list_mean_h_po_sub_Kel_round[1])

                    list_mean_h_po_sub_AUC0_inf_round= ["%.2f" % round(v,2) for v in list_parametr_mean_h_po_sub[3]]
                    parametr_round_mean_h_AUC0_inf=str(list_mean_h_po_sub_AUC0_inf_round[0]) +"±"+str(list_mean_h_po_sub_AUC0_inf_round[1]) 

                    list_mean_h_po_sub_half_live_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_sub[4]]
                    parametr_round_mean_h_half_live=str(list_mean_h_po_sub_half_live_round[0]) +"±"+str(list_mean_h_po_sub_half_live_round[1])

                    list_mean_h_po_sub_AUMC0_inf_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_sub[5]] 
                    parametr_round_mean_h_AUMC0_inf=str(list_mean_h_po_sub_AUMC0_inf_round[0]) +"±"+str(list_mean_h_po_sub_AUMC0_inf_round[1]) 

                    list_mean_h_po_sub_MRT0_inf_round=["%.3f" % round(v,3) for v in list_parametr_mean_h_po_sub[6]]
                    parametr_round_mean_h_MRT0_inf=str(list_mean_h_po_sub_MRT0_inf_round[0]) +"±"+str(list_mean_h_po_sub_MRT0_inf_round[1])

                    list_mean_h_po_sub_Сmax_dev_AUC0_t_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_po_sub[7]]
                    parametr_round_mean_h_Сmax_dev_AUC0_t=str(list_mean_h_po_sub_Сmax_dev_AUC0_t_round[0]) +"±"+str(list_mean_h_po_sub_Сmax_dev_AUC0_t_round[1])

                    list_parametr_round_mean_h_po_sub= [parametr_round_mean_h_Cmax,parametr_round_mean_h_AUC0_t,parametr_round_mean_h_Kel,parametr_round_mean_h_AUC0_inf,parametr_round_mean_h_half_live,parametr_round_mean_h_AUMC0_inf,parametr_round_mean_h_MRT0_inf,parametr_round_mean_h_Сmax_dev_AUC0_t]

                    t_mean_po_sub = str("%.2f" % round(np.mean(list_PK_Tmax_1_not_round),2))     
                    list_parametr_round_mean_h_po_sub.insert(1,t_mean_po_sub)

              ##############################################################################################################

              st.title('Пероральное введение ГЛФ')
              
              uploaded_file_3 = st.file_uploader("Выбрать файл перорального введения ГЛФ (формат XLSX)", key='Файл перорального введения ГЛФ при изучении абсолютной и относительной биодоступности препарата')
              
              #сохранение файла
              if uploaded_file_3 is not None:
                 save_uploadedfile(uploaded_file_3)
                 st.session_state["uploaded_file_3"] = uploaded_file_3.name

              dose_po_rdf = st.text_input("Доза при пероральном введении ГЛФ", key='Доза при пероральном введении ГЛФ при изучении абсолютной и относительной биодоступности препарата', value = st.session_state["dose_po_rdf"])
              
              st.session_state["dose_po_rdf"] = dose_po_rdf

              if "uploaded_file_3" in st.session_state and dose_po_rdf and measure_unit:

                 df = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_3"]))
                 st.subheader('Индивидуальные значения концентраций в крови после перорального введения ГЛФ')
                 
                 ###интерактивная таблица
                 df = edit_frame(df,st.session_state["uploaded_file_3"])

                 ###количество животных 
                 count_rows_number_rdf= len(df.axes[0])
           
                 table_heading='Индивидуальные и усредненные значения концентраций в крови после перорального введения ГЛФ'
                 list_heading_word.append(table_heading)

                 ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                 df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']
                 
                 list_table_word.append(df_concat_round_str_transpose)

              ########### графики    

              ######индивидуальные    

                 # в линейных координатах
                 col_mapping = df.columns.tolist()
                 col_mapping.remove('Номер')

                 count_row_df = len(df.axes[0])

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 for r in range(0,count_row_df):

                     list_concentration=df.iloc[r].tolist()

                     numer_animal=list_concentration[0]

                     list_concentration.pop(0) #удаление номера животного

                     list_concentration = [float(v) for v in list_concentration]


                     fig, ax = plt.subplots()
                     plt.plot(list_time,list_concentration,marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)
                    
                     list_graphics_word.append(fig)  

                     graphic='График индивидуального фармакокинетического профиля в крови (в линейных координатах) после перорального введения ГЛФ,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

                  #в полулогарифмических координатах методом удаления точек
                     count_for_0_1=len(list_concentration)
                     list_range_for_0_1=range(0,count_for_0_1)

                     list_time_0=[]
                     list_for_log_1=[]
                     for i in list_range_for_0_1:
                         if list_concentration[i] !=0:
                            list_for_log_1.append(list_concentration[i])
                            list_time_0.append(list_time[i]) 

                     fig, ax = plt.subplots()
                     plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                     ax.set_yscale("log")
                     plt.xlabel("Время, ч")
                     plt.ylabel("Концентрация, "+measure_unit)

                     list_graphics_word.append(fig) 

                     graphic='График индивидуального фармакокинетического профиля в крови (в полулогарифмических координатах) после перорального введения ГЛФ,  '+numer_animal
                     list_heading_graphics_word.append(graphic)

              # объединенные индивидуальные в линейных координатах

                 df_for_plot_conc=df.drop(['Номер'], axis=1)
                 df_for_plot_conc_1 = df_for_plot_conc.transpose()
                 list_numer_animal_for_plot=df['Номер'].tolist()
                 count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды
                 list_color = [] ## генерация 500 цветов
                 for i in range(0,500):
                     hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                     list_color.append(hexadecimal)

                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 

                 graphic="Сравнение индивидуальных фармакокинетических профилей (в линейных координатах) после перорального введения ГЛФ"
                 list_heading_graphics_word.append(graphic)    
              # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                 df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)


                 fig, ax = plt.subplots()
                 
                 ax.set_prop_cycle(cycler(color=list_color))

                 plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                 ax.set_xlabel("Время, ч")
                 ax.set_ylabel("Концентрация, "+measure_unit)
                 ax.set_yscale("log")
                 if count_numer_animal > 20:
                    ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                 else:
                    ax.legend(bbox_to_anchor=(1, 1))

                 list_graphics_word.append(fig) 
          
                 graphic="Сравнение индивидуальных фармакокинетических профилей (в полулогарифмических координатах) после перорального введения ГЛФ"
                 list_heading_graphics_word.append(graphic) 

              ### усреденные    
              #в линейных    

                 list_time = []
                 for i in col_mapping:
                     numer=float(i)
                     list_time.append(numer)

                 df_averaged_concentrations=df.describe()
                 list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                 err_y_2=df_averaged_concentrations.loc['std'].tolist()


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_2, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в линейных координатах) после перорального введения ГЛФ'
                 list_heading_graphics_word.append(graphic)  

              #в полулогарифмических координатах
                 list_time.remove(0)
                 list_concentration.remove(0)
                 err_y_2.remove(0) 


                 fig, ax = plt.subplots()
                 plt.errorbar(list_time,list_concentration,yerr=err_y_2, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                 ax.set_yscale("log")
                 plt.xlabel("Время, ч")
                 plt.ylabel("Концентрация, "+measure_unit)

                 list_graphics_word.append(fig) 

                 graphic='График усредненного фармакокинетического профиля в крови (в полулогарифмических координатах) после перорального введения ГЛФ'
                 list_heading_graphics_word.append(graphic) 

                 ############### Параметры ФК

                 df_without_numer=df.drop(['Номер'],axis=1)
                 count_row=df_without_numer.shape[0]

                 list_count_row=range(count_row)
                 
                 ###Cmax
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                 if st.session_state["agree_cmax2 - ИБ"] == False:
                    list_cmax_1_rdf=[]
                    for i in range(0,count_row):
                        cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                        list_cmax_1_rdf.append(cmax)
                 
                 #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                 if st.session_state["agree_cmax2 - ИБ"] == True:

                    ###создание состояния
                    if "selected_value_rdf" not in st.session_state:
                       st.session_state["selected_value_rdf"] = []
                    
                    if "feature_disable_selected_value_rdf" not in st.session_state:
                        st.session_state["feature_disable_selected_value_rdf"] = True

                    ###создание состояния
                    st.info('Выбери Cmax:')
                    list_columns_without_numer = df.columns.tolist()
                    list_columns_without_numer.remove('Номер')
                    selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax перорального введения ГЛФ',max_selections=1)
                    st.session_state["selected_columns_rdf"] = selected_columns 

                    list_keys_cmax = st.session_state["selected_value_rdf"]
                    if selected_columns != [] and st.session_state["feature_disable_selected_value_rdf"]:
                       selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax перорального введения ГЛФ',max_selections=1)
                       list_keys_cmax.append(selected_value)

                    if list_keys_cmax != []:
                       st.session_state["selected_value_rdf"] = list_keys_cmax

                    list_keys_cmax = st.session_state["selected_value_rdf"]
                    list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]

                    if st.button('Очистить список Cmax', key="Очистка списка Cmax перорального введения ГЛФ"):
                       del st.session_state["selected_value_rdf"]
                       list_keys_cmax_sample = []
                       selected_columns = st.session_state["selected_columns_rdf"]
                       st.session_state["feature_disable_selected_value_rdf"] = True
                                           
                    st.write("Список Cmax:")
                    st.write(list_keys_cmax_sample)
                    

                    list_cmax_1_rdf=list_keys_cmax_sample 

                    list_cmax_2_rdf=[]

                 if len(list_cmax_1_rdf) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    st.session_state["feature_disable_selected_value_rdf"] = False

                    ######Cmax2

                    if "feature_disable_selected_value_2" not in st.session_state:
                     st.session_state["feature_disable_selected_value_rdf_2"] = True

                    st.info('Выбери Cmax(2):')
                    
                    selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 перорального введения ГЛФ', max_selections=1)
                    st.session_state["selected_columns_2_rdf"] = selected_columns_2

                    ###создание состояния
                    if "selected_value_2_rdf" not in st.session_state:
                       st.session_state["selected_value_2_rdf"] = []

                    list_keys_cmax_2 = st.session_state["selected_value_2_rdf"]
                    if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_rdf_2"]:
                       selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 перорального введения ГЛФ', max_selections=1)
                       list_keys_cmax_2.append(selected_value_2)

                    if list_keys_cmax_2 != []:
                       st.session_state["selected_value_2_rdf"] = list_keys_cmax_2

                    list_keys_cmax_2 = st.session_state["selected_value_2_rdf"]
                    list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                    if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) перорального введения ГЛФ"):
                       del st.session_state["selected_value_2_rdf"]
                       list_keys_cmax_sample_2 = []
                       selected_columns_2 = st.session_state["selected_columns_2_rdf"]
                       st.session_state["feature_disable_selected_value_rdf_2"] = True

                    st.write("Список Cmax(2):")
                    st.write(list_keys_cmax_sample_2)

                    list_cmax_2_rdf= list_keys_cmax_sample_2

                    if len(list_cmax_2_rdf) == len(df.index.tolist()):
                       st.session_state["feature_disable_selected_value_rdf_2"] = False
                 
                 if (len(list_cmax_1_rdf) == len(df.index.tolist())):
                    
                    ###Tmax   
                    list_Tmax_1=[]
                    for cmax in list_cmax_1_rdf:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_1.append(f"{column}")
                  
                    list_Tmax_float_1=[]           
                    for i in list_Tmax_1:
                        Tmax=float(i)
                        list_Tmax_float_1.append(Tmax)

                 if (len(list_cmax_1_rdf) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - ИБ"] == True):
                    
                    list_Tmax_2=[]
                    for cmax in list_cmax_2_rdf:
                        for column in df.columns:
                            for num, row in df.iterrows():
                                if df.iloc[num][column] == cmax:
                                   list_Tmax_2.append(f"{column}")
                  
                    list_Tmax_float_2=[]           
                    for i in list_Tmax_2:
                        Tmax=float(i)
                        list_Tmax_float_2.append(Tmax)  

                 if (len(list_cmax_1_rdf) == len(df.index.tolist())):
                    
                    ###AUC0-t
                    list_AUC_0_T=[]
                    if method_auc == 'linear':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                           list_AUC_0_T.append(AUC_0_T)

                    if method_auc == 'linear-up/log-down':
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################
                           
                           list_c = list_concentration
                           list_t = list_columns_T
                           
                           count_i = len(list_c)
                           list_range= range(0,count_i)
                           
                           list_AUC_0_T_ascending=[]
                           list_AUC_0_T_descending = []
                           AUC_0_T_ascending=0
                           AUC_0_T_descending = 0
                           a=0
                           a1=0
                           d=0
                           d1=0
                           for i in list_range:
                               if a1<count_i-1:
                                  if list_c[i+1] > list_c[i]:
                                     if a<count_i-1:
                                         AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                         a+=1
                                         list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                               if d1<count_i-1:
                                  if list_c[i+1] < list_c[i]:      
                                     if d<count_i-1:
                                         AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                         d+=1
                                         list_AUC_0_T_descending.append(AUC_0_T_descending)
                                  a1+=1
                                  d1+=1
            
                           AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                           
                           list_AUC_0_T.append(AUC_O_T)

                    ####Сmax/AUC0-t
                    list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_rdf,list_AUC_0_T)
                    list_Сmax_division_AUC0_t=[]
                    for i,j in list_Сmax_division_AUC0_t_for_division:
                            list_Сmax_division_AUC0_t.append(i/j)


                    ####KEL
                    list_kel_total=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c=list_concentration

                        list_time=df_without_numer.columns.tolist()
                        list_time.remove(0) 

                        list_t=[]
                        for i in list_time:
                            i=float(i)
                            list_t.append(i)

                        #срез_без_cmax
                        max_value_c=max(list_c)
                        index_cmax=list_c.index(max_value_c)

                        list_c_without_cmax=list_c[index_cmax+1:]
                        list_t_without_cmax=list_t[index_cmax+1:]

                        #удаление всех нулей из массивов
                        count_for_0_1=len(list_c_without_cmax)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_conc_0=[]
                        for i in list_range_for_0_1:
                            if list_c_without_cmax[i] !=0:
                               list_conc_0.append(list_c_without_cmax[i])
                               list_time_0.append(list_t_without_cmax[i]) 
                        ################################

                        n_points=len(list_conc_0)
                        list_n_points = range(0,n_points)

                        #создание списков с поочередно уменьщающемся кол, точек
                        list_for_kel_c=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_c_new=list_conc_0[j:n_points]
                               list_for_kel_c.append(list_c_new)
                        list_for_kel_c.pop(-1) #удаление списка с одной точкой
                        list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                        list_for_kel_t=[]
                        for j in list_n_points:
                            if j<n_points:
                               list_t_new=list_time_0[j:n_points]
                               list_for_kel_t.append(list_t_new)
                        list_for_kel_t.pop(-1) #удаление списка с одной точкой
                        list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                        list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                        list_kel=[]
                        list_r=[]
                        for i,j in list_ct_zip:

                            n_points_r=len(i)

                            np_c=np.asarray(i)
                            np_t_1=np.asarray(j).reshape((-1,1))

                            np_c_log=np.log(np_c)

                            model = LinearRegression().fit(np_t_1,np_c_log)

                            np_t=np.asarray(j)
                            a=np.corrcoef(np_t, np_c_log)
                            cor=((a[0])[1])
                            r_sq=cor**2

                            adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                            ########################################
                            kel=abs(model.coef_[0])
                            list_kel.append(kel)
                            list_r.append(adjusted_r_sq)

                        #делаем срезы списоков до rmax
                        max_r=max(list_r)

                        index_max_r= list_r.index(max_r)

                        list_r1=list_r
                        list_kel1=list_kel

                        number_elem_list_r1=len(list_r1)

                        list_range_kel=range(0,number_elem_list_r1) 

                        list_kel_total_1=[]
                        for i in list_range_kel:

                            if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                               list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                               break #самая ранняя удовлетовряющая условию

                        for i in list_kel_total_1:
                            list_kel_total.append(i) 


                    ####T1/2
                    list_half_live=[]
                    for i in list_kel_total:
                        half_live=math.log(2)/i
                        list_half_live.append(half_live)


                    ###AUC0-inf 

                    list_auc0_inf=[] 

                    list_of_list_c=[]
                    for i in range(0,count_row):
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                        list_concentration.remove(0)
                        list_c = list_concentration
                        list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                        list_of_list_c.append(list_c)

                    list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                        #AUCt-inf 
                    list_auc_t_inf=[]     
                    for i,j in list_zip_c_AUCt_inf:
                        for clast in j:
                            if clast != 0:
                               clast_true=clast
                               break
                        auc_t_inf=clast_true/i
                        list_auc_t_inf.append(auc_t_inf)

                    list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                    for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                        auc0_inf=i+j    
                        list_auc0_inf.append(auc0_inf)


                    ####CL
                    list_cl=[]

                    for i in list_auc0_inf:
                        cl = float(dose_po_rdf)/i * 1000
                        list_cl.append(cl)


                    ####Vd
                    list_Vd=[]

                    list_zip_kel_cl=zip(list_kel_total,list_cl)

                    for i,j in list_zip_kel_cl:
                        Vd = j/i
                        list_Vd.append(Vd)


                    ###AUMC
                    list_AUMCO_inf=[]

                    list_AUMC0_t=[]

                    list_C_last=[]
                    list_T_last=[]
                    for i in range(0,count_row):
                        list_columns_T=[]
                        for column in df_without_numer.columns:
                            list_columns_T.append(float(column))
                        list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                        ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                        cmax = max(list_concentration)
                        index_cmax = list_concentration.index(cmax)
                        list_before_cmax = list_concentration[0:index_cmax]
                        list_after_cmax = list_concentration[index_cmax:]
                        list_before_cmax_t = list_columns_T[0:index_cmax]
                        list_after_cmax_t = list_columns_T[index_cmax:]

                        count_list_concentration = len(list_after_cmax)
                        list_range_for_remove_0 = range(0,count_list_concentration)

                        list_conc_without_0=[]
                        list_t_without_0=[]
                        for i in list_range_for_remove_0:
                            if list_after_cmax[i] !=0:
                               list_conc_without_0.append(list_after_cmax[i])
                               list_t_without_0.append(list_after_cmax_t[i])

                        list_concentration = list_before_cmax + list_conc_without_0
                        list_columns_T = list_before_cmax_t + list_t_without_0
                        ######################

                        list_C_last.append(list_concentration[-1]) 
                        list_T_last.append(list_columns_T[-1]) 

                        list_len=len(list_concentration)

                        list_aumc_i=[]
                        for i in range(0,list_len):
                            AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                            list_aumc_i.append(AUMC)

                        list_aumc_i.pop(0)

                        a=0
                        list_AUMC0_t_1=[]
                        for i in list_aumc_i:
                            a+=i
                            list_AUMC0_t_1.append(a)
                        list_AUMC0_t.append(list_AUMC0_t_1[-1])

                    list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                    list_AUMCt_inf=[]
                    for k,c,t in list_zip_for_AUMC_inf:
                        AUMCt_inf=c*t/k+c/(k*k)
                        list_AUMCt_inf.append(AUMCt_inf)


                    list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                    for i,j in list_AUMC_zip:
                        AUMCO_inf=i+j
                        list_AUMCO_inf.append(AUMCO_inf)

                    ###MRT0-inf
                    list_MRT0_inf=[]

                    list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                    for i,j in list_zip_AUMCO_inf_auc0_inf:
                        MRT0_inf=i/j
                        list_MRT0_inf.append(MRT0_inf)
                 
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    #####Cmax условие для дальнейшего кода
                    if len(list_cmax_1_rdf) == len(df.index.tolist()) and len(list_cmax_2_rdf) == len(df.index.tolist()):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_rdf,list_Tmax_float_1,list_cmax_2_rdf,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 
                 
                 if len(list_cmax_1_rdf) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False):
                    
                    ##################### Фрейм ФК параметров

                    ### пользовательский индекс
                    list_for_index=df["Номер"].tolist()
                    df_PK=pd.DataFrame(list(zip(list_cmax_1_rdf,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 
                 
                 checking_condition_cmax2 = False

                 if st.session_state["agree_cmax2 - ИБ"] == True:
                     
                    checking_condition_cmax2 = len(list_cmax_1_rdf) == len(df.index.tolist()) and len(list_cmax_2_rdf) == len(df.index.tolist()) and st.session_state["agree_cmax2 - ИБ"] == True

                 if checking_condition_cmax2 or (len(list_cmax_1_rdf) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - ИБ"] == False)):
                    
                    ###описательная статистика

                    col_mapping_PK = df_PK.columns.tolist()

                    list_gmean_PK=[]

                    list_cv_PK=[] 

                    for i in col_mapping_PK:

                        list_ser_PK=df_PK[i].tolist()

                        def g_mean(list_ser_PK):
                            a=np.log(list_ser_PK)
                            return np.exp(a.mean())
                        Gmean_PK=g_mean(list_ser_PK)
                        list_gmean_PK.append(Gmean_PK)

                        cv_std_PK=lambda x: np.std(x, ddof= 1 )
                        cv_mean_PK=lambda x: np.mean(x)

                        CV_std_PK=cv_std_PK(list_ser_PK)
                        CV_mean_PK=cv_mean_PK(list_ser_PK)

                        CV_PK=(CV_std_PK/CV_mean_PK * 100)
                        list_cv_PK.append(CV_PK)


                    df_averaged_concentrations_PK=df_PK.describe()
                    df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                    df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                    df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                    df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                    df_round_without_CV_PK=df_averaged_3_PK
                    df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                    df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                    df_concat_PK_po_rdf= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)

                    ###округление описательной статистики и ФК параметров

                    series_Cmax=df_concat_PK_po_rdf['Cmax']
                    list_Cmax_str_f=["%.2f" % round(v,2) for v in series_Cmax.tolist()]
                    series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='Cmax ' +"("+measure_unit+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Cmax_2=df_concat_PK_po_rdf['Cmax(2)']
                       list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                       series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_po_rdf.index.tolist(), name='Cmax(2) ' +"("+measure_unit+")")

                    series_Tmax=df_concat_PK_po_rdf['Tmax']
                    list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                    series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='Tmax ' +"("+"ч"+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       series_Tmax_2=df_concat_PK_po_rdf['Tmax(2)']
                       list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                       series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_po_rdf.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                    series_MRT0_inf= df_concat_PK_po_rdf['MRT0→∞']
                    list_MRT0_inf_str_f=["%.3f" % round(v,3) for v in series_MRT0_inf.tolist()]
                    series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                    series_half_live= df_concat_PK_po_rdf['T1/2']
                    list_half_live_str_f=["%.2f" % round(v,2) for v in series_half_live.tolist()]
                    series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='T1/2 '+"("+"ч"+")")

                    series_AUC0_t= df_concat_PK_po_rdf['AUC0-t']
                    list_AUC0_t_str_f=["%.2f" % round(v,2) for v in series_AUC0_t.tolist()]
                    series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='AUC0-t '+"("+measure_unit+"×ч" +")")

                    series_AUC0_inf= df_concat_PK_po_rdf['AUC0→∞']
                    list_AUC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUC0_inf.tolist()]
                    series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='AUC0→∞ '+"("+measure_unit+"×ч" +")")

                    series_AUMC0_inf= df_concat_PK_po_rdf['AUMC0-∞']
                    list_AUMC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUMC0_inf.tolist()]
                    series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='AUMC0-∞ '+"("+measure_unit+"×ч\u00B2" +")")

                    series_Сmax_dev_AUC0_t= df_concat_PK_po_rdf['Сmax/AUC0-t']
                    list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                    series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                    series_Kel= df_concat_PK_po_rdf['Kel']
                    list_Kel_str_f=["%.4f" % round(v,4) for v in series_Kel.tolist()]
                    series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                    series_CL= df_concat_PK_po_rdf['CL/F']
                    list_CL_str_f=["%.2f" % round(v,2) for v in series_CL.tolist()]
                    series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                    series_Vd= df_concat_PK_po_rdf['Vd']
                    list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                    series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_po_rdf.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")
                    
                    if st.session_state["agree_cmax2 - ИБ"] == True:
                       df_total_PK_po_rdf = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    else:
                       df_total_PK_po_rdf = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                    
                    df_total_PK_po_rdf.index.name = 'Номер'

                    ##изменение названий параметров описательной статистики

                    df_total_PK_po_rdf1=df_total_PK_po_rdf.copy()
                    df_total_PK_po_rdf1.iloc[-6,:],df_total_PK_po_rdf1.iloc[-2,:]=df_total_PK_po_rdf.iloc[-2,:],df_total_PK_po_rdf.iloc[-6,:]

                    df_total_PK_po_rdf=df_total_PK_po_rdf1

                    df_total_PK_po_rdf1=df_total_PK_po_rdf.copy()
                    df_total_PK_po_rdf1.iloc[-4,:],df_total_PK_po_rdf1.iloc[-5,:]=df_total_PK_po_rdf.iloc[-5,:],df_total_PK_po_rdf.iloc[-4,:]

                    df_total_PK_po_rdf=df_total_PK_po_rdf1

                    df_total_PK_po_rdf = df_total_PK_po_rdf.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')


                    table_heading='Фармакокинетические показатели в крови после перорального введения ГЛФ'
                    list_heading_word.append(table_heading)
                    
                    list_table_word.append(df_total_PK_po_rdf)

                    ####получение интервала для средних ФК параметров
                    list_PK_Cmax_1_not_round = df_PK['Cmax'].tolist()
                    list_PK_Tmax_1_not_round = df_PK['Tmax'].tolist() 
                    list_PK_MRT0_inf_not_round = df_PK['MRT0→∞'].tolist() 
                    list_PK_half_live_not_round = df_PK['T1/2'].tolist() 
                    list_PK_AUC0_t_not_round = df_PK['AUC0-t'].tolist()
                    list_PK_AUC0_inf_not_round = df_PK['AUC0→∞'].tolist()
                    list_PK_AUMC0_inf_not_round = df_PK['AUMC0-∞'].tolist()
                    list_PK_Сmax_dev_AUC0_t_not_round = df_PK['Сmax/AUC0-t'].tolist()
                    list_PK_Kel_not_round = df_PK['Kel'].tolist()

                    list_list_PK_parametr_po_rdf=[list_PK_Cmax_1_not_round,list_PK_AUC0_t_not_round,list_PK_Kel_not_round,list_PK_AUC0_inf_not_round,list_PK_half_live_not_round,list_PK_AUMC0_inf_not_round,list_PK_MRT0_inf_not_round,list_PK_Сmax_dev_AUC0_t_not_round]
                    list_parametr_mean_h_po_rdf=[]
                    for i in list_list_PK_parametr_po_rdf:
                         n=len(i)

                         def confidential_interval(i):
                             if n < 30:
                                h = statistics.stdev(i)
                                mean = np.mean(i)
                             else:
                                h = statistics.stdev(i)  ### прояснить момент с n-1
                                mean = np.mean(i)
                             return ([mean,h]) 
                         func_mean_h = confidential_interval(i)

                         list_parametr_mean_h_po_rdf.append(func_mean_h)


                    list_mean_h_po_rdf_Cmax_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_rdf[0]]
                    parametr_round_mean_h_Cmax=str(list_mean_h_po_rdf_Cmax_round[0]) +"±"+str(list_mean_h_po_rdf_Cmax_round[1])

                    list_mean_h_po_rdf_AUC0_t_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_rdf[1]] 
                    parametr_round_mean_h_AUC0_t=str(list_mean_h_po_rdf_AUC0_t_round[0]) +"±"+str(list_mean_h_po_rdf_AUC0_t_round[1]) 

                    list_mean_h_po_rdf_Kel_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_po_rdf[2]]
                    parametr_round_mean_h_Kel=str(list_mean_h_po_rdf_Kel_round[0]) +"±"+str(list_mean_h_po_rdf_Kel_round[1])

                    list_mean_h_po_rdf_AUC0_inf_round= ["%.2f" % round(v,2) for v in list_parametr_mean_h_po_rdf[3]]
                    parametr_round_mean_h_AUC0_inf=str(list_mean_h_po_rdf_AUC0_inf_round[0]) +"±"+str(list_mean_h_po_rdf_AUC0_inf_round[1]) 

                    list_mean_h_po_rdf_half_live_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_rdf[4]]
                    parametr_round_mean_h_half_live=str(list_mean_h_po_rdf_half_live_round[0]) +"±"+str(list_mean_h_po_rdf_half_live_round[1])

                    list_mean_h_po_rdf_AUMC0_inf_round=["%.2f" % round(v,2) for v in list_parametr_mean_h_po_rdf[5]] 
                    parametr_round_mean_h_AUMC0_inf=str(list_mean_h_po_rdf_AUMC0_inf_round[0]) +"±"+str(list_mean_h_po_rdf_AUMC0_inf_round[1]) 

                    list_mean_h_po_rdf_MRT0_inf_round=["%.3f" % round(v,3) for v in list_parametr_mean_h_po_rdf[6]]
                    parametr_round_mean_h_MRT0_inf=str(list_mean_h_po_rdf_MRT0_inf_round[0]) +"±"+str(list_mean_h_po_rdf_MRT0_inf_round[1])

                    list_mean_h_po_rdf_Сmax_dev_AUC0_t_round=["%.4f" % round(v,4) for v in list_parametr_mean_h_po_rdf[7]]
                    parametr_round_mean_h_Сmax_dev_AUC0_t=str(list_mean_h_po_rdf_Сmax_dev_AUC0_t_round[0]) +"±"+str(list_mean_h_po_rdf_Сmax_dev_AUC0_t_round[1])

                    list_parametr_round_mean_h_po_rdf= [parametr_round_mean_h_Cmax,parametr_round_mean_h_AUC0_t,parametr_round_mean_h_Kel,parametr_round_mean_h_AUC0_inf,parametr_round_mean_h_half_live,parametr_round_mean_h_AUMC0_inf,parametr_round_mean_h_MRT0_inf,parametr_round_mean_h_Сmax_dev_AUC0_t]

                    t_mean_po_rdf = str("%.2f" % round(np.mean(list_PK_Tmax_1_not_round),2))     
                    list_parametr_round_mean_h_po_rdf.insert(1,t_mean_po_rdf)

              ###Биодоступность
              button_calculation_bioavailability = False
              
              if ("uploaded_file_1" in st.session_state) and ("uploaded_file_2" in st.session_state) and ("uploaded_file_3" in st.session_state) and measure_unit and dose_iv and dose_po_sub and dose_po_rdf:
                 
                 condition_iv_cmax1 =  len(list_cmax_1_iv) == count_rows_number_iv
                 condition_sub_cmax1 = len(list_cmax_1_sub) == count_rows_number_sub
                 condition_rdf_cmax1 = len(list_cmax_1_rdf) == count_rows_number_rdf
                 
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    condition_iv_cmax2 =  len(list_cmax_2_iv) == count_rows_number_iv
                    condition_sub_cmax2 = len(list_cmax_2_sub) == count_rows_number_sub
                    condition_rdf_cmax2 = len(list_cmax_2_rdf) == count_rows_number_rdf
                 
                 if st.session_state["agree_cmax2 - ИБ"] == True:
                    if (condition_iv_cmax2 and condition_sub_cmax2 and condition_rdf_cmax2):
                       button_calculation_bioavailability = True
                 if st.session_state["agree_cmax2 - ИБ"] == False:
                    if (condition_iv_cmax1 and condition_sub_cmax1 and condition_rdf_cmax1):
                       button_calculation_bioavailability = True

                 if button_calculation_bioavailability == True:
                    st.write('👩🏽‍💻Биодоступность подсчитана!')
                 else:   
                    st.write('🔧Заполните все поля ввода и загрузите файлы!')

              if ("uploaded_file_1" in st.session_state) and ("uploaded_file_2" in st.session_state) and ("uploaded_file_3" in st.session_state) and measure_unit and dose_iv and dose_po_sub and dose_po_rdf and button_calculation_bioavailability:
                  
                  table_heading='Усредненные фармакокинетические параметры в крови после внутривенного введения субстанции, перорального введения субстанции и перорального введения ГЛФ, а также абсолютная и относительная биодоступность'
                  list_heading_word.append(table_heading)

                  AUCT_inf_mean_iv = df_concat_PK_iv["AUC0-t"].loc["mean"]
                  AUCT_inf_mean_po_sub = df_concat_PK_po_sub["AUC0-t"].loc["mean"]
                  AUCT_inf_mean_po_rdf = df_concat_PK_po_rdf["AUC0-t"].loc["mean"]

                  #абсолютная биодоступность

                  F_po_sub_iv=round((AUCT_inf_mean_po_sub * float(dose_iv))/(AUCT_inf_mean_iv*float(dose_po_sub))*100,2)
                  F_po_rdf_iv=round((AUCT_inf_mean_po_rdf * float(dose_iv))/(AUCT_inf_mean_iv*float(dose_po_rdf))*100,2)

                  #относительная биодоступность
                  RF_po_sub_rdf=round((AUCT_inf_mean_po_rdf*float(dose_po_sub))/(AUCT_inf_mean_po_sub*float(dose_po_rdf))*100,2)

                  df_intravenous_substance = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_1"]))
                  df_oral_substance = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_2"]))
                  df_oral_rdf = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_3"]))

                  df_averaged_concentrations_intravenous_substance=df_intravenous_substance.describe()
                  list_concentration__intravenous_substance=df_averaged_concentrations_intravenous_substance.loc['mean'].tolist()

                  df_averaged_concentrations_oral_substance=df_oral_substance.describe()
                  list_concentration__oral_substance=df_averaged_concentrations_oral_substance.loc['mean'].tolist()

                  df_averaged_concentrations_oral_rdf=df_oral_rdf.describe()
                  list_concentration__oral_rdf=df_averaged_concentrations_oral_rdf.loc['mean'].tolist()

              ### итоговый фрейм по PK параметрам

                  list_index_for_df_total_PK_mean = ['Cmax ' +"("+measure_unit+")",'Tmax ' +"("+"ч"+")",'AUC0-t '+"("+measure_unit+"×ч" +")",'Kel '+"("+"ч\u207B\u00B9"+")",'AUC0→∞ '+"("+measure_unit+"×ч" +")",'T1/2 '+"("+"ч"+")",'AUMC0-∞ '+"("+measure_unit+"×ч\u00B2"+")",'MRT0→∞ '+"("+"ч"+")",'Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")","F(абсолютная биодоступность),%","Относительная биодоступность,% (по сравнению с пероральным введением субстанции)"]
                  
                  #добавление значений биодоступности
                  list_parametr_round_mean_h_iv.append("-")
                  list_parametr_round_mean_h_iv.append("-")

                  list_parametr_round_mean_h_po_sub.append(F_po_sub_iv)
                  list_parametr_round_mean_h_po_sub.append("-")

                  list_parametr_round_mean_h_po_rdf.append(F_po_rdf_iv)
                  list_parametr_round_mean_h_po_rdf.append(RF_po_sub_rdf)


                  df_total_PK_mean = pd.DataFrame(list(zip(list_parametr_round_mean_h_iv,list_parametr_round_mean_h_po_sub,list_parametr_round_mean_h_po_rdf)),columns=['Внутривенное введение субстанции','Пероральное введение субстанции','Пероральное введение ГЛФ'],index=list_index_for_df_total_PK_mean)
                  df_total_PK_mean.index.name = 'Параметры, размерность'
                  
                  list_table_word.append(df_total_PK_mean)

              #####объединенные графики

              ### в линейных координатах
                  col_mapping = df_intravenous_substance.columns.tolist() ### можно указать любой фрейм
                  col_mapping.remove('Номер')
                  list_time = []
                  for i in col_mapping:
                      numer=float(i)
                      list_time.append(numer)

                  err_y_1=df_averaged_concentrations_intravenous_substance.loc['std'].tolist()
                  err_y_2=df_averaged_concentrations_oral_substance.loc['std'].tolist()
                  err_y_3=df_averaged_concentrations_oral_rdf.loc['std'].tolist()
                  
                  df_total_injection = pd.DataFrame(list(zip(list_concentration__intravenous_substance, list_concentration__oral_substance, list_concentration__oral_rdf)),columns =['внутривенное введение','пероральное введение субстанции','пероральное введение ГЛФ'])
                  df_total_injection.loc[df_total_injection["внутривенное введение"] == 0, "внутривенное введение"] = np.nan #т.к. внутривенное введение
                  
                  df_total_error = pd.DataFrame(list(zip(err_y_1, err_y_2, err_y_3)),columns =['внутривенное введение','пероральное введение субстанции','пероральное введение ГЛФ'])
                  df_total_error.loc[df_total_injection["внутривенное введение"] == 0, "внутривенное введение"] = np.nan #т.к. внутривенное введение
                  list_name_injection = ['внутривенное введение','пероральное введение субстанции','пероральное введение ГЛФ']
                  list_name_colors = ["black","red","blue"]
                  zip_injection_colors_error = zip(list_name_injection,list_name_colors)


                  fig, ax = plt.subplots()
                  
                  for injection,color in zip_injection_colors_error:
                      plt.errorbar(list_time,df_total_injection[injection],yerr=df_total_error[injection],color= color, marker='o',markersize=4.0,markeredgecolor=color,markerfacecolor=color,ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0, label = injection)
                      ax.set_xlabel("Время, ч")
                      ax.set_ylabel("Концентрация, "+measure_unit)
                      ax.legend()

                  list_graphics_word.append(fig) 

                  graphic="Сравнение фармакокинетических профилей (в линейных координатах) после внутривенного введения субстанции, перорального введения субстанции и перорального введения ГЛФ"
                  list_heading_graphics_word.append(graphic) 
              ### в полулогарифмических координатах
                  list_time.remove(0)
                  
                  list_concentration__intravenous_substance.remove(0)
                  list_concentration__oral_substance.remove(0)
                  list_concentration__oral_rdf.remove(0)
                  
                  err_y_1.remove(0)
                  err_y_2.remove(0) 
                  err_y_3.remove(0) 

                  fig, ax = plt.subplots()    

                  plt.errorbar(list_time,list_concentration__intravenous_substance,yerr=err_y_1,color="black", marker='o',markersize=4.0,markeredgecolor="black",markerfacecolor="black",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0, label = 'внутривенное введение')
                  plt.errorbar(list_time,list_concentration__oral_substance,yerr=err_y_2,color= "red", marker='o',markersize=4.0,markeredgecolor="red",markerfacecolor="red",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0, label = 'пероральное введение субстанции')
                  plt.errorbar(list_time,list_concentration__oral_rdf,yerr=err_y_3,color= "blue", marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0, label = 'пероральное введение ГЛФ')

                  ax.set_yscale("log")
                  ax.set_xlabel("Время, ч")
                  ax.set_ylabel("Концентрация, "+measure_unit)
                  ax.legend()

                  list_graphics_word.append(fig) 

                  graphic="Сравнение фармакокинетических профилей (в полулогарифмических координатах) после внутривенного введения субстанции, перорального введения субстанции и перорального введения ГЛФ"
                  list_heading_graphics_word.append(graphic)
              else:
                  st.write("")

              ##############################################################################################################

              ###сохранение состояния 
              st.session_state["list_heading_word"] = list_heading_word
              st.session_state["list_table_word"] = list_table_word
              st.session_state["list_graphics_word"] = list_graphics_word
              st.session_state["list_heading_graphics_word"] = list_heading_graphics_word
              
           
           #####Создание word отчета
           if panel == "Таблицы":

                 list_heading_word = st.session_state["list_heading_word"]
                 list_table_word = st.session_state["list_table_word"]

                 ###вызов функции создания таблицы
                 create_table(list_heading_word,list_table_word)

           if panel == "Графики":
                 
                 list_graphics_word = st.session_state["list_graphics_word"]
                 list_heading_graphics_word = st.session_state["list_heading_graphics_word"]
                 
                 ###вызов функции создания графика
                 create_graphic(list_graphics_word,list_heading_graphics_word)

                 #######визуализация

                 #классификация графиков по кнопкам
                 type_graphics = st.selectbox('Выберите вид графиков',
           ('Индивидуальные фармакокинетические профили', 'Сравнение индивидуальных фармакокинетических профилей', 'Графики усредненного фармакокинетического профиля', "Сравнение фармакокинетических профилей при разных видах введения"),disabled = False, key = "Вид графика - ИБ" )

                 count_graphics_for_visual = len(list_heading_graphics_word)
                 list_range_count_graphics_for_visual = range(0,count_graphics_for_visual)
                 
                 for i in list_range_count_graphics_for_visual:
                     if list_heading_graphics_word[i].__contains__("индивидуального"): 
                        if type_graphics == 'Индивидуальные фармакокинетические профили':
                           st.pyplot(list_graphics_word[i])
                           st.subheader(list_heading_graphics_word[i])
                     if list_heading_graphics_word[i].__contains__("Сравнение индивидуальных"):   
                        if type_graphics == 'Сравнение индивидуальных фармакокинетических профилей':
                           st.pyplot(list_graphics_word[i])
                           st.subheader(list_heading_graphics_word[i])
                     if list_heading_graphics_word[i].__contains__("усредненного"):
                        if type_graphics == 'Графики усредненного фармакокинетического профиля':
                           st.pyplot(list_graphics_word[i])
                           st.subheader(list_heading_graphics_word[i])
                     if list_heading_graphics_word[i].__contains__("Сравнение фармакокинетических"):
                        if type_graphics == 'Сравнение фармакокинетических профилей при разных видах введения':
                           st.pyplot(list_graphics_word[i])
                           st.subheader(list_heading_graphics_word[i])
                     
   #####################################################################        
   if option == 'Изучение фармакокинетики в органах животных':
      
      st.title('Исследование ФК параметров для органов животных')
      
      col1, col2 = st.columns([0.66, 0.34])
      
      ######### боковое меню справа
      with col2:
           selected = option_menu(None, ["Включение параметров в исследование"], 
           icons=['menu-button'], 
           menu_icon="cast", default_index=0, orientation="vertical",
           styles={
               "container": {"padding": "0!important", "background-color": "#24769C"},
               "icon": {"color": "#5DAED3", "font-size": "13px"}, 
               "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
               "nav-link-selected": {"background-color": "#335D70"},
           })

           if selected == "Включение параметров в исследование":
              type_parameter = st.selectbox('Выберите параметр',
           ('Cmax(2)',"Вид введения"),disabled = False, key = "Вид параметра - органы")
              

           if type_parameter == 'Cmax(2)':
              
              if "agree_cmax2 - органы" not in st.session_state:
                 st.session_state["agree_cmax2 - органы"] = False

              st.session_state["agree_cmax2 - органы"] = st.checkbox('Добавить возможность выбора Cmax(2)', key = "Возможность добавления Cmax2 - органы", value = st.session_state["agree_cmax2 - органы"])
              
              if st.session_state["agree_cmax2 - органы"] == True:
                 st.write('🧠Параметр добавлен!')

      with col1:
          
         panel = st.radio(
              "⚙️Панель управления",
              ("Загрузка файлов", "Таблицы","Графики"),
              horizontal=True, key= "Загрузка файлов - Исследование ФК параметров для органов животных"
          )

         ###создание состояния
         if "measure_unit_org_blood" not in st.session_state:
            st.session_state["measure_unit_org_blood"] = ""
         if "measure_unit_org_organs" not in st.session_state:
            st.session_state["measure_unit_org_organs"] = ""
         if "dose_org" not in st.session_state:
            st.session_state["dose_org"] = ""

         #cписки для word-отчета
         list_heading_word=[]
         list_table_word=[]
         list_graphics_word=[]
         list_heading_graphics_word=[]
          
         if panel == "Загрузка файлов":

            measure_unit_org_blood = st.text_input("Введите единицы измерения концентрации в крови", key='Единицы измерения при изучении фармакокинетики в органах животных в крови', value = st.session_state["measure_unit_org_blood"])
            
            st.session_state["measure_unit_org_blood"] = measure_unit_org_blood

            measure_unit_org_organs = st.text_input("Введите единицы измерения концентрации в органах", key='Единицы измерения при изучении фармакокинетики в органах животных в органах', value = st.session_state["measure_unit_org_organs"])
            
            st.session_state["measure_unit_org_organs"] = measure_unit_org_organs

            dose = st.text_input("Доза препарата", key='Доза препарата при изучении фармакокинетики в органах животных', value = st.session_state["dose_org"])
            
            st.session_state["dose_org"] = dose

            #cостояние радио-кнопки "method_auc"
            if "index_method_auc - ИО" not in st.session_state:
                st.session_state["index_method_auc - ИО"] = 0

            method_auc = st.radio("📌Метод подсчёта AUC0-t",('linear',"linear-up/log-down"),key = "Метод подсчёта AUC0-t - ИО", index = st.session_state["index_method_auc - ИО"])
            
            if st.session_state["Метод подсчёта AUC0-t - ИО"] == 'linear':
               st.session_state["index_method_auc - ИО"] = 0
            if st.session_state["Метод подсчёта AUC0-t - ИО"] == "linear-up/log-down":
               st.session_state["index_method_auc - ИО"] = 1

            file_uploader = st.file_uploader("Выберите нужное количество файлов соответственно количеству исследуемых органов(в том числе файл для крови); файл должен быть назван соотвественно органу;исходный файл крови должен быть назван 'Кровь'",accept_multiple_files=True, key='Файлы при изучении фармакокинетики в органах животных')

            ###сохранение файла
            if file_uploader is not None:
               for i in file_uploader:
                   save_uploadedfile(i)
                   st.session_state[str(i.name)] = i.name
            
            list_keys_file_org = []
            for i in st.session_state.keys():
                if i.__contains__("xlsx") and (not i.__contains__("Дозировка")) and (not i.__contains__("edited_df")):### чтобы не перекрывалось с lin; #обрезаем фразу ненужного добавления названия "edited_df"
                   list_keys_file_org.append(i)
            
            if (list_keys_file_org != []) and dose and measure_unit_org_blood and measure_unit_org_organs:

                list_name_organs=[]
                list_df_unrounded=[]
                list_df_for_mean_unround_for_graphics=[]
                list_t_graph=[]
                
                for i in list_keys_file_org:
                    df = pd.read_excel(os.path.join("Папка для сохранения файлов",i))

                    file_name=st.session_state[i][:-5]

                    st.subheader('Индивидуальные значения концентраций ' + "("+file_name+")")
                    
                    ###интерактивная таблица
                    df = edit_frame(df,i)

                    ###количество животных 
                    count_rows_number_org = len(df.axes[0])

                    table_heading='Индивидуальные и усредненные значения концентраций ' + "("+file_name+")"
                    list_heading_word.append(table_heading)

                    ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                    df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']

                    list_table_word.append(df_concat_round_str_transpose) 

                    if file_name == "Кровь":
                       measure_unit_org = measure_unit_org_blood
                    else:
                       measure_unit_org = measure_unit_org_organs
                    ########### графики    

                    ######индивидуальные    

                    # в линейных координатах 
                    col_mapping = df.columns.tolist()
                    col_mapping.remove('Номер')

                    count_row_df = len(df.axes[0])

                    list_time = []
                    for i in col_mapping:
                        numer=float(i)
                        list_time.append(numer)
                    list_t_graph.append(list_time) 

                    for r in range(0,count_row_df):

                        list_concentration=df.iloc[r].tolist()

                        numer_animal=list_concentration[0]

                        list_concentration.pop(0) #удаление номера животного

                        list_concentration = [float(v) for v in list_concentration]


                        fig, ax = plt.subplots()
                        plt.plot(list_time,list_concentration,marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+measure_unit_org)
         
                        list_graphics_word.append(fig)

                        graphic='График индивидуального фармакокинетического профиля в линейных координатах '  + "("+file_name+")"',  '+numer_animal
                        list_heading_graphics_word.append(graphic)  
                        

                     #в полулогарифмических координатах методом удаления точек
                        count_for_0_1=len(list_concentration)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_for_log_1=[]
                        for i in list_range_for_0_1:
                            if list_concentration[i] !=0:
                               list_for_log_1.append(list_concentration[i])
                               list_time_0.append(list_time[i]) 

                        fig, ax = plt.subplots()
                        plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                        ax.set_yscale("log")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+measure_unit_org)

                        
                        list_graphics_word.append(fig)

                        graphic='График индивидуального фармакокинетического профиля в полулогарифмических координатах ' + "("+file_name+")"',  '+numer_animal
                        list_heading_graphics_word.append(graphic) 
          
                 # объединенные индивидуальные в линейных координатах

                    df_for_plot_conc=df.drop(['Номер'], axis=1)
                    df_for_plot_conc_1 = df_for_plot_conc.transpose()
                    list_numer_animal_for_plot=df['Номер'].tolist()
                    count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды
                    list_color = [] ## генерация 500 цветов
                    for i in range(0,500):
                        hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                        list_color.append(hexadecimal)

                    fig, ax = plt.subplots()
                    
                    ax.set_prop_cycle(cycler(color=list_color))

                    plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                    ax.set_xlabel("Время, ч")
                    ax.set_ylabel("Концентрация, "+measure_unit_org)
                    if count_numer_animal > 20:
                       ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                    else:
                       ax.legend(bbox_to_anchor=(1, 1))
                    
                    list_graphics_word.append(fig)

                    graphic="Сравнение индивидуальных фармакокинетических профилей в линейных координатах " + "("+file_name+")"
                    list_heading_graphics_word.append(graphic)     
                 # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                    df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)


                    fig, ax = plt.subplots()

                    ax.set_prop_cycle(cycler(color=list_color))

                    plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                    ax.set_xlabel("Время, ч")
                    ax.set_ylabel("Концентрация, "+measure_unit_org)
                    ax.set_yscale("log")
                    if count_numer_animal > 20:
                       ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                    else:
                       ax.legend(bbox_to_anchor=(1, 1))
                    
                    list_graphics_word.append(fig)

                    graphic="Сравнение индивидуальных фармакокинетических профилей в полулогарифмических координатах " + "("+file_name+")"
                    list_heading_graphics_word.append(graphic)
                     ###усредненные    
                 # в линейных координатах
                    list_time = []
                    for i in col_mapping:
                        numer=float(i)
                        list_time.append(numer)

                    df_averaged_concentrations=df.describe()
                    list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                    err_y_1=df_averaged_concentrations.loc['std'].tolist()


                    fig, ax = plt.subplots()
                    plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                    plt.xlabel("Время, ч")
                    plt.ylabel("Концентрация, "+measure_unit_org)
                    
                    list_graphics_word.append(fig)

                    graphic='График усредненного фармакокинетического профиля в линейных координатах ' + "("+file_name+")"
                    list_heading_graphics_word.append(graphic)

                 #в полулогарифмических координатах
                    #для полулогарифм. посторим без нуля
                    list_time.remove(0)
                    list_concentration.remove(0)
                    err_y_1.remove(0) 


                    fig, ax = plt.subplots()
                    plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                    ax.set_yscale("log")
                    plt.xlabel("Время, ч")
                    plt.ylabel("Концентрация, "+measure_unit_org)

                    
                    list_graphics_word.append(fig)

                    graphic='График усредненного фармакокинетического профиля в полулогарифмических координатах ' + "("+file_name+")"
                    list_heading_graphics_word.append(graphic)

                    ############ Параметры ФК

                    df_without_numer=df.drop(['Номер'],axis=1)
                    count_row=df_without_numer.shape[0]

                    list_count_row=range(count_row)
          
                    ###Cmax
                    #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                    if st.session_state["agree_cmax2 - органы"] == False:
                       list_cmax_1_org=[]
                       for i in range(0,count_row):
                           cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                           list_cmax_1_org.append(cmax)
                 
                    #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                    if st.session_state["agree_cmax2 - органы"] == True:
                        
                       ###создание состояния
                       if ("selected_value_org" + file_name) not in st.session_state:
                          st.session_state["selected_value_org"+ file_name] = []
                       
                       if ("feature_disable_selected_value_org" + file_name) not in st.session_state:
                           st.session_state["feature_disable_selected_value_org" + file_name] = True

                       ###создание состояния
                       st.info('Выбери Cmax:')
                       list_columns_without_numer = df.columns.tolist()
                       list_columns_without_numer.remove('Номер')
                       selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax органы ' + file_name, max_selections=1)
                       st.session_state["selected_columns_org"+ file_name] = selected_columns 

                       list_keys_cmax = st.session_state["selected_value_org"+ file_name]
                       if selected_columns != [] and st.session_state["feature_disable_selected_value_org"+ file_name]:
                          selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax органы ' + file_name, max_selections=1)
                          list_keys_cmax.append(selected_value)

                       if list_keys_cmax != []:
                          st.session_state["selected_value_org"+ file_name] = list_keys_cmax

                       list_keys_cmax = st.session_state["selected_value_org"+ file_name]
                       list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]
                       
                       if st.button('Очистить список Cmax', key="Очистка списка Cmax органы " + file_name):
                          del st.session_state["selected_value_org"+ file_name]
                          list_keys_cmax_sample = []
                          selected_columns = st.session_state["selected_columns_org"+ file_name]
                          st.session_state["feature_disable_selected_value_org"+ file_name] = True
                       
                       st.write("Список Cmax:")
                       st.write(list_keys_cmax_sample)
                       
                       list_cmax_1_org=list_keys_cmax_sample 
                       
                       list_cmax_2_org=[]

                    if len(list_cmax_1_org) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - органы"] == True):
                       
                       st.session_state["feature_disable_selected_value_org"+ file_name] = False

                       ######Cmax2

                       if ("feature_disable_selected_value_org_2"+ file_name) not in st.session_state:
                        st.session_state["feature_disable_selected_value_org_2"+ file_name] = True

                       st.info('Выбери Cmax(2):')
                       
                       selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 органы ' + file_name, max_selections=1)
                       st.session_state["selected_columns_2_org"+ file_name] = selected_columns_2

                       ###создание состояния
                       if ("selected_value_2_org"+ file_name) not in st.session_state:
                          st.session_state["selected_value_2_org"+ file_name] = []

                       list_keys_cmax_2 = st.session_state["selected_value_2_org"+ file_name]
                       if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_org_2"+ file_name]:
                          selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 органы '  + file_name, max_selections=1)
                          list_keys_cmax_2.append(selected_value_2)

                       if list_keys_cmax_2 != []:
                          st.session_state["selected_value_2_org"+ file_name] = list_keys_cmax_2

                       list_keys_cmax_2 = st.session_state["selected_value_2_org"+ file_name]
                       list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                       if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) органы " + file_name):
                          del st.session_state["selected_value_2_org"+ file_name]
                          list_keys_cmax_sample_2 = []
                          selected_columns_2 = st.session_state["selected_columns_2_org"+ file_name]
                          st.session_state["feature_disable_selected_value_org_2"+ file_name] = True

                       st.write("Список Cmax(2):")
                       st.write(list_keys_cmax_sample_2)

                       list_cmax_2_org= list_keys_cmax_sample_2

                       if len(list_cmax_2_org) == len(df.index.tolist()):
                          st.session_state["feature_disable_selected_value_org_2"+ file_name] = False
                    
                    if (len(list_cmax_1_org) == len(df.index.tolist())):
                       
                       ###Tmax   
                       list_Tmax_1=[]
                       for cmax in list_cmax_1_org:
                           for column in df.columns:
                               for num, row in df.iterrows():
                                   if df.iloc[num][column] == cmax:
                                      list_Tmax_1.append(f"{column}")
                     
                       list_Tmax_float_1=[]           
                       for i in list_Tmax_1:
                           Tmax=float(i)
                           list_Tmax_float_1.append(Tmax)

                    if (len(list_cmax_1_org) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - органы"] == True):
                       
                       list_Tmax_2=[]
                       for cmax in list_cmax_2_org:
                           for column in df.columns:
                               for num, row in df.iterrows():
                                   if df.iloc[num][column] == cmax:
                                      list_Tmax_2.append(f"{column}")
                     
                       list_Tmax_float_2=[]           
                       for i in list_Tmax_2:
                           Tmax=float(i)
                           list_Tmax_float_2.append(Tmax)  

                    if (len(list_cmax_1_org) == len(df.index.tolist())):
                       
                       ###AUC0-t
                       list_AUC_0_T=[]
                       if method_auc == 'linear':
                          for i in range(0,count_row):
                              list_columns_T=[]
                              for column in df_without_numer.columns:
                                  list_columns_T.append(float(column))
                              list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                              ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                              cmax = max(list_concentration)
                              index_cmax = list_concentration.index(cmax)
                              list_before_cmax = list_concentration[0:index_cmax]
                              list_after_cmax = list_concentration[index_cmax:]
                              list_before_cmax_t = list_columns_T[0:index_cmax]
                              list_after_cmax_t = list_columns_T[index_cmax:]

                              count_list_concentration = len(list_after_cmax)
                              list_range_for_remove_0 = range(0,count_list_concentration)

                              list_conc_without_0=[]
                              list_t_without_0=[]
                              for i in list_range_for_remove_0:
                                  if list_after_cmax[i] !=0:
                                     list_conc_without_0.append(list_after_cmax[i])
                                     list_t_without_0.append(list_after_cmax_t[i])

                              list_concentration = list_before_cmax + list_conc_without_0
                              list_columns_T = list_before_cmax_t + list_t_without_0
                              ######################

                              AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                              list_AUC_0_T.append(AUC_0_T)

                       if method_auc == 'linear-up/log-down':
                          for i in range(0,count_row):
                              list_columns_T=[]
                              for column in df_without_numer.columns:
                                  list_columns_T.append(float(column))
                              list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                              ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                              cmax = max(list_concentration)
                              index_cmax = list_concentration.index(cmax)
                              list_before_cmax = list_concentration[0:index_cmax]
                              list_after_cmax = list_concentration[index_cmax:]
                              list_before_cmax_t = list_columns_T[0:index_cmax]
                              list_after_cmax_t = list_columns_T[index_cmax:]

                              count_list_concentration = len(list_after_cmax)
                              list_range_for_remove_0 = range(0,count_list_concentration)

                              list_conc_without_0=[]
                              list_t_without_0=[]
                              for i in list_range_for_remove_0:
                                  if list_after_cmax[i] !=0:
                                     list_conc_without_0.append(list_after_cmax[i])
                                     list_t_without_0.append(list_after_cmax_t[i])

                              list_concentration = list_before_cmax + list_conc_without_0
                              list_columns_T = list_before_cmax_t + list_t_without_0
                              ######################
                              
                              list_c = list_concentration
                              list_t = list_columns_T
                              
                              count_i = len(list_c)
                              list_range= range(0,count_i)
                              
                              list_AUC_0_T_ascending=[]
                              list_AUC_0_T_descending = []
                              AUC_0_T_ascending=0
                              AUC_0_T_descending = 0
                              a=0
                              a1=0
                              d=0
                              d1=0
                              for i in list_range:
                                  if a1<count_i-1:
                                     if list_c[i+1] > list_c[i]:
                                        if a<count_i-1:
                                            AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                            a+=1
                                            list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                                  if d1<count_i-1:
                                     if list_c[i+1] < list_c[i]:      
                                        if d<count_i-1:
                                            AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                            d+=1
                                            list_AUC_0_T_descending.append(AUC_0_T_descending)
                                     a1+=1
                                     d1+=1
               
                              AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                              
                              list_AUC_0_T.append(AUC_O_T)

                       ####Сmax/AUC0-t
                       list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_org,list_AUC_0_T)
                       list_Сmax_division_AUC0_t=[]
                       for i,j in list_Сmax_division_AUC0_t_for_division:
                               list_Сmax_division_AUC0_t.append(i/j)


                       ####KEL
                       list_kel_total=[]
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                           list_concentration.remove(0)
                           list_c=list_concentration

                           list_time=df_without_numer.columns.tolist()
                           list_time.remove(0) 

                           list_t=[]
                           for i in list_time:
                               i=float(i)
                               list_t.append(i)

                           #срез_без_cmax
                           max_value_c=max(list_c)
                           index_cmax=list_c.index(max_value_c)

                           list_c_without_cmax=list_c[index_cmax+1:]
                           list_t_without_cmax=list_t[index_cmax+1:]

                           #удаление всех нулей из массивов
                           count_for_0_1=len(list_c_without_cmax)
                           list_range_for_0_1=range(0,count_for_0_1)

                           list_time_0=[]
                           list_conc_0=[]
                           for i in list_range_for_0_1:
                               if list_c_without_cmax[i] !=0:
                                  list_conc_0.append(list_c_without_cmax[i])
                                  list_time_0.append(list_t_without_cmax[i]) 
                           ################################

                           n_points=len(list_conc_0)
                           list_n_points = range(0,n_points)

                           #создание списков с поочередно уменьщающемся кол, точек
                           list_for_kel_c=[]
                           for j in list_n_points:
                               if j<n_points:
                                  list_c_new=list_conc_0[j:n_points]
                                  list_for_kel_c.append(list_c_new)
                           list_for_kel_c.pop(-1) #удаление списка с одной точкой
                           list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                           list_for_kel_t=[]
                           for j in list_n_points:
                               if j<n_points:
                                  list_t_new=list_time_0[j:n_points]
                                  list_for_kel_t.append(list_t_new)
                           list_for_kel_t.pop(-1) #удаление списка с одной точкой
                           list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                           list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                           list_kel=[]
                           list_r=[]
                           for i,j in list_ct_zip:

                               n_points_r=len(i)

                               np_c=np.asarray(i)
                               np_t_1=np.asarray(j).reshape((-1,1))

                               np_c_log=np.log(np_c)

                               model = LinearRegression().fit(np_t_1,np_c_log)

                               np_t=np.asarray(j)
                               a=np.corrcoef(np_t, np_c_log)
                               cor=((a[0])[1])
                               r_sq=cor**2

                               adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                               ########################################
                               kel=abs(model.coef_[0])
                               list_kel.append(kel)
                               list_r.append(adjusted_r_sq)

                           #делаем срезы списоков до rmax
                           max_r=max(list_r)

                           index_max_r= list_r.index(max_r)

                           list_r1=list_r
                           list_kel1=list_kel

                           number_elem_list_r1=len(list_r1)

                           list_range_kel=range(0,number_elem_list_r1) 

                           list_kel_total_1=[]
                           for i in list_range_kel:

                               if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                                  list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                                  break #самая ранняя удовлетовряющая условию

                           for i in list_kel_total_1:
                               list_kel_total.append(i) 


                       ####T1/2
                       list_half_live=[]
                       for i in list_kel_total:
                           half_live=math.log(2)/i
                           list_half_live.append(half_live)


                       ###AUC0-inf 

                       list_auc0_inf=[] 

                       list_of_list_c=[]
                       for i in range(0,count_row):
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                           list_concentration.remove(0)
                           list_c = list_concentration
                           list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                           list_of_list_c.append(list_c)

                       list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                           #AUCt-inf 
                       list_auc_t_inf=[]     
                       for i,j in list_zip_c_AUCt_inf:
                           for clast in j:
                               if clast != 0:
                                  clast_true=clast
                                  break
                           auc_t_inf=clast_true/i
                           list_auc_t_inf.append(auc_t_inf)

                       list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                       for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                           auc0_inf=i+j    
                           list_auc0_inf.append(auc0_inf)


                       ####CL
                       list_cl=[]

                       for i in list_auc0_inf:
                           cl = float(dose)/i * 1000
                           list_cl.append(cl)


                       ####Vd
                       list_Vd=[]

                       list_zip_kel_cl=zip(list_kel_total,list_cl)

                       for i,j in list_zip_kel_cl:
                           Vd = j/i
                           list_Vd.append(Vd)


                       ###AUMC
                       list_AUMCO_inf=[]

                       list_AUMC0_t=[]

                       list_C_last=[]
                       list_T_last=[]
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           list_C_last.append(list_concentration[-1]) 
                           list_T_last.append(list_columns_T[-1]) 

                           list_len=len(list_concentration)

                           list_aumc_i=[]
                           for i in range(0,list_len):
                               AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                               list_aumc_i.append(AUMC)

                           list_aumc_i.pop(0)

                           a=0
                           list_AUMC0_t_1=[]
                           for i in list_aumc_i:
                               a+=i
                               list_AUMC0_t_1.append(a)
                           list_AUMC0_t.append(list_AUMC0_t_1[-1])

                       list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                       list_AUMCt_inf=[]
                       for k,c,t in list_zip_for_AUMC_inf:
                           AUMCt_inf=c*t/k+c/(k*k)
                           list_AUMCt_inf.append(AUMCt_inf)


                       list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                       for i,j in list_AUMC_zip:
                           AUMCO_inf=i+j
                           list_AUMCO_inf.append(AUMCO_inf)

                       ###MRT0-inf
                       list_MRT0_inf=[]

                       list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                       for i,j in list_zip_AUMCO_inf_auc0_inf:
                           MRT0_inf=i/j
                           list_MRT0_inf.append(MRT0_inf)
                    
                    if st.session_state["agree_cmax2 - органы"] == True:
                       #####Cmax условие для дальнейшего кода  ####
                       if len(list_cmax_1_org) == len(df.index.tolist()) and len(list_cmax_2_org) == len(df.index.tolist()):

                          ##################### Фрейм ФК параметров

                          ### пользовательский индекс
                          list_for_index=df["Номер"].tolist()
                          df_PK=pd.DataFrame(list(zip(list_cmax_1_org,list_Tmax_float_1,list_cmax_2_org,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 
                    
                    if len(list_cmax_1_org) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - органы"] == False):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_org,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd'],index=list_for_index) 
                    
                    checking_condition_cmax2 = False

                    if st.session_state["agree_cmax2 - органы"] == True:
                     
                       checking_condition_cmax2 = len(list_cmax_1_org) == len(df.index.tolist()) and len(list_cmax_2_org) == len(df.index.tolist()) and st.session_state["agree_cmax2 - органы"] == True

                    if checking_condition_cmax2 or (len(list_cmax_1_org) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - органы"] == False)):
                       
                       ###описательная статистика

                       col_mapping_PK = df_PK.columns.tolist()

                       list_gmean_PK=[]

                       list_cv_PK=[] 

                       for i in col_mapping_PK:

                           list_ser_PK=df_PK[i].tolist()

                           def g_mean(list_ser_PK):
                               a=np.log(list_ser_PK)
                               return np.exp(a.mean())
                           Gmean_PK=g_mean(list_ser_PK)
                           list_gmean_PK.append(Gmean_PK)

                           cv_std_PK=lambda x: np.std(x, ddof= 1 )
                           cv_mean_PK=lambda x: np.mean(x)

                           CV_std_PK=cv_std_PK(list_ser_PK)
                           CV_mean_PK=cv_mean_PK(list_ser_PK)

                           CV_PK=(CV_std_PK/CV_mean_PK * 100)
                           list_cv_PK.append(CV_PK)


                       df_averaged_concentrations_PK=df_PK.describe()
                       df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                       df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                       df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                       df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                       df_round_without_CV_PK=df_averaged_3_PK
                       df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                       df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                       df_concat_PK_org= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)

                       ###округление описательной статистики и ФК параметров

                       series_Cmax=df_concat_PK_org['Cmax']
                       list_Cmax_str_f=["%.2f" % round(v,2) for v in series_Cmax.tolist()]
                       series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_org.index.tolist(), name='Cmax ' +"("+measure_unit_org +")")
                       
                       if st.session_state["agree_cmax2 - органы"] == True:
                          series_Cmax_2=df_concat_PK_org['Cmax(2)']
                          list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                          series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_org.index.tolist(), name='Cmax(2) ' +"("+measure_unit_org +")")

                       series_Tmax=df_concat_PK_org['Tmax']
                       list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                       series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_org.index.tolist(), name='Tmax ' +"("+"ч"+")")
                       
                       if st.session_state["agree_cmax2 - органы"] == True:
                          series_Tmax_2=df_concat_PK_org['Tmax(2)']
                          list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                          series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_org.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                       series_MRT0_inf= df_concat_PK_org['MRT0→∞']
                       list_MRT0_inf_str_f=["%.3f" % round(v,3) for v in series_MRT0_inf.tolist()]
                       series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_org.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                       series_half_live= df_concat_PK_org['T1/2']
                       list_half_live_str_f=["%.2f" % round(v,2) for v in series_half_live.tolist()]
                       series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_org.index.tolist(), name='T1/2 '+"("+"ч"+")")

                       series_AUC0_t= df_concat_PK_org['AUC0-t']
                       list_AUC0_t_str_f=["%.2f" % round(v,2) for v in series_AUC0_t.tolist()]
                       series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_org.index.tolist(), name='AUC0-t '+"("+measure_unit_org +"×ч" +")")

                       series_AUC0_inf= df_concat_PK_org['AUC0→∞']
                       list_AUC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUC0_inf.tolist()]
                       series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_org.index.tolist(), name='AUC0→∞ '+"("+measure_unit_org +"×ч" +")")

                       series_AUMC0_inf= df_concat_PK_org['AUMC0-∞']
                       list_AUMC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUMC0_inf.tolist()]
                       series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_org.index.tolist(), name='AUMC0-∞ '+"("+measure_unit_org +"×ч\u00B2" +")")

                       series_Сmax_dev_AUC0_t= df_concat_PK_org['Сmax/AUC0-t']
                       list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                       series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_org.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                       series_Kel= df_concat_PK_org['Kel']
                       list_Kel_str_f=["%.4f" % round(v,4) for v in series_Kel.tolist()]
                       series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_org.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                       series_CL= df_concat_PK_org['CL/F']
                       list_CL_str_f=["%.2f" % round(v,2) for v in series_CL.tolist()]
                       series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_org.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                       series_Vd= df_concat_PK_org['Vd']
                       list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                       series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_org.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")
                       
                       if st.session_state["agree_cmax2 - органы"] == True:
                          df_total_PK_org = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                       else:
                          df_total_PK_org = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1)

                       df_total_PK_org.index.name = 'Номер'

                       ##изменение названий параметров описательной статистики

                       df_total_PK_org1=df_total_PK_org.copy()
                       df_total_PK_org1.iloc[-6,:],df_total_PK_org1.iloc[-2,:]=df_total_PK_org.iloc[-2,:],df_total_PK_org.iloc[-6,:]

                       df_total_PK_org=df_total_PK_org1

                       df_total_PK_org1=df_total_PK_org.copy()
                       df_total_PK_org1.iloc[-4,:],df_total_PK_org1.iloc[-5,:]=df_total_PK_org.iloc[-5,:],df_total_PK_org.iloc[-4,:]

                       df_total_PK_org=df_total_PK_org1

                       df_total_PK_org = df_total_PK_org.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')

                       table_heading='Фармакокинетические показатели ' + "("+file_name+")"
                       list_heading_word.append(table_heading)
                       
                       list_table_word.append(df_total_PK_org) 
                       #создание списков фреймов, названий органов и т.д.

                       ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                       df_concat = create_table_descriptive_statistics(df)['df_concat']

                       list_name_organs.append(file_name)
                       list_df_unrounded.append(df_concat_PK_org)
                       list_df_for_mean_unround_for_graphics.append(df_concat)

                ###Кнопка активации дальнейших действий
                button_calculation = False
                
                if (list_keys_file_org != []) and dose and measure_unit_org_blood and measure_unit_org_organs:
                 
                   condition_cmax1 =  len(list_cmax_1_org) == count_rows_number_org
                   
                   if st.session_state["agree_cmax2 - органы"] == True:
                      condition_cmax2 =  len(list_cmax_2_org) == count_rows_number_org
                   
                   if st.session_state["agree_cmax2 - органы"] == True:
                      if (condition_cmax2):
                         button_calculation = True
                   if st.session_state["agree_cmax2 - органы"] == False:
                      if (condition_cmax1):
                         button_calculation = True

                   if button_calculation == True:
                      st.write('👩🏽‍💻Расчеты произведены!')
                   else:   
                      st.write('🔧Заполните все поля ввода и загрузите файлы!')
                
                if (list_keys_file_org != []) and dose and measure_unit_org_blood and measure_unit_org_organs and button_calculation:
                   
                   if st.session_state["agree_cmax2 - органы"] == True:
                      list_list_PK_par_mean=[]
                      for i in list_df_unrounded: 
                          mean_сmax=i['Cmax'].loc['mean']
                          mean_tmax=i['Tmax'].loc['mean']
                          mean_сmax2=i['Cmax(2)'].loc['mean']
                          mean_tmax2=i['Tmax(2)'].loc['mean']
                          mean_mrt0inf=i['MRT0→∞'].loc['mean']
                          mean_thalf=i['T1/2'].loc['mean']
                          mean_auc0t=i['AUC0-t'].loc['mean']
                          mean_auc0inf=i['AUC0→∞'].loc['mean']
                          mean_aumc0inf=i['AUMC0-∞'].loc['mean']
                          mean_kel=i['Kel'].loc['mean']
                          list_list_PK_par_mean.append([mean_сmax,mean_tmax,mean_сmax2,mean_tmax2,mean_mrt0inf,mean_thalf,mean_auc0t,mean_auc0inf,mean_aumc0inf,mean_kel])
                   else:
                      list_list_PK_par_mean=[]
                      for i in list_df_unrounded: 
                          mean_сmax=i['Cmax'].loc['mean']
                          mean_tmax=i['Tmax'].loc['mean']
                          mean_mrt0inf=i['MRT0→∞'].loc['mean']
                          mean_thalf=i['T1/2'].loc['mean']
                          mean_auc0t=i['AUC0-t'].loc['mean']
                          mean_auc0inf=i['AUC0→∞'].loc['mean']
                          mean_aumc0inf=i['AUMC0-∞'].loc['mean']
                          mean_kel=i['Kel'].loc['mean']
                          list_list_PK_par_mean.append([mean_сmax,mean_tmax,mean_mrt0inf,mean_thalf,mean_auc0t,mean_auc0inf,mean_aumc0inf,mean_kel])

                   ### получение итогового фрейма ФК параметров органов
                   if st.session_state["agree_cmax2 - органы"] == True:
                      df_PK_organs_total = pd.DataFrame(list_list_PK_par_mean, columns =['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Kel'],index=list_name_organs)
                   else:
                      df_PK_organs_total = pd.DataFrame(list_list_PK_par_mean, columns =['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Kel'],index=list_name_organs) 
                   
                   df_PK_organs_total_transpose=df_PK_organs_total.transpose()

                   index_blood = list_name_organs.index("Кровь")
                   ###ft
                   list_aucot_for_ft=[]
                   list_columns_df_PK_organs_total_transpose=df_PK_organs_total_transpose.columns.tolist()
                   list_columns_df_PK_organs_total_transpose.remove('Кровь') #исходный файл крови должен быть назван так "Кровь"
                   for i in list_columns_df_PK_organs_total_transpose:
                       aucot=df_PK_organs_total_transpose[i].loc['AUC0-t']
                       list_aucot_for_ft.append(aucot)

                   list_ft=[] ## для диаграммы
                   list_ft_round=[]
                   for i in list_aucot_for_ft:
                       ft=i/df_PK_organs_total_transpose["Кровь"].loc['AUC0-t']
                       list_ft.append(ft)
                       list_ft_round.append("%.2f" % round(ft,2))
                   list_ft_round.insert(index_blood, "-")

                   df_PK_organs_total_transpose.loc[ len(df_PK_organs_total_transpose.index )] = list_ft_round

                   if st.session_state["agree_cmax2 - органы"] == True:
                      df_PK_organs_total_transpose.index=['Cmax ' +"("+measure_unit_org_blood+")",'Tmax ' +"("+"ч"+")",'Cmax(2) ' +"("+measure_unit_org_blood+")",'Tmax(2) ' +"("+"ч"+")",'MRT0→∞ '+"("+"ч"+")",'T1/2 '+"("+"ч"+")",'AUC0-t '+"("+measure_unit_org_blood+"×ч" +")",'AUC0→∞ '+"("+measure_unit_org_blood+"×ч" +")",'AUMC0-∞ '+"("+measure_unit_org_blood+"×ч\u00B2" +")",'Kel '+"("+"ч\u207B\u00B9"+")",'fт']
                   else:
                      df_PK_organs_total_transpose.index=['Cmax ' +"("+measure_unit_org_blood+")",'Tmax ' +"("+"ч"+")",'MRT0→∞ '+"("+"ч"+")",'T1/2 '+"("+"ч"+")",'AUC0-t '+"("+measure_unit_org_blood+"×ч" +")",'AUC0→∞ '+"("+measure_unit_org_blood+"×ч" +")",'AUMC0-∞ '+"("+measure_unit_org_blood+"×ч\u00B2" +")",'Kel '+"("+"ч\u207B\u00B9"+")",'fт']
                   
                   #округление фрейма df_PK_organs_total_transpose

                   df_organs_trans_trans=df_PK_organs_total_transpose.transpose()


                   series_Cmax=df_organs_trans_trans['Cmax ' +"("+measure_unit_org_blood+")"].tolist() 
                   series_Cmax=pd.Series(["%.2f" % round(v,2) for v in series_Cmax])

                   series_Tmax=df_organs_trans_trans['Tmax ' +"("+"ч"+")"].tolist()       
                   series_Tmax=pd.Series(["%.2f" % round(v,2) for v in series_Tmax]) 
                   
                   if st.session_state["agree_cmax2 - органы"] == True:
                      series_Cmax2=df_organs_trans_trans['Cmax(2) ' +"("+measure_unit_org_blood+")"].tolist() 
                      series_Cmax2=pd.Series(["%.2f" % round(v,2) for v in series_Cmax2])

                      series_Tmax2=df_organs_trans_trans['Tmax(2) ' +"("+"ч"+")"].tolist()       
                      series_Tmax2=pd.Series(["%.2f" % round(v,2) for v in series_Tmax2])

                   series_MRT0_inf= df_organs_trans_trans['MRT0→∞ '+"("+"ч"+")"].tolist()   
                   series_MRT0_inf=pd.Series(["%.3f" % round(v,3) for v in series_MRT0_inf])

                   series_half_live= df_organs_trans_trans['T1/2 '+"("+"ч"+")"].tolist()   
                   series_half_live=pd.Series(["%.2f" % round(v,2) for v in series_half_live]) 

                   series_AUC0_t= df_organs_trans_trans['AUC0-t '+"("+measure_unit_org_blood+"×ч" +")"].tolist()   
                   series_AUC0_t=pd.Series(["%.2f" % round(v,2) for v in series_AUC0_t])

                   series_AUC0_inf= df_organs_trans_trans['AUC0→∞ '+"("+measure_unit_org_blood+"×ч" +")"].tolist()  
                   series_AUC0_inf=pd.Series(["%.2f" % round(v,2) for v in series_AUC0_inf]) 

                   series_AUMC0_inf= df_organs_trans_trans['AUMC0-∞ '+"("+measure_unit_org_blood+"×ч\u00B2" +")"].tolist()   
                   series_AUMC0_inf=pd.Series(["%.2f" % round(v,2) for v in series_AUMC0_inf])
             
                   series_Kel= df_organs_trans_trans['Kel '+"("+"ч\u207B\u00B9"+")"].tolist()   
                   series_Kel=pd.Series(["%.4f" % round(v,4) for v in series_Kel])

                   series_ft= df_organs_trans_trans['fт'].tolist() ##уже округлен
                   series_ft=pd.Series(series_ft)
                   
                   if st.session_state["agree_cmax2 - органы"] == True:
                      df_total_total_organs = pd.concat([series_Cmax,series_Tmax,series_Cmax2,series_Tmax2,series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Kel,series_ft], axis= 1)
                   else:
                      df_total_total_organs = pd.concat([series_Cmax,series_Tmax,series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Kel,series_ft], axis= 1)

                   df_total_total_organs.index=df_PK_organs_total_transpose.columns.tolist()
                   df_total_total_organs.columns=df_PK_organs_total_transpose.index.tolist() 

                   df_total_total_organs_total= df_total_total_organs.transpose()
                   df_total_total_organs_total.index.name = 'Параметры, размерность'

                   table_heading='Фармакокинетические параметры в различных тканях'
                   list_heading_word.append(table_heading) 

                   list_table_word.append(df_total_total_organs_total) 

                   ###построение графика "Фармакокинетический профиль в органах"

                   ### в линейных координатах

                   list_list_mean_conc=[]
                   list_list_std_conc=[]
                   for i in list_df_for_mean_unround_for_graphics: 
                       mean_conc_list=i.loc['mean'].tolist()
                       std_conc_list=i.loc['std'].tolist()
                       list_list_mean_conc.append(mean_conc_list)
                       list_list_std_conc.append(std_conc_list)

                   list_name_organs_std=[]
                   for i in list_name_organs:
                    j= i + " std"
                    list_name_organs_std.append(j)

                   df_mean_conc_graph = pd.DataFrame(list_list_mean_conc, columns =list_t_graph[0],index=list_name_organs)
                   df_mean_conc_graph_1=df_mean_conc_graph.transpose()
                   df_std_conc_graph = pd.DataFrame(list_list_std_conc, columns =list_t_graph[0],index=list_name_organs_std)
                   df_std_conc_graph_1=df_std_conc_graph.transpose()
                   df_concat_mean_std= pd.concat([df_mean_conc_graph_1,df_std_conc_graph_1],sort=False,axis=1)

                   list_colors = [] ## генерация 500 цветов
                   for i in range(0,500):
                       hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                       list_colors.append(hexadecimal)
                   
                   list_zip_mean_std_colors=zip(list_name_organs,list_name_organs_std,list_colors)    

                   fig, ax = plt.subplots()
                   for i,j,c in list_zip_mean_std_colors:
                        plt.errorbar(list(df_concat_mean_std.index),df_concat_mean_std[i],yerr=df_concat_mean_std[j],color= c, marker='o',markersize=4.0,markeredgecolor=c,markerfacecolor=c,ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0,label=i)
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+ measure_unit_org_blood)
                        ax.legend(fontsize = 5)
                   
                   list_graphics_word.append(fig)

                   graphic='Сравнение фармакокинетических профилей (в линейных координатах) в органах'
                   list_heading_graphics_word.append(graphic)

                   ### в полулог. координатах

                   list_t_organs=list(df_concat_mean_std.index)
                   list_t_organs.remove(0)
                   df_concat_mean_std_without_0=df_concat_mean_std.drop([0])
                   list_zip_mean_std_colors=zip(list_name_organs,list_name_organs_std,list_colors)

                   fig, ax = plt.subplots()
                   for i,j,c in list_zip_mean_std_colors:
                        plt.errorbar(list_t_organs,df_concat_mean_std_without_0[i],yerr=df_concat_mean_std_without_0[j],color= c, marker='o',markersize=4.0,markeredgecolor=c,markerfacecolor=c,ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0,label=i)
                        ax.set_yscale("log")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+ measure_unit_org_blood)
                        ax.legend(fontsize = 5)
                   
                   list_graphics_word.append(fig)

                   graphic='Сравнение фармакокинетических профилей (в полулогарифмических координатах) в органах'
                   list_heading_graphics_word.append(graphic)

                   ###построение диаграммы для тканевой доступности

                   #list_zip_list_ft_list_name_organs=zip(list_ft,list_name_organs)
                   list_name_organs.remove("Кровь")

                   fig, ax = plt.subplots()

                   sns.barplot(x=list_name_organs, y=list_ft,color='blue',width=0.3)

                   plt.ylabel("Тканевая доступность")

                   ax.set_xticklabels(list_name_organs,fontdict={'fontsize': 6.0})

                   list_graphics_word.append(fig)
                   
                   graphic='Тканевая доступность в органах'
                   list_heading_graphics_word.append(graphic) 
                   

            ###сохранение состояния 
            st.session_state["list_heading_word"] = list_heading_word
            st.session_state["list_table_word"] = list_table_word
            st.session_state["list_graphics_word"] = list_graphics_word
            st.session_state["list_heading_graphics_word"] = list_heading_graphics_word


      #####Создание word отчета
         if panel == "Таблицы": 
            
            list_heading_word = st.session_state["list_heading_word"]
            list_table_word = st.session_state["list_table_word"]

            ###вызов функции создания таблицы
            create_table(list_heading_word,list_table_word)

         if panel == "Графики":

            list_graphics_word = st.session_state["list_graphics_word"]
            list_heading_graphics_word = st.session_state["list_heading_graphics_word"]

            ###вызов функции создания графика
            create_graphic(list_graphics_word,list_heading_graphics_word)
            
            #######визуализация

            #классификация графиков по кнопкам
            type_graphics = st.selectbox('Выберите вид графиков',
      ('Индивидуальные фармакокинетические профили', 'Сравнение индивидуальных фармакокинетических профилей', 'Графики усредненного фармакокинетического профиля', "Сравнение фармакокинетических профилей в различных органах", "Тканевая доступность в органах"),disabled = False, key = "Вид графика - ИО" )

            count_graphics_for_visual = len(list_heading_graphics_word)
            list_range_count_graphics_for_visual = range(0,count_graphics_for_visual)
            
            for i in list_range_count_graphics_for_visual:
                if list_heading_graphics_word[i].__contains__("индивидуального"): 
                   if type_graphics == 'Индивидуальные фармакокинетические профили':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Сравнение индивидуальных"):   
                   if type_graphics == 'Сравнение индивидуальных фармакокинетических профилей':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("усредненного"):
                   if type_graphics == 'Графики усредненного фармакокинетического профиля':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Сравнение фармакокинетических"):
                   if type_graphics == 'Сравнение фармакокинетических профилей в различных органах':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Тканевая"):
                   if type_graphics == 'Тканевая доступность в органах':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])

################################################################################################

   if option == 'Линейность дозирования':
      
      st.title('Исследование линейности дозирования')
      
      col1, col2 = st.columns([0.66, 0.34])

      ######### боковое меню справа
      with col2:
           selected = option_menu(None, ["Включение параметров в исследование"], 
           icons=['menu-button'], 
           menu_icon="cast", default_index=0, orientation="vertical",
           styles={
               "container": {"padding": "0!important", "background-color": "#24769C"},
               "icon": {"color": "#5DAED3", "font-size": "13px"}, 
               "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
               "nav-link-selected": {"background-color": "#335D70"},
           })

           if selected == "Включение параметров в исследование":
              type_parameter = st.selectbox('Выберите параметр',
           ('Cmax(2)',"Вид введения"),disabled = False, key = "Вид параметра - линейность")
              

           if type_parameter == 'Cmax(2)':
              
              if "agree_cmax2 - линейность" not in st.session_state:
                 st.session_state["agree_cmax2 - линейность"] = False

              st.session_state["agree_cmax2 - линейность"] = st.checkbox('Добавить возможность выбора Cmax(2)', key = "Возможность добавления Cmax2 - линейность", value = st.session_state["agree_cmax2 - линейность"])
              
              if st.session_state["agree_cmax2 - линейность"] == True:
                 st.write('🧠Параметр добавлен!')

      with col1:

         panel = st.radio(
              "⚙️Панель управления",
              ("Загрузка файлов", "Таблицы","Графики"),
              horizontal=True, key= "Загрузка файлов - Исследование ФК параметров для линейности дозирования"
          )
         
         ###создание состояния
         if "measure_unit_lin" not in st.session_state:
            st.session_state["measure_unit_lin"] = ""
         if "measure_unit_dose_lin" not in st.session_state:
            st.session_state["measure_unit_dose_lin"] = ""

         #cписки для word-отчета
         list_heading_word=[]
         list_table_word=[]
         list_graphics_word=[]
         list_heading_graphics_word=[]

         if panel == "Загрузка файлов":

            measure_unit_lin = st.text_input("Введите единицы измерения концентрации", key="Единицы измерения концентрации при исследовании линейности дозирования", value = st.session_state["measure_unit_lin"])
            st.session_state["measure_unit_lin"] = measure_unit_lin

            measure_unit_dose_lin = st.text_input("Введите единицы измерения дозировки", key="Единицы измерения дозировки при исследовании линейности дозирования", value = st.session_state["measure_unit_dose_lin"])
            st.session_state["measure_unit_dose_lin"] = measure_unit_dose_lin 

            #cостояние радио-кнопки "method_auc"
            if "index_method_auc - ЛД" not in st.session_state:
                st.session_state["index_method_auc - ЛД"] = 0

            method_auc = st.radio("📌Метод подсчёта AUC0-t",('linear',"linear-up/log-down"),key = "Метод подсчёта AUC0-t - ЛД", index = st.session_state["index_method_auc - ЛД"])
            
            if st.session_state["Метод подсчёта AUC0-t - ЛД"] == 'linear':
               st.session_state["index_method_auc - ЛД"] = 0
            if st.session_state["Метод подсчёта AUC0-t - ЛД"] == "linear-up/log-down":
               st.session_state["index_method_auc - ЛД"] = 1

            file_uploader = st.file_uploader("Выберите нужное количество файлов соответственно количеству исследуемых дозировок(не менее 3-х файлов); файл должен быть назван соотвественно своей дозировке, например: 'Дозировка 50'. Если дозировка предcтавляет из себя дробное число, дробь писать через '.' Слово 'Дозировка' в верхнем регистре обязательно",accept_multiple_files=True, key='Файлы при исследовании линейности дозирования')
            
            ###сохранение файла
            if file_uploader is not None:
               for i in file_uploader:
                   save_uploadedfile(i)
                   st.session_state[str(i.name)] = i.name
            
            list_keys_file_lin = []
            for i in st.session_state.keys():
                if i.__contains__("xlsx") and i.__contains__("Дозировка") and (not i.__contains__("edited_df")): ###слово дозировка нужно, чтобы отличать файлы от других xlsx органов, т.к там тоже ключи имя файла; #обрезаем фразу ненужного добавления названия "edited_df"
                   list_keys_file_lin.append(i)

            if (list_keys_file_lin != []) and measure_unit_lin and measure_unit_dose_lin:

                list_name_doses=[]
                list_df_unrounded=[]
                list_df_for_mean_unround_for_graphics=[]
                list_t_graph=[]

                for i in list_keys_file_lin:
                    df = pd.read_excel(os.path.join("Папка для сохранения файлов",i))

                    file_name=i[10:-5]

                    st.subheader('Индивидуальные значения концентраций в дозировке ' +file_name+" "+ measure_unit_lin)
                    
                    ###интерактивная таблица
                    df = edit_frame(df,i)

                    ###количество животных 
                    count_rows_number_lin= len(df.axes[0])

                    table_heading='Индивидуальные и усредненные значения концентраций в дозировке ' +file_name+" "+ measure_unit_lin
                    list_heading_word.append(table_heading)

                    ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                    df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']

                    list_table_word.append(df_concat_round_str_transpose)
                    ########### графики    
                    
                    
                    ######индивидуальные    

                    # в линейных координатах
                    col_mapping = df.columns.tolist()
                    col_mapping.remove('Номер')

                    count_row_df = len(df.axes[0])

                    list_time = []
                    for i in col_mapping:
                        numer=float(i)
                        list_time.append(numer)
                    list_t_graph.append(list_time) 

                    for r in range(0,count_row_df):

                        list_concentration=df.iloc[r].tolist()

                        numer_animal=list_concentration[0]

                        list_concentration.pop(0) #удаление номера животного

                        list_concentration = [float(v) for v in list_concentration]


                        fig, ax = plt.subplots()
                        plt.plot(list_time,list_concentration,marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+measure_unit_lin)
                        
                        list_graphics_word.append(fig)
           
                        graphic='График индивидуального фармакокинетического профиля в линейных координатах в дозировке '  +file_name+" "+ measure_unit_lin+',  '+numer_animal
                        list_heading_graphics_word.append(graphic) 

                     #в полулогарифмических координатах методом удаления точек
                        count_for_0_1=len(list_concentration)
                        list_range_for_0_1=range(0,count_for_0_1)

                        list_time_0=[]
                        list_for_log_1=[]
                        for i in list_range_for_0_1:
                            if list_concentration[i] !=0:
                               list_for_log_1.append(list_concentration[i])
                               list_time_0.append(list_time[i]) 

                        fig, ax = plt.subplots()
                        plt.plot(list_time_0,list_for_log_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue")
                        ax.set_yscale("log")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+measure_unit_lin)

                        
                        list_graphics_word.append(fig)
                        
                        graphic='График индивидуального фармакокинетического профиля в полулогарифмических координатах в дозировке ' +file_name+" "+ measure_unit_lin+',  '+numer_animal
                        list_heading_graphics_word.append(graphic) 

                 # объединенные индивидуальные в линейных координатах

                    df_for_plot_conc=df.drop(['Номер'], axis=1)
                    df_for_plot_conc_1 = df_for_plot_conc.transpose()
                    list_numer_animal_for_plot=df['Номер'].tolist()
                    count_numer_animal = len(list_numer_animal_for_plot) ### для регулирования пропорции легенды
                    list_color = [] ## генерация 500 цветов
                    for i in range(0,500):
                        hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                        list_color.append(hexadecimal)

                    fig, ax = plt.subplots()

                    ax.set_prop_cycle(cycler(color=list_color))

                    plt.plot(df_for_plot_conc_1,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                    ax.set_xlabel("Время, ч")
                    ax.set_ylabel("Концентрация, "+measure_unit_lin)
                    if count_numer_animal > 20:
                       ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                    else:
                       ax.legend(bbox_to_anchor=(1, 1))
                    
                    list_graphics_word.append(fig)
                    
                    graphic="Сравнение индивидуальных фармакокинетических профилей в линейных координатах в дозировке " +file_name+" "+ measure_unit_lin
                    list_heading_graphics_word.append(graphic) 
             
                 # объединенные индивидуальные в полулогарифмических координатах методом замены 0 на None
                    df_for_plot_conc_1_log=df_for_plot_conc_1.replace(0, None)


                    fig, ax = plt.subplots()

                    ax.set_prop_cycle(cycler(color=list_color))

                    plt.plot(df_for_plot_conc_1_log,marker='o',markersize=4.0,label = list_numer_animal_for_plot)

                    ax.set_xlabel("Время, ч")
                    ax.set_ylabel("Концентрация, "+measure_unit_lin)
                    ax.set_yscale("log")
                    if count_numer_animal > 20:
                       ax.legend(fontsize=(160/count_numer_animal),bbox_to_anchor=(1, 1))
                    else:
                       ax.legend(bbox_to_anchor=(1, 1))
                    
                    list_graphics_word.append(fig)
                    
                    graphic="Сравнение индивидуальных фармакокинетических профилей в полулогарифмических координатах в дозировке " +file_name+" "+ measure_unit_lin
                    list_heading_graphics_word.append(graphic) 
                     ###усредненные    
                 # в линейных координатах
                    list_time = []
                    for i in col_mapping:
                        numer=float(i)
                        list_time.append(numer)

                    df_averaged_concentrations=df.describe()
                    list_concentration=df_averaged_concentrations.loc['mean'].tolist()
                    err_y_1=df_averaged_concentrations.loc['std'].tolist()


                    fig, ax = plt.subplots()
                    plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                    plt.xlabel("Время, ч")
                    plt.ylabel("Концентрация, "+measure_unit_lin)
                     
                    list_graphics_word.append(fig)
                    
                    graphic='График усредненного фармакокинетического профиля в линейных координатах в дозировке ' +file_name+" "+ measure_unit_lin
                    list_heading_graphics_word.append(graphic)



                 #в полулогарифмических координатах
                    #для полулогарифм. посторим без нуля
                    list_time.remove(0)
                    list_concentration.remove(0)
                    err_y_1.remove(0) 


                    fig, ax = plt.subplots()
                    plt.errorbar(list_time,list_concentration,yerr=err_y_1, marker='o',markersize=4.0,markeredgecolor="blue",markerfacecolor="blue",ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0)
                    ax.set_yscale("log")
                    plt.xlabel("Время, ч")
                    plt.ylabel("Концентрация, "+measure_unit_lin)

                    list_graphics_word.append(fig)
                    
                    graphic='График усредненного фармакокинетического профиля в полулогарифмических координатах ' +file_name+" "+ measure_unit_lin
                    list_heading_graphics_word.append(graphic)

                    ############ Параметры ФК

                    df_without_numer=df.drop(['Номер'],axis=1)
                    count_row=df_without_numer.shape[0]

                    list_count_row=range(count_row)
          
                    ###Cmax
                    #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (выкл)
                    if st.session_state["agree_cmax2 - линейность"] == False:
                       list_cmax_1_lin=[]
                       for i in range(0,count_row):
                           cmax=float(max(df_without_numer.iloc[[i]].iloc[0].tolist()))
                           list_cmax_1_lin.append(cmax)

                    #выбор метода подсчета Сmax в зависимости от надобности Cmax2 (вкл)
                    if st.session_state["agree_cmax2 - линейность"] == True:

                       ###создание состояния
                       if ("selected_value_lin" + file_name) not in st.session_state:
                          st.session_state["selected_value_lin"+ file_name] = []
                       
                       if ("feature_disable_selected_value_lin" + file_name) not in st.session_state:
                           st.session_state["feature_disable_selected_value_lin" + file_name] = True

                       ###создание состояния
                       st.info('Выбери Cmax:')
                       list_columns_without_numer = df.columns.tolist()
                       list_columns_without_numer.remove('Номер')
                       selected_columns = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax линейность дозирования ' + file_name, max_selections=1)
                       st.session_state["selected_columns_lin"+ file_name] = selected_columns 

                       list_keys_cmax = st.session_state["selected_value_lin"+ file_name]
                       if selected_columns != [] and st.session_state["feature_disable_selected_value_lin"+ file_name]:
                          selected_value = st.multiselect('Выбери значение концентрации:', df[selected_columns], key='Выбери значение концентрации Cmax линейность дозирования ' + file_name, max_selections=1)
                          list_keys_cmax.append(selected_value)

                       if list_keys_cmax != []:
                          st.session_state["selected_value_lin"+ file_name] = list_keys_cmax

                       list_keys_cmax = st.session_state["selected_value_lin"+ file_name]
                       list_keys_cmax_sample = [item for sublist in list_keys_cmax for item in sublist]
                       
                       if st.button('Очистить список Cmax', key="Очистка списка Cmax линейность дозирования " + file_name):
                          del st.session_state["selected_value_lin"+ file_name]
                          list_keys_cmax_sample = []
                          selected_columns = st.session_state["selected_columns_lin"+ file_name]
                          st.session_state["feature_disable_selected_value_lin"+ file_name] = True
                       
                       st.write("Список Cmax:")
                       st.write(list_keys_cmax_sample)
                       
                       list_cmax_1_lin =list_keys_cmax_sample 
                       
                       list_cmax_2_lin =[]

                    if len(list_cmax_1_lin) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - линейность"] == True):
                       
                       st.session_state["feature_disable_selected_value_lin"+ file_name] = False

                       ######Cmax2

                       if ("feature_disable_selected_value_lin_2"+ file_name) not in st.session_state:
                        st.session_state["feature_disable_selected_value_lin_2"+ file_name] = True

                       st.info('Выбери Cmax(2):')
                       
                       selected_columns_2 = st.multiselect('Выбери временную точку:', list_columns_without_numer, key='Выбери временную точку Cmax2 линейность дозирования ' + file_name, max_selections=1)
                       st.session_state["selected_columns_2_lin"+ file_name] = selected_columns_2

                       ###создание состояния
                       if ("selected_value_2_lin"+ file_name) not in st.session_state:
                          st.session_state["selected_value_2_lin"+ file_name] = []

                       list_keys_cmax_2 = st.session_state["selected_value_2_lin"+ file_name]
                       if selected_columns_2 != [] and st.session_state["feature_disable_selected_value_lin_2"+ file_name]:
                          selected_value_2 = st.multiselect('Выбери значение концентрации:', df[selected_columns_2], key='Выбери значение концентрации Cmax2 линейность дозирования '  + file_name, max_selections=1)
                          list_keys_cmax_2.append(selected_value_2)

                       if list_keys_cmax_2 != []:
                          st.session_state["selected_value_2_lin"+ file_name] = list_keys_cmax_2

                       list_keys_cmax_2 = st.session_state["selected_value_2_lin"+ file_name]
                       list_keys_cmax_sample_2 = [item for sublist in list_keys_cmax_2 for item in sublist]

                       if st.button('Очистить список Cmax(2)', key="Очистка списка Cmax(2) линейность дозирования " + file_name):
                          del st.session_state["selected_value_2_lin"+ file_name]
                          list_keys_cmax_sample_2 = []
                          selected_columns_2 = st.session_state["selected_columns_2_lin"+ file_name]
                          st.session_state["feature_disable_selected_value_lin_2"+ file_name] = True

                       st.write("Список Cmax(2):")
                       st.write(list_keys_cmax_sample_2)

                       list_cmax_2_lin = list_keys_cmax_sample_2

                       if len(list_cmax_2_lin) == len(df.index.tolist()):
                          st.session_state["feature_disable_selected_value_lin_2"+ file_name] = False
                    
                    if (len(list_cmax_1_lin) == len(df.index.tolist())):
                       
                       ###Tmax   
                       list_Tmax_1=[]
                       for cmax in list_cmax_1_lin:
                           for column in df.columns:
                               for num, row in df.iterrows():
                                   if df.iloc[num][column] == cmax:
                                      list_Tmax_1.append(f"{column}")
                     
                       list_Tmax_float_1=[]           
                       for i in list_Tmax_1:
                           Tmax=float(i)
                           list_Tmax_float_1.append(Tmax)

                    if (len(list_cmax_1_lin) == len(df.index.tolist())) and (st.session_state["agree_cmax2 - линейность"] == True):
                       
                       list_Tmax_2=[]
                       for cmax in list_cmax_2_lin:
                           for column in df.columns:
                               for num, row in df.iterrows():
                                   if df.iloc[num][column] == cmax:
                                      list_Tmax_2.append(f"{column}")
                     
                       list_Tmax_float_2=[]           
                       for i in list_Tmax_2:
                           Tmax=float(i)
                           list_Tmax_float_2.append(Tmax)  

                    if (len(list_cmax_1_lin) == len(df.index.tolist())):
                       
                       ###AUC0-t
                       list_AUC_0_T=[]
                       if method_auc == 'linear':
                          for i in range(0,count_row):
                              list_columns_T=[]
                              for column in df_without_numer.columns:
                                  list_columns_T.append(float(column))
                              list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                              ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                              cmax = max(list_concentration)
                              index_cmax = list_concentration.index(cmax)
                              list_before_cmax = list_concentration[0:index_cmax]
                              list_after_cmax = list_concentration[index_cmax:]
                              list_before_cmax_t = list_columns_T[0:index_cmax]
                              list_after_cmax_t = list_columns_T[index_cmax:]

                              count_list_concentration = len(list_after_cmax)
                              list_range_for_remove_0 = range(0,count_list_concentration)

                              list_conc_without_0=[]
                              list_t_without_0=[]
                              for i in list_range_for_remove_0:
                                  if list_after_cmax[i] !=0:
                                     list_conc_without_0.append(list_after_cmax[i])
                                     list_t_without_0.append(list_after_cmax_t[i])

                              list_concentration = list_before_cmax + list_conc_without_0
                              list_columns_T = list_before_cmax_t + list_t_without_0
                              ######################

                              AUC_0_T=np.trapz(list_concentration,x=list_columns_T)
                              list_AUC_0_T.append(AUC_0_T)

                       if method_auc == 'linear-up/log-down':
                          for i in range(0,count_row):
                              list_columns_T=[]
                              for column in df_without_numer.columns:
                                  list_columns_T.append(float(column))
                              list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                              ###удаление всех нулей сзади массива, т.к. AUC0-t это AUClast (до последней определяемой точки, а не наблюдаемой)
                              cmax = max(list_concentration)
                              index_cmax = list_concentration.index(cmax)
                              list_before_cmax = list_concentration[0:index_cmax]
                              list_after_cmax = list_concentration[index_cmax:]
                              list_before_cmax_t = list_columns_T[0:index_cmax]
                              list_after_cmax_t = list_columns_T[index_cmax:]

                              count_list_concentration = len(list_after_cmax)
                              list_range_for_remove_0 = range(0,count_list_concentration)

                              list_conc_without_0=[]
                              list_t_without_0=[]
                              for i in list_range_for_remove_0:
                                  if list_after_cmax[i] !=0:
                                     list_conc_without_0.append(list_after_cmax[i])
                                     list_t_without_0.append(list_after_cmax_t[i])

                              list_concentration = list_before_cmax + list_conc_without_0
                              list_columns_T = list_before_cmax_t + list_t_without_0
                              ######################
                              
                              list_c = list_concentration
                              list_t = list_columns_T
                              
                              count_i = len(list_c)
                              list_range= range(0,count_i)
                              
                              list_AUC_0_T_ascending=[]
                              list_AUC_0_T_descending = []
                              AUC_0_T_ascending=0
                              AUC_0_T_descending = 0
                              a=0
                              a1=0
                              d=0
                              d1=0
                              for i in list_range:
                                  if a1<count_i-1:
                                     if list_c[i+1] > list_c[i]:
                                        if a<count_i-1:
                                            AUC_0_T_ascending += ((list_c[i]+list_c[i+1])*(list_t[i+1]-list_t[i]))/2
                                            a+=1
                                            list_AUC_0_T_ascending.append(AUC_0_T_ascending)
                                  if d1<count_i-1:
                                     if list_c[i+1] < list_c[i]:      
                                        if d<count_i-1:
                                            AUC_0_T_descending+=(list_t[i+1]-list_t[i])/(np.log(np.asarray(list_c[i])/np.asarray(list_c[i+1]))) *(list_c[i]-list_c[i+1])
                                            d+=1
                                            list_AUC_0_T_descending.append(AUC_0_T_descending)
                                     a1+=1
                                     d1+=1
               
                              AUC_O_T = list_AUC_0_T_ascending[-1]+list_AUC_0_T_descending[-1]
                              
                              list_AUC_0_T.append(AUC_O_T)

                       ####Сmax/AUC0-t
                       list_Сmax_division_AUC0_t_for_division=zip(list_cmax_1_lin,list_AUC_0_T)
                       list_Сmax_division_AUC0_t=[]
                       for i,j in list_Сmax_division_AUC0_t_for_division:
                               list_Сmax_division_AUC0_t.append(i/j)


                       ####KEL
                       list_kel_total=[]
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                           list_concentration.remove(0)
                           list_c=list_concentration

                           list_time=df_without_numer.columns.tolist()
                           list_time.remove(0) 

                           list_t=[]
                           for i in list_time:
                               i=float(i)
                               list_t.append(i)

                           #срез_без_cmax
                           max_value_c=max(list_c)
                           index_cmax=list_c.index(max_value_c)

                           list_c_without_cmax=list_c[index_cmax+1:]
                           list_t_without_cmax=list_t[index_cmax+1:]

                           #удаление всех нулей из массивов
                           count_for_0_1=len(list_c_without_cmax)
                           list_range_for_0_1=range(0,count_for_0_1)

                           list_time_0=[]
                           list_conc_0=[]
                           for i in list_range_for_0_1:
                               if list_c_without_cmax[i] !=0:
                                  list_conc_0.append(list_c_without_cmax[i])
                                  list_time_0.append(list_t_without_cmax[i]) 
                           ################################

                           n_points=len(list_conc_0)
                           list_n_points = range(0,n_points)

                           #создание списков с поочередно уменьщающемся кол, точек
                           list_for_kel_c=[]
                           for j in list_n_points:
                               if j<n_points:
                                  list_c_new=list_conc_0[j:n_points]
                                  list_for_kel_c.append(list_c_new)
                           list_for_kel_c.pop(-1) #удаление списка с одной точкой
                           list_for_kel_c.pop(-1)  #удаление списка с двумя точками     

                           list_for_kel_t=[]
                           for j in list_n_points:
                               if j<n_points:
                                  list_t_new=list_time_0[j:n_points]
                                  list_for_kel_t.append(list_t_new)
                           list_for_kel_t.pop(-1) #удаление списка с одной точкой
                           list_for_kel_t.pop(-1) #удаление списка с двумя точками 

                           list_ct_zip=zip(list_for_kel_c,list_for_kel_t)

                           list_kel=[]
                           list_r=[]
                           for i,j in list_ct_zip:

                               n_points_r=len(i)

                               np_c=np.asarray(i)
                               np_t_1=np.asarray(j).reshape((-1,1))

                               np_c_log=np.log(np_c)

                               model = LinearRegression().fit(np_t_1,np_c_log)

                               np_t=np.asarray(j)
                               a=np.corrcoef(np_t, np_c_log)
                               cor=((a[0])[1])
                               r_sq=cor**2

                               adjusted_r_sq=1-((1-r_sq)*((n_points_r-1))/(n_points_r-2))

                               ########################################
                               kel=abs(model.coef_[0])
                               list_kel.append(kel)
                               list_r.append(adjusted_r_sq)

                           #делаем срезы списоков до rmax
                           max_r=max(list_r)

                           index_max_r= list_r.index(max_r)

                           list_r1=list_r
                           list_kel1=list_kel

                           number_elem_list_r1=len(list_r1)

                           list_range_kel=range(0,number_elem_list_r1) 

                           list_kel_total_1=[]
                           for i in list_range_kel:

                               if abs(list_r[index_max_r] - list_r1[i]) < 0.0001: #проверяем все точки слева и справа от rmax
                                  list_kel_total.append(list_kel1[i]*math.log(math.exp(1))) #отдаю предпочтение rmax с большим количеством точек
                                  break #самая ранняя удовлетовряющая условию

                           for i in list_kel_total_1:
                               list_kel_total.append(i) 


                       ####T1/2
                       list_half_live=[]
                       for i in list_kel_total:
                           half_live=math.log(2)/i
                           list_half_live.append(half_live)


                       ###AUC0-inf 

                       list_auc0_inf=[] 

                       list_of_list_c=[]
                       for i in range(0,count_row):
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()
                           list_concentration.remove(0)
                           list_c = list_concentration
                           list_c.reverse() ### переворачиваем, для дальнейшей итерации с конца списка и поиска Clast не равное нулю
                           list_of_list_c.append(list_c)

                       list_zip_c_AUCt_inf=zip(list_kel_total,list_of_list_c)

                           #AUCt-inf 
                       list_auc_t_inf=[]     
                       for i,j in list_zip_c_AUCt_inf:
                           for clast in j:
                               if clast != 0:
                                  clast_true=clast
                                  break
                           auc_t_inf=clast_true/i
                           list_auc_t_inf.append(auc_t_inf)

                       list_auc_t_inf_and_AUC_0_T_zip=zip(list_AUC_0_T,list_auc_t_inf)

                       for i,j in list_auc_t_inf_and_AUC_0_T_zip:
                           auc0_inf=i+j    
                           list_auc0_inf.append(auc0_inf)


                       ####CL
                       list_cl=[]
                       
                       for i in list_auc0_inf:
                           cl = float(file_name)/i * 1000
                           list_cl.append(cl)


                       ####Vd
                       list_Vd=[]

                       list_zip_kel_cl=zip(list_kel_total,list_cl)

                       for i,j in list_zip_kel_cl:
                           Vd = j/i
                           list_Vd.append(Vd)


                       ###AUMC
                       list_AUMCO_inf=[]

                       list_AUMC0_t=[]

                       list_C_last=[]
                       list_T_last=[]
                       for i in range(0,count_row):
                           list_columns_T=[]
                           for column in df_without_numer.columns:
                               list_columns_T.append(float(column))
                           list_concentration=df_without_numer.iloc[[i]].iloc[0].tolist()

                           ###удаление всех нулей сзади массива, т.к. AUMC0-t это AUMClast (до последней определяемой точки, а не наблюдаемой)
                           cmax = max(list_concentration)
                           index_cmax = list_concentration.index(cmax)
                           list_before_cmax = list_concentration[0:index_cmax]
                           list_after_cmax = list_concentration[index_cmax:]
                           list_before_cmax_t = list_columns_T[0:index_cmax]
                           list_after_cmax_t = list_columns_T[index_cmax:]

                           count_list_concentration = len(list_after_cmax)
                           list_range_for_remove_0 = range(0,count_list_concentration)

                           list_conc_without_0=[]
                           list_t_without_0=[]
                           for i in list_range_for_remove_0:
                               if list_after_cmax[i] !=0:
                                  list_conc_without_0.append(list_after_cmax[i])
                                  list_t_without_0.append(list_after_cmax_t[i])

                           list_concentration = list_before_cmax + list_conc_without_0
                           list_columns_T = list_before_cmax_t + list_t_without_0
                           ######################

                           list_C_last.append(list_concentration[-1]) 
                           list_T_last.append(list_columns_T[-1]) 

                           list_len=len(list_concentration)

                           list_aumc_i=[]
                           for i in range(0,list_len):
                               AUMC=(list_columns_T[i] - list_columns_T[i-1]) *  ((list_concentration[i] * list_columns_T[i] + list_concentration[i-1] * list_columns_T[i-1])/2)
                               list_aumc_i.append(AUMC)

                           list_aumc_i.pop(0)

                           a=0
                           list_AUMC0_t_1=[]
                           for i in list_aumc_i:
                               a+=i
                               list_AUMC0_t_1.append(a)
                           list_AUMC0_t.append(list_AUMC0_t_1[-1])

                       list_zip_for_AUMC_inf=zip(list_kel_total,list_C_last,list_T_last)

                       list_AUMCt_inf=[]
                       for k,c,t in list_zip_for_AUMC_inf:
                           AUMCt_inf=c*t/k+c/(k*k)
                           list_AUMCt_inf.append(AUMCt_inf)


                       list_AUMC_zip=zip(list_AUMC0_t,list_AUMCt_inf)

                       for i,j in list_AUMC_zip:
                           AUMCO_inf=i+j
                           list_AUMCO_inf.append(AUMCO_inf)

                       ###MRT0-inf
                       list_MRT0_inf=[]

                       list_zip_AUMCO_inf_auc0_inf = zip(list_AUMCO_inf,list_auc0_inf)

                       for i,j in list_zip_AUMCO_inf_auc0_inf:
                           MRT0_inf=i/j
                           list_MRT0_inf.append(MRT0_inf)
                    
                    if st.session_state["agree_cmax2 - линейность"] == True:
                       #####Cmax условие для дальнейшего кода
                       if len(list_cmax_1_lin) == len(df.index.tolist()) and len(list_cmax_2_lin) == len(df.index.tolist()):
                       
                          ##################### Фрейм ФК параметров

                          ### пользовательский индекс
                          list_for_index=df["Номер"].tolist()
                          df_PK=pd.DataFrame(list(zip(list_cmax_1_lin,list_Tmax_float_1,list_cmax_2_lin,list_Tmax_float_2,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','Cmax(2)','Tmax(2)','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd/F'],index=list_for_index) 
                    
                    if len(list_cmax_1_lin) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - линейность"] == False):

                       ##################### Фрейм ФК параметров

                       ### пользовательский индекс
                       list_for_index=df["Номер"].tolist()
                       df_PK=pd.DataFrame(list(zip(list_cmax_1_lin,list_Tmax_float_1,list_MRT0_inf,list_half_live,list_AUC_0_T,list_auc0_inf,list_AUMCO_inf,list_Сmax_division_AUC0_t,list_kel_total,list_cl,list_Vd)),columns=['Cmax','Tmax','MRT0→∞','T1/2','AUC0-t','AUC0→∞','AUMC0-∞','Сmax/AUC0-t','Kel','CL/F','Vd/F'],index=list_for_index)

                    checking_condition_cmax2 = False

                    if st.session_state["agree_cmax2 - линейность"] == True:
                     
                       checking_condition_cmax2 = len(list_cmax_1_lin) == len(df.index.tolist()) and len(list_cmax_2_lin) == len(df.index.tolist()) and st.session_state["agree_cmax2 - линейность"] == True

                    if checking_condition_cmax2 or (len(list_cmax_1_lin) == len(df.index.tolist()) and (st.session_state["agree_cmax2 - линейность"] == False)):

                       ###описательная статистика

                       col_mapping_PK = df_PK.columns.tolist()

                       list_gmean_PK=[]

                       list_cv_PK=[] 

                       for i in col_mapping_PK:

                           list_ser_PK=df_PK[i].tolist()

                           def g_mean(list_ser_PK):
                               a=np.log(list_ser_PK)
                               return np.exp(a.mean())
                           Gmean_PK=g_mean(list_ser_PK)
                           list_gmean_PK.append(Gmean_PK)

                           cv_std_PK=lambda x: np.std(x, ddof= 1 )
                           cv_mean_PK=lambda x: np.mean(x)

                           CV_std_PK=cv_std_PK(list_ser_PK)
                           CV_mean_PK=cv_mean_PK(list_ser_PK)

                           CV_PK=(CV_std_PK/CV_mean_PK * 100)
                           list_cv_PK.append(CV_PK)


                       df_averaged_concentrations_PK=df_PK.describe()
                       df_averaged_concentrations_1_PK= df_averaged_concentrations_PK.drop(['count', '25%','75%'],axis=0)
                       df_averaged_concentrations_2_PK= df_averaged_concentrations_1_PK.rename(index={"50%": "median"})
                       df_averaged_concentrations_2_PK.loc[len(df_averaged_concentrations_2_PK.index )] = list_gmean_PK
                       df_averaged_3_PK = df_averaged_concentrations_2_PK.rename(index={5 : "Gmean"})
                       df_round_without_CV_PK=df_averaged_3_PK
                       df_round_without_CV_PK.loc[len(df_round_without_CV_PK.index )] = list_cv_PK
                       df_averaged_3_PK = df_round_without_CV_PK.rename(index={6 : "CV, %"})


                       df_concat_PK_lin= pd.concat([df_PK,df_averaged_3_PK],sort=False,axis=0)


                       ###округление описательной статистики и ФК параметров

                       series_Cmax=df_concat_PK_lin['Cmax']
                       list_Cmax_str_f=["%.2f" % round(v,2) for v in series_Cmax.tolist()]
                       series_Cmax=pd.Series(list_Cmax_str_f, index = df_concat_PK_lin.index.tolist(), name='Cmax ' +"("+measure_unit_lin +")")

                       if st.session_state["agree_cmax2 - линейность"] == True:
                          series_Cmax_2=df_concat_PK_lin['Cmax(2)']
                          list_Cmax_str_f_2=["%.2f" % round(v,2) for v in series_Cmax_2.tolist()]
                          series_Cmax_2=pd.Series(list_Cmax_str_f_2, index = df_concat_PK_lin.index.tolist(), name='Cmax(2) ' +"("+measure_unit_lin +")")

                       series_Tmax=df_concat_PK_lin['Tmax']
                       list_Tmax_str_f=["%.2f" % round(v,2) for v in series_Tmax.tolist()]
                       series_Tmax=pd.Series(list_Tmax_str_f, index = df_concat_PK_lin.index.tolist(), name='Tmax ' +"("+"ч"+")")

                       if st.session_state["agree_cmax2 - линейность"] == True:
                          series_Tmax_2=df_concat_PK_lin['Tmax(2)']
                          list_Tmax_str_f_2=["%.2f" % round(v,2) for v in series_Tmax_2.tolist()]
                          series_Tmax_2=pd.Series(list_Tmax_str_f_2, index = df_concat_PK_lin.index.tolist(), name='Tmax(2) ' +"("+"ч"+")")

                       series_MRT0_inf= df_concat_PK_lin['MRT0→∞']
                       list_MRT0_inf_str_f=["%.3f" % round(v,3) for v in series_MRT0_inf.tolist()]
                       series_MRT0_inf=pd.Series(list_MRT0_inf_str_f, index = df_concat_PK_lin.index.tolist(), name='MRT0→∞ '+"("+"ч"+")")

                       series_half_live= df_concat_PK_lin['T1/2']
                       list_half_live_str_f=["%.2f" % round(v,2) for v in series_half_live.tolist()]
                       series_half_live=pd.Series(list_half_live_str_f, index = df_concat_PK_lin.index.tolist(), name='T1/2 '+"("+"ч"+")")

                       series_AUC0_t= df_concat_PK_lin['AUC0-t']
                       list_AUC0_t_str_f=["%.2f" % round(v,2) for v in series_AUC0_t.tolist()]
                       series_AUC0_t=pd.Series(list_AUC0_t_str_f, index = df_concat_PK_lin.index.tolist(), name='AUC0-t '+"("+measure_unit_lin +"×ч" +")")

                       series_AUC0_inf= df_concat_PK_lin['AUC0→∞']
                       list_AUC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUC0_inf.tolist()]
                       series_AUC0_inf=pd.Series(list_AUC0_inf_str_f, index = df_concat_PK_lin.index.tolist(), name='AUC0→∞ '+"("+measure_unit_lin +"×ч" +")")

                       series_AUMC0_inf= df_concat_PK_lin['AUMC0-∞']
                       list_AUMC0_inf_str_f=["%.2f" % round(v,2) for v in series_AUMC0_inf.tolist()]
                       series_AUMC0_inf=pd.Series(list_AUMC0_inf_str_f, index = df_concat_PK_lin.index.tolist(), name='AUMC0-∞ '+"("+measure_unit_lin +"×ч\u00B2" +")")

                       series_Сmax_dev_AUC0_t= df_concat_PK_lin['Сmax/AUC0-t']
                       list_Сmax_dev_AUC0_t_str_f=["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t.tolist()]
                       series_Сmax_dev_AUC0_t=pd.Series(list_Сmax_dev_AUC0_t_str_f, index = df_concat_PK_lin.index.tolist(), name='Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")")

                       series_Kel= df_concat_PK_lin['Kel']
                       list_Kel_str_f=["%.4f" % round(v,4) for v in series_Kel.tolist()]
                       series_Kel=pd.Series(list_Kel_str_f, index = df_concat_PK_lin.index.tolist(), name='Kel '+"("+"ч\u207B\u00B9"+")")

                       series_CL= df_concat_PK_lin['CL/F']
                       list_CL_str_f=["%.2f" % round(v,2) for v in series_CL.tolist()]
                       series_CL=pd.Series(list_CL_str_f, index = df_concat_PK_lin.index.tolist(), name='CL/F ' +"("+"л/ч"+")")

                       series_Vd= df_concat_PK_lin['Vd/F']
                       list_Vd_str_f=["%.1f" % round(v,1) for v in series_Vd.tolist()]
                       series_Vd=pd.Series(list_Vd_str_f, index = df_concat_PK_lin.index.tolist(), name='Vd/F ' +"("+"л/кг"+")")
                       
                       if st.session_state["agree_cmax2 - линейность"] == True:
                          df_total_PK_lin = pd.concat([series_Cmax, series_Tmax, series_Cmax_2, series_Tmax_2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1) 
                       else:
                          df_total_PK_lin = pd.concat([series_Cmax, series_Tmax, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1)

                       df_total_PK_lin.index.name = 'Номер'

                       ##изменение названий параметров описательной статистики

                       df_total_PK_lin1=df_total_PK_lin.copy()
                       df_total_PK_lin1.iloc[-6,:],df_total_PK_lin1.iloc[-2,:]=df_total_PK_lin.iloc[-2,:],df_total_PK_lin.iloc[-6,:]

                       df_total_PK_lin=df_total_PK_lin1

                       df_total_PK_lin1=df_total_PK_lin.copy()
                       df_total_PK_lin1.iloc[-4,:],df_total_PK_lin1.iloc[-5,:]=df_total_PK_lin.iloc[-5,:],df_total_PK_lin.iloc[-4,:]

                       df_total_PK_lin=df_total_PK_lin1

                       df_total_PK_lin = df_total_PK_lin.rename({'Gmean': 'SD', 'std': 'Gmean','median': 'Минимум', 'min': 'Медиана','max': 'Максимум','mean': 'Mean'}, axis='index')

                       table_heading='Фармакокинетические показатели препарата в дозировке ' +file_name +" "+ measure_unit_lin
                       list_heading_word.append(table_heading)

                       list_table_word.append(df_total_PK_lin)
                       #создание списков фреймов, доз и т.д.

                       ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                       df_concat = create_table_descriptive_statistics(df)['df_concat']

                       list_name_doses.append(file_name)
                       list_df_unrounded.append(df_concat_PK_lin)
                       list_df_for_mean_unround_for_graphics.append(df_concat)
                
                ###Кнопка активации дальнейших действий
                button_calculation = False
                
                if (list_keys_file_lin != []) and measure_unit_lin and measure_unit_dose_lin:
                 
                   condition_cmax1 =  len(list_cmax_1_lin) == count_rows_number_lin
                   
                   if st.session_state["agree_cmax2 - линейность"] == True:
                      condition_cmax2 =  len(list_cmax_2_lin) == count_rows_number_lin
                   
                   if st.session_state["agree_cmax2 - линейность"] == True:
                      if (condition_cmax2):
                         button_calculation = True
                   if st.session_state["agree_cmax2 - линейность"] == False:
                      if (condition_cmax1):
                         button_calculation = True

                   if button_calculation == True:
                      st.write('👩🏽‍💻Расчеты произведены!')
                   else:   
                      st.write('🔧Заполните все поля ввода и загрузите файлы!')
                
                if (list_keys_file_lin != []) and measure_unit_lin and measure_unit_dose_lin and button_calculation:
                   
                   if st.session_state["agree_cmax2 - линейность"] == True:
                      list_list_PK_par_mean=[]
                      for i in list_df_unrounded: 
                          mean_сmax=i['Cmax'].loc['mean']
                          mean_tmax=i['Tmax'].loc['mean']
                          mean_сmax2=i['Cmax(2)'].loc['mean']
                          mean_tmax2=i['Tmax(2)'].loc['mean']
                          mean_mrt0inf=i['MRT0→∞'].loc['mean']
                          mean_thalf=i['T1/2'].loc['mean']
                          mean_auc0t=i['AUC0-t'].loc['mean']
                          mean_auc0inf=i['AUC0→∞'].loc['mean']
                          mean_aumc0inf=i['AUMC0-∞'].loc['mean']
                          mean_сmaxdevaucot=i['Сmax/AUC0-t'].loc['mean']
                          mean_kel=i['Kel'].loc['mean']
                          mean_cl=i['CL/F'].loc['mean']
                          mean_vd=i['Vd/F'].loc['mean']
                          list_list_PK_par_mean.append([mean_сmax,mean_tmax,mean_сmax2,mean_tmax2,mean_mrt0inf,mean_thalf,mean_auc0t,mean_auc0inf,mean_aumc0inf,mean_сmaxdevaucot,mean_kel,mean_cl,mean_vd])
                   else:
                      list_list_PK_par_mean=[]
                      for i in list_df_unrounded: 
                          mean_сmax=i['Cmax'].loc['mean']
                          mean_tmax=i['Tmax'].loc['mean']
                          mean_mrt0inf=i['MRT0→∞'].loc['mean']
                          mean_thalf=i['T1/2'].loc['mean']
                          mean_auc0t=i['AUC0-t'].loc['mean']
                          mean_auc0inf=i['AUC0→∞'].loc['mean']
                          mean_aumc0inf=i['AUMC0-∞'].loc['mean']
                          mean_сmaxdevaucot=i['Сmax/AUC0-t'].loc['mean']
                          mean_kel=i['Kel'].loc['mean']
                          mean_cl=i['CL/F'].loc['mean']
                          mean_vd=i['Vd/F'].loc['mean']
                          list_list_PK_par_mean.append([mean_сmax,mean_tmax,mean_mrt0inf,mean_thalf,mean_auc0t,mean_auc0inf,mean_aumc0inf,mean_сmaxdevaucot,mean_kel,mean_cl,mean_vd]) 

                   list_name_doses_with_measure_unit=[]
                   for i in list_name_doses:
                    j= i + " " + measure_unit_lin
                    list_name_doses_with_measure_unit.append(j)

                   ### получение итогового фрейма ФК параметров доз
                   if st.session_state["agree_cmax2 - линейность"] == True:
                      df_PK_doses_total = pd.DataFrame(list_list_PK_par_mean, columns =['Cmax ' +"("+measure_unit_lin+")",'Cmax(2) ' +"("+measure_unit_lin+")",'Tmax ' +"("+"ч"+")",'Tmax(2) ' +"("+"ч"+")",'MRT0→∞ '+"("+"ч"+")",'T1/2 '+"("+"ч"+")",'AUC0-t '+"("+measure_unit_lin+"×ч" +")",'AUC0→∞ '+"("+measure_unit_lin+"×ч" +")",'AUMC0-∞ '+"("+measure_unit_lin+"×ч\u00B2" +")",'Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")",'Kel '+"("+"ч\u207B\u00B9"+")",'CL/F ' +"("+"л/ч"+")",'Vd/F ' +"("+"л/кг"+")"],index=list_name_doses_with_measure_unit)
                   else:
                      df_PK_doses_total = pd.DataFrame(list_list_PK_par_mean, columns =['Cmax ' +"("+measure_unit_lin+")",'Tmax ' +"("+"ч"+")",'MRT0→∞ '+"("+"ч"+")",'T1/2 '+"("+"ч"+")",'AUC0-t '+"("+measure_unit_lin+"×ч" +")",'AUC0→∞ '+"("+measure_unit_lin+"×ч" +")",'AUMC0-∞ '+"("+measure_unit_lin+"×ч\u00B2" +")",'Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")",'Kel '+"("+"ч\u207B\u00B9"+")",'CL/F ' +"("+"л/ч"+")",'Vd/F ' +"("+"л/кг"+")"],index=list_name_doses_with_measure_unit)

                   df_PK_doses_total_transpose=df_PK_doses_total.transpose()

                   #округление фрейма df_PK_doses_total_transpose

                   df_doses_trans_trans=df_PK_doses_total_transpose.transpose()

                   series_Cmax=df_doses_trans_trans['Cmax ' +"("+measure_unit_lin+")"].tolist() 
                   series_Cmax=pd.Series(["%.2f" % round(v,2) for v in series_Cmax])

                   series_Tmax=df_doses_trans_trans['Tmax ' +"("+"ч"+")"].tolist()       
                   series_Tmax=pd.Series(["%.2f" % round(v,2) for v in series_Tmax])

                   if st.session_state["agree_cmax2 - линейность"] == True:
                      series_Cmax2=df_doses_trans_trans['Cmax(2) ' +"("+measure_unit_lin+")"].tolist() 
                      series_Cmax2=pd.Series(["%.2f" % round(v,2) for v in series_Cmax2])

                      series_Tmax2=df_doses_trans_trans['Tmax(2) ' +"("+"ч"+")"].tolist()       
                      series_Tmax2=pd.Series(["%.2f" % round(v,2) for v in series_Tmax2])

                   series_MRT0_inf= df_doses_trans_trans['MRT0→∞ '+"("+"ч"+")"].tolist()   
                   series_MRT0_inf=pd.Series(["%.3f" % round(v,3) for v in series_MRT0_inf])

                   series_half_live= df_doses_trans_trans['T1/2 '+"("+"ч"+")"].tolist()   
                   series_half_live=pd.Series(["%.2f" % round(v,2) for v in series_half_live]) 

                   series_AUC0_t= df_doses_trans_trans['AUC0-t '+"("+measure_unit_lin+"×ч" +")"].tolist()   
                   series_AUC0_t=pd.Series(["%.2f" % round(v,2) for v in series_AUC0_t])

                   series_AUC0_inf= df_doses_trans_trans['AUC0→∞ '+"("+measure_unit_lin+"×ч" +")"].tolist()  
                   series_AUC0_inf=pd.Series(["%.2f" % round(v,2) for v in series_AUC0_inf]) 

                   series_AUMC0_inf= df_doses_trans_trans['AUMC0-∞ '+"("+measure_unit_lin+"×ч\u00B2" +")"].tolist()   
                   series_AUMC0_inf=pd.Series(["%.2f" % round(v,2) for v in series_AUMC0_inf])

                   series_Сmax_dev_AUC0_t= df_doses_trans_trans['Сmax/AUC0-t '+"("+"ч\u207B\u00B9"+")"].tolist()  
                   series_Сmax_dev_AUC0_t=pd.Series(["%.4f" % round(v,4) for v in series_Сmax_dev_AUC0_t]) 

                   series_Kel= df_doses_trans_trans['Kel '+"("+"ч\u207B\u00B9"+")"].tolist()   
                   series_Kel=pd.Series(["%.4f" % round(v,4) for v in series_Kel])

                   series_CL= df_doses_trans_trans['CL/F ' +"("+"л/ч"+")"].tolist()  
                   series_CL=pd.Series(["%.2f" % round(v,2) for v in series_CL]) 

                   series_Vd= df_doses_trans_trans['Vd/F ' +"("+"л/кг"+")"].tolist()   
                   series_Vd=pd.Series(["%.1f" % round(v,1) for v in series_Vd])
                   
                   if st.session_state["agree_cmax2 - линейность"] == True:
                      df_total_total_doses = pd.concat([series_Cmax, series_Tmax,series_Cmax2, series_Tmax2, series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1)
                   else:
                      df_total_total_doses = pd.concat([series_Cmax, series_Tmax,series_MRT0_inf,series_half_live,series_AUC0_t,series_AUC0_inf,series_AUMC0_inf,series_Сmax_dev_AUC0_t,series_Kel,series_CL,series_Vd], axis= 1)

                   df_total_total_doses.index=df_PK_doses_total_transpose.columns.tolist()
                   df_total_total_doses.columns=df_PK_doses_total_transpose.index.tolist() 

                   df_total_total_doses_total= df_total_total_doses.transpose()
                   df_total_total_doses_total.index.name = 'Параметры, размерность'
                
                   table_heading='Фармакокинетические параметры препарата в различных дозировках'
                   list_heading_word.append(table_heading)

                   list_table_word.append(df_total_total_doses_total)
                   ###построение графика "Фармакокинетический профиль в различных дозировках"

                   ### в линейных координатах
                   list_list_mean_conc=[]
                   list_list_std_conc=[]
                   for i in list_df_for_mean_unround_for_graphics: 
                       mean_conc_list=i.loc['mean'].tolist()
                       std_conc_list=i.loc['std'].tolist()
                       list_list_mean_conc.append(mean_conc_list)
                       list_list_std_conc.append(std_conc_list)

                   list_name_doses_with_measure_unit_std=[]
                   for i in list_name_doses_with_measure_unit:
                    j= i + " std"
                    list_name_doses_with_measure_unit_std.append(j)

                   df_mean_conc_graph = pd.DataFrame(list_list_mean_conc, columns =list_t_graph[0],index=list_name_doses_with_measure_unit)
                   df_mean_conc_graph_1=df_mean_conc_graph.transpose()
                   df_std_conc_graph = pd.DataFrame(list_list_std_conc, columns =list_t_graph[0],index=list_name_doses_with_measure_unit_std)
                   df_std_conc_graph_1=df_std_conc_graph.transpose()
                   df_concat_mean_std= pd.concat([df_mean_conc_graph_1,df_std_conc_graph_1],sort=False,axis=1)

                   list_colors = [] ## генерация 500 цветов
                   for i in range(0,500):
                       hexadecimal = "#"+''.join([random.choice('ABCDEF0123456789') for i in range(6)])
                       list_colors.append(hexadecimal)
                       
                   list_zip_mean_std_colors=zip(list_name_doses_with_measure_unit,list_name_doses_with_measure_unit_std,list_colors)

                   fig, ax = plt.subplots()
                   for i,j,c in list_zip_mean_std_colors:
                        plt.errorbar(list(df_concat_mean_std.index),df_concat_mean_std[i],yerr=df_concat_mean_std[j],color= c, marker='o',markersize=4.0,markeredgecolor=c,markerfacecolor=c,ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0,label=i)
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+ measure_unit_lin)
                        ax.legend(fontsize = 8)
                  
                   list_graphics_word.append(fig)

                   graphic='Сравнение фармакокинетических профилей (в линейных координатах) в различных дозировках'
                   list_heading_graphics_word.append(graphic) 

                   ### в полулог. координатах
                   
                   list_t_doses=list(df_concat_mean_std.index)
                   list_t_doses.remove(0)
                   df_concat_mean_std_without_0=df_concat_mean_std.drop([0])
                   list_zip_mean_std_colors=zip(list_name_doses_with_measure_unit,list_name_doses_with_measure_unit_std,list_colors)

                   fig, ax = plt.subplots()
                   for i,j,c in list_zip_mean_std_colors:
                        plt.errorbar(list_t_doses,df_concat_mean_std_without_0[i],yerr=df_concat_mean_std_without_0[j],color= c, marker='o',markersize=4.0,markeredgecolor=c,markerfacecolor=c,ecolor="black",elinewidth=0.8,capsize=2.0,capthick=1.0,label=i)
                        ax.set_yscale("log")
                        plt.xlabel("Время, ч")
                        plt.ylabel("Концентрация, "+ measure_unit_lin)
                        ax.legend(fontsize = 8)
                   
                   list_graphics_word.append(fig)

                   graphic='Сравнение фармакокинетических профилей (в полулогарифмических координатах) в различных дозировках'
                   list_heading_graphics_word.append(graphic)
                   #линейность

                   list_AUC0_inf_lin_mean=[]
                   for i in list_df_unrounded: 
                       mean_auc0inf=i['AUC0→∞'].loc['mean']
                       list_AUC0_inf_lin_mean.append(mean_auc0inf)

                   list_name_doses_lin_float=[]
                   for i in list_name_doses:
                    j= float(i)
                    list_name_doses_lin_float.append(j)
                   
                   df_for_lin = pd.DataFrame(list(zip(list_AUC0_inf_lin_mean,list_name_doses_lin_float)), columns =['AUC0→∞_mean', 'doses'])

                   doses = df_for_lin['doses']
                   AUC0_inf_mean = df_for_lin['AUC0→∞_mean']

                   doses = sm.add_constant(doses)
                   model = sm.OLS(AUC0_inf_mean, doses).fit()
                   predictions = model.predict(doses) 
                   print_model = model.summary()
                   
                   graphic='Зависимость значений AUC0→∞ от величин вводимых доз'
                   list_heading_graphics_word.append(graphic) 
                   ###график
                   fig, ax = plt.subplots()
                   sns.regplot(x='doses',y='AUC0→∞_mean',data=df_for_lin, color="black",ci=None,scatter_kws = {'s': 30}, line_kws = {'linewidth': 1})
                   plt.xlabel("Дозировка, " +measure_unit_dose_lin)
                   plt.ylabel("AUC0→∞, "+ measure_unit_lin + "*ч")
                   plt.annotate('y = ' + "%.4f" % round(model.params[1],4) +'x ' + "%.4f" % round(model.params[0],4), xy =(110, 530),xytext =(110, 530),fontsize=10)
                   
                   list_graphics_word.append(fig)

                   graphic='Коэффициент линейной регрессии и критерий Фишера значимости линейной регрессии для параметра AUC0→∞'
                   list_heading_graphics_word.append(graphic) 
                   # параметры линейной регрессии
                   fig, ax = plt.subplots()
                   table_data_first=[
                    ["R²","F","Df Residuals","Df Model","p"],
                    ["%.3f" % round(model.rsquared,3), int(round(model.fvalue,0)),int(round(model.df_resid,0)),int(round(model.df_model,0)),"%.3f" % round(model.pvalues[1],3)]
                    ]
                   table = ax.table(cellText=table_data_first,cellLoc='left',bbox = [0, 0.7, 0.7, 0.1])
                   plt.annotate('Model Fit Measures', xy =(0, 0.9),xytext =(0, 0.9),fontsize=10)
                   plt.annotate('Overall Model Test', xy =(0, 0.85),xytext =(0, 0.85),fontsize=10)
                   table_data_second=[
                    ['Predictor','Estimate','SE','t','p'],
                    ["Intercept","%.2f" % round(model.params[0],2),"%.3f" % round(model.HC2_se[0],3),"%.3f" % round(model.tvalues[0],3),"%.3f" % round(model.pvalues[0],3)],
                    ["B","%.2f" % round(model.params[1],2),"%.3f" % round(model.HC2_se[1],3),"%.3f" % round(model.tvalues[1],3),"%.3f" % round(model.pvalues[1],3)]
                    ]
                   table = ax.table(cellText=table_data_second,cellLoc='left',bbox = [0, 0.35, 0.7, 0.2])
                   plt.annotate('Model Coefficients', xy =(0, 0.6),xytext =(0, 0.6),fontsize=10)
                   plt.axis('off')
                   
                   list_graphics_word.append(fig)

            ###сохранение состояния 
            st.session_state["list_heading_word"] = list_heading_word
            st.session_state["list_table_word"] = list_table_word
            st.session_state["list_graphics_word"] = list_graphics_word
            st.session_state["list_heading_graphics_word"] = list_heading_graphics_word
            
         
      #####Создание word отчета
         if panel == "Таблицы": 
         
            list_heading_word = st.session_state["list_heading_word"]
            list_table_word = st.session_state["list_table_word"]
            
            ###вызов функции создания таблицы
            create_table(list_heading_word,list_table_word)
            

         if panel == "Графики":
         
            list_graphics_word = st.session_state["list_graphics_word"]
            list_heading_graphics_word = st.session_state["list_heading_graphics_word"]

            ###вызов функции создания графика
            create_graphic(list_graphics_word,list_heading_graphics_word)
                
            #######визуализация

            #классификация графиков по кнопкам
            type_graphics = st.selectbox('Выберите вид графиков',
      ('Индивидуальные фармакокинетические профили', 'Сравнение индивидуальных фармакокинетических профилей', 'Графики усредненного фармакокинетического профиля', "Сравнение фармакокинетических профилей в различных дозировках", "Зависимость значений AUC0→∞ от величин вводимых доз", "Коэффициент линейной регрессии и критерий Фишера значимости линейной регрессии для параметра AUC0→∞"),disabled = False, key = "Вид графика - ИО" )

            count_graphics_for_visual = len(list_heading_graphics_word)
            list_range_count_graphics_for_visual = range(0,count_graphics_for_visual)
            
            for i in list_range_count_graphics_for_visual:
                if list_heading_graphics_word[i].__contains__("индивидуального"): 
                   if type_graphics == 'Индивидуальные фармакокинетические профили':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Сравнение индивидуальных"):   
                   if type_graphics == 'Сравнение индивидуальных фармакокинетических профилей':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("усредненного"):
                   if type_graphics == 'Графики усредненного фармакокинетического профиля':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Сравнение фармакокинетических"):
                   if type_graphics == 'Сравнение фармакокинетических профилей в различных дозировках':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Зависимость"):
                   if type_graphics == 'Зависимость значений AUC0→∞ от величин вводимых доз':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
                if list_heading_graphics_word[i].__contains__("Коэффициент"):
                   if type_graphics == 'Коэффициент линейной регрессии и критерий Фишера значимости линейной регрессии для параметра AUC0→∞':
                      st.pyplot(list_graphics_word[i])
                      st.subheader(list_heading_graphics_word[i])
   
   ###########################################################################################
   if option == 'Изучение экскреции препарата':
       
       st.title('Изучение экскреции препарата')

       col1, col2 = st.columns([0.66, 0.34])
       
       with col2:
            selected = option_menu(None, ["Включение параметров в исследование"], 
            icons=['menu-button'], 
            menu_icon="cast", default_index=0, orientation="vertical",
            styles={
                "container": {"padding": "0!important", "background-color": "#24769C"},
                "icon": {"color": "#5DAED3", "font-size": "13px"}, 
                "nav-link": {"font-size": "13px", "text-align": "left", "margin":"0px", "--hover-color": "#eee"},
                "nav-link-selected": {"background-color": "#335D70"},
            })

       ####### основной экран
       with col1:         
            panel = st.radio(
               "⚙️Панель управления",
               ("Загрузка файлов", "Таблицы","Графики"),
               horizontal=True, key= "Загрузка файлов - Изучение экскреции препарата"
            )
            
            ###создание состояния
            if "measure_unit_ex" not in st.session_state:   
               st.session_state["measure_unit_ex"] = ""
               
            #cписки для word-отчета
            list_heading_word=[]
            list_table_word=[]
            list_graphics_word=[]
            list_heading_graphics_word=[]

            if panel == "Загрузка файлов":
               
               #cостояние радио-кнопки "type_ex"
               if "index_type_ex" not in st.session_state:
                   st.session_state["index_type_ex"] = 0

               type_excretion = st.radio('💩Выберите вид экскреции',('Кал', 'Моча', 'Желчь'), key = "Вид экскреции",index = st.session_state["index_type_ex"])
               
               if st.session_state["Вид экскреции"] == 'Кал':
                  st.session_state["index_type_ex"] = 0
               if st.session_state["Вид экскреции"] == 'Моча':
                  st.session_state["index_type_ex"] = 1
               if st.session_state["Вид экскреции"] == 'Желчь':
                  st.session_state["index_type_ex"] = 2

               if type_excretion == 'Кал':
                  excretion_tv = "калом"
                  excretion_pr = "кале"
               if type_excretion == 'Моча':
                  excretion_tv = "мочой"
                  excretion_pr = "моче"
               if type_excretion == 'Желчь':
                  excretion_tv = "желчью"
                  excretion_pr = "желчи"

               st.title('Исследование экскреции с ' + excretion_tv)

               measure_unit_ex = st.text_input("Введите единицы измерения концентрации", key='Единицы измерения при изучении экскреции препарата', value = st.session_state["measure_unit_ex"])
                   
               st.session_state["measure_unit_ex"] = measure_unit_ex

               uploaded_file_excrement = st.file_uploader("Выбрать файл экскреции (формат XLSX)", key="Файл экскреции")

               if uploaded_file_excrement is not None:
                   save_uploadedfile(uploaded_file_excrement)
                   st.session_state["uploaded_file_excrement"] = uploaded_file_excrement.name

               if "uploaded_file_excrement" in st.session_state and measure_unit_ex:
                   
                   df = pd.read_excel(os.path.join("Папка для сохранения файлов",st.session_state["uploaded_file_excrement"]))
                   st.subheader('Индивидуальные значения концентраций в ' + excretion_pr)
                   
                   ###интерактивная таблица
                   df = edit_frame(df,st.session_state["uploaded_file_excrement"])

                   table_heading='Индивидуальные и усредненные значения концентраций в ' + excretion_pr
                   list_heading_word.append(table_heading) 

                   ## вызов функции подсчета опистательной статистики и создания соотвествующей таблицы с округлениями
                   df_concat_round_str_transpose = create_table_descriptive_statistics(df)['df_concat_round_str_transpose']

                   list_table_word.append(df_concat_round_str_transpose) 

                   ########### диаграмма    
                   
                   col_mapping = df.columns.tolist()
                   col_mapping.remove('Номер')

                   list_time = []
                   for i in col_mapping:
                       numer=float(i)
                       list_time.append(numer)
                   
                   df_averaged_concentrations=df.describe()
                   list_concentration=df_averaged_concentrations.loc['mean'].tolist()

                   list_concentration.remove(0)
                   list_time.remove(0)

                   fig, ax = plt.subplots()

                   sns.barplot(x=list_time, y=list_concentration,color='blue',width=0.5)
                   plt.xlabel("Время, ч")
                   plt.ylabel("Концентрация, "+measure_unit_ex)

                   list_graphics_word.append(fig)

                   graphic='Выведение с ' + excretion_tv
                   list_heading_graphics_word.append(graphic)
               else:
                  st.write("")    
               
               ##############################################################################################################

               ###сохранение состояния 
               st.session_state["list_heading_word"] = list_heading_word
               st.session_state["list_table_word"] = list_table_word
               st.session_state["list_graphics_word"] = list_graphics_word
               st.session_state["list_heading_graphics_word"] = list_heading_graphics_word

            #####Создание word отчета
            if panel == "Таблицы":

                  list_heading_word = st.session_state["list_heading_word"]
                  list_table_word = st.session_state["list_table_word"]

                  ###вызов функции создания таблицы
                  create_table(list_heading_word,list_table_word)

            if panel == "Графики":
                  
                  list_graphics_word = st.session_state["list_graphics_word"]
                  list_heading_graphics_word = st.session_state["list_heading_graphics_word"]
                  
                  ###вызов функции создания графика
                  create_graphic(list_graphics_word,list_heading_graphics_word)

                  #######визуализация

                  count_graphics_for_visual = len(list_heading_graphics_word)
                  list_range_count_graphics_for_visual = range(0,count_graphics_for_visual)
                  
                  for i in list_range_count_graphics_for_visual:
                      if list_heading_graphics_word[i].__contains__("Выведение"):
                         st.pyplot(list_graphics_word[i])
                         st.subheader(list_heading_graphics_word[i])
                      