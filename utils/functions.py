import streamlit as st
import os
import pandas as pd
from io import BytesIO
from docx import Document

import tempfile
import math

from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from streamlit_option_menu import option_menu

import networkx as nx
from pyvis.network import Network
import random
import string
import hashlib

def sort_by_keys_with_indices(list_data, list_keys):
    indexed_list = [(i, item) for i, item in enumerate(list_data)]  # Сохраняем изначальные индексы
    sorted_list = sorted(indexed_list, key=lambda x: next((i for i, key in enumerate(list_keys) if key in x[1]), float('inf')))
    
    sorted_data = [item[1] for item in sorted_list]  # Отсортированные строки
    index_mapping = {old_idx: new_idx for new_idx, (old_idx, _) in enumerate(sorted_list)}  # Сопоставление индексов
    
    return sorted_data, index_mapping

def reorder_list_by_mapping(original_list, index_mapping):
    reordered_list = [None] * len(original_list)  # Создаем пустой список нужного размера
    for old_idx, new_idx in index_mapping.items():
        reordered_list[new_idx] = original_list[old_idx]  # Переставляем элементы
    return reordered_list

#функция соотвествий для биодоступности
# Функция визуализации
# Вызов функции визуализации
def get_color(name):
    """Генерирует уникальный цвет на основе имени узла."""
    hash_object = hashlib.md5(name.encode())  # Генерируем хеш на основе имени
    random.seed(int(hash_object.hexdigest(), 16))  # Используем хеш как seed
    return "#{:06x}".format(random.randint(0, 0xFFFFFF))  # Генерируем цвет

def visualize_mapping(list_keys_file_bioavailability):
    st.title("Дизайн исследования")

    if "graph" not in st.session_state:
        st.session_state.graph = nx.DiGraph()

    # Автоматическое добавление узлов и связей
    for drug in list_keys_file_bioavailability:
        if drug:  # Проверяем, не пустой ли элемент
            st.session_state.graph.add_node(drug)

    # Визуализация графа
    def render_graph():
        net = Network(notebook=True, height="500px", width="100%", directed=True)

        for node in st.session_state.graph.nodes:
            net.add_node(node, title=node, color=get_color(node))  # Уникальный цвет для каждого узла

        for edge in st.session_state.graph.edges:
            net.add_edge(edge[0], edge[1], title=f"{edge[0]} → {edge[1]}")

        net.show("graph.html")
        with open("graph.html", "r", encoding="utf-8") as f:
            html_content = f.read()
        st.components.v1.html(html_content, height=500)

    st.subheader("Управление узлами и связями")
    col1, col2 = st.columns(2)
    col3, col4 = st.columns([0.30,0.70])


    nodes = list(st.session_state.graph.nodes)
    if len(nodes) >= 2:
       with col1:
          source = st.selectbox("Референт", nodes, key="source_node")
          target = st.selectbox("Исследуемый", nodes, key="target_node")
       with col2:
             
             with col3:
                 if st.button("Добавить связь") and source and target and source != target:
                     st.session_state.graph.add_edge(source, target)
             with col4:
                 if st.button("Очистить граф"):
                     st.session_state.graph.clear()

    render_graph()

    # Отображение связей
    edges = list(st.session_state.graph.edges)

    if len(edges) != 0:
       cols = st.columns(len(edges))
       for i, edge in enumerate(edges,start=0):
           with cols[i]:  # Каждая кнопка в своей колонке
               if st.button(f"❌ Связь №{i+1}", key=f"del_{edge[0]}_{edge[1]}", help=f"{edge[0]} → {edge[1]}"):
                   st.session_state.graph.remove_edge(edge[0], edge[1])
                   st.rerun()


        
    selected_edges = [f"{edge[0]} → {edge[1]}" for edge in st.session_state.graph.edges]
    return selected_edges


#основная радиокнопка исследования
def main_radio_button_study(option):
    panel = st.radio(
            "⚙️Панель управления",
            ("Загрузка файлов", "Таблицы","Графики"),
            horizontal=True, key= f"Загрузка файлов - {option}"
        )
    
    return panel

#инициализация состояния дозы и времени инфузии
def initialization_dose_infusion_time_session(option,file_name=None):
    
    if file_name is None:
        if f"dose_{option}" not in st.session_state:
            st.session_state[f"dose_{option}"] = ""

        if f"infusion_time_{option}" not in st.session_state:
            st.session_state[f"infusion_time_{option}"] = ""
    else:
        if f"dose_{option}_{file_name}" not in st.session_state:
            st.session_state[f"dose_{option}_{file_name}"] = ""

        if f"infusion_time_{option}_{file_name}" not in st.session_state:
            st.session_state[f"infusion_time_{option}_{file_name}"] = ""

def settings_additional_research_parameters(option,custom_success,key=None,file_name=None):
    
    if key is None and file_name is None:
       #оформительский элемент настройки дополнительных параметров исследования
       selected = style_icon_setting_additional_parameters(key,file_name)
       
       if selected == "Настройка дополнительных параметров":
           type_parameter = st.selectbox('Выберите параметр',
           ("Вид введения",'Двойные пики'),disabled = False, key = f"Вид параметра - {option}")
    else:
       #оформительский элемент настройки дополнительных параметров исследования
       selected = style_icon_setting_additional_parameters(key,file_name)

       if selected == f"Настройка дополнительных параметров для «{file_name}»":
            type_parameter = st.selectbox('Выберите параметр',
            ('Вид введения','-'),disabled = False, key = f"Вид параметра - {option}_{file_name}")   
    

    if key is None and file_name is None:

       if f"agree_cmax2 - {option}" not in st.session_state:
               st.session_state[f"agree_cmax2 - {option}"] = False
       
       if type_parameter == 'Двойные пики':

           st.session_state[f"agree_cmax2 - {option}"] = st.checkbox('В зависимости "Концентрация-Время" отчетливо наблюдаются двойные пики', key = f"Возможность добавления Cmax2 - {option}", value = st.session_state[f"agree_cmax2 - {option}"])
           
           if st.session_state[f"agree_cmax2 - {option}"] == True:
               custom_success('Параметр добавлен!')

       if f"agree_injection - {option}" not in st.session_state:
               st.session_state[f"agree_injection - {option}"] = "extravascular"

    else:
        
       if f"agree_injection - {option}_{file_name}" not in st.session_state:
               st.session_state[f"agree_injection - {option}_{file_name}"] = "extravascular"


    if key is None and file_name is None:
       
       if type_parameter == "Вид введения":

           # Проверка наличия значения в сессии, если его нет, устанавливаем значение по умолчанию
           if f"injection_choice - {option}" not in st.session_state:
               st.session_state[f"injection_choice - {option}"] = 0  # Значение по умолчанию

           # Радиокнопка для выбора типа введения
           injection_type = st.radio(
               "Выберите тип введения:",
               options=["Внутривенный болюс", "Внесосудистое введение", "Инфузионное введение"],
               index=st.session_state[f"injection_choice - {option}"],
               key=f"injection_choice_{option}",  # Ключ для сохранения выбора в сессии
           )

           # Логика для обновления состояния сессии
           if injection_type == "Внутривенный болюс":
               st.session_state[f"agree_injection - {option}"] = "intravenously"
               st.session_state[f"injection_choice - {option}"] = 0
           elif injection_type == "Внесосудистое введение":
               st.session_state[f"agree_injection - {option}"] = "extravascular"
               st.session_state[f"injection_choice - {option}"] = 1
           else:
               st.session_state[f"agree_injection - {option}"] = "infusion"
               st.session_state[f"injection_choice - {option}"] = 2

           # Сообщение в зависимости от выбора
           if st.session_state[f"agree_injection - {option}"] == "intravenously":
               custom_success("Выбрано: Внутривенный болюс!")
           elif st.session_state[f"agree_injection - {option}"] == "extravascular":
               custom_success("Выбрано: Внесосудистое введение!")
           else:
               custom_success("Выбрано: Инфузионное введение!")
    else: 
       
      if type_parameter == "Вид введения":

           # Проверка наличия значения в сессии, если его нет, устанавливаем значение по умолчанию
           if f"injection_choice - {option}_{file_name}" not in st.session_state:
               st.session_state[f"injection_choice - {option}_{file_name}"] = 0  # Значение по умолчанию

           # Радиокнопка для выбора типа введения
           injection_type = st.radio(
               "Выберите тип введения:",
               options=["Внутривенный болюс", "Внесосудистое введение", "Инфузионное введение"],
               index=st.session_state[f"injection_choice - {option}_{file_name}"],
               key=f"injection_choice_{option}_{file_name}",  # Ключ для сохранения выбора в сессии
           )

           # Логика для обновления состояния сессии
           if injection_type == "Внутривенный болюс":
               st.session_state[f"agree_injection - {option}_{file_name}"] = "intravenously"
               st.session_state[f"injection_choice - {option}_{file_name}"] = 0
           elif injection_type == "Внесосудистое введение":
               st.session_state[f"agree_injection - {option}_{file_name}"] = "extravascular"
               st.session_state[f"injection_choice - {option}_{file_name}"] = 1
           else:
               st.session_state[f"agree_injection - {option}_{file_name}"] = "infusion"
               st.session_state[f"injection_choice - {option}_{file_name}"] = 2

           # Сообщение в зависимости от выбора
           if st.session_state[f"agree_injection - {option}_{file_name}"] == "intravenously":
               custom_success("Выбрано: Внутривенный болюс!")
           elif st.session_state[f"agree_injection - {option}_{file_name}"] == "extravascular":
               custom_success("Выбрано: Внесосудистое введение!")
           else:
               custom_success("Выбрано: Инфузионное введение!")



#чтобы не добавлять по несколько раз в session_state
def add_or_replace_df_graph(list_heading,list_element,heading,element):
    try:
        index = list_heading.index(heading)  # Ищем индекс элемента
        list_element[index] = element  # Заменяем его
    except IndexError:
        list_element.append(element)  # Добавляем в конец, если элемента нет

#чтобы не добавлять названия графиков, таблиц по несколько раз в session_state
def add_or_replace(list, element):
    try:
        index = list.index(element)  # Ищем индекс элемента
        list[index] = element  # Заменяем его
    except ValueError:
        list.append(element)  # Добавляем в конец, если элемента нет

#сохранение состояния единиц измерения исследований после выбора их пользователем
def save_session_state_measure_unit_value(measure_unit_time,measure_unit_concentration,key,measure_unit_dose=None,measure_unit_org_organs=None):
    st.session_state[f'measure_unit_{key}_time'] = measure_unit_time
    st.session_state[f'measure_unit_{key}_concentration'] = measure_unit_concentration
    st.session_state[f'measure_unit_{key}_dose'] = measure_unit_dose
    if key == 'Распределение по органам' and measure_unit_org_organs is not None:
       st.session_state[f'measure_unit_{key}_organs'] = measure_unit_org_organs


#Инизиализация состояния фреймов с результатами исследований
def initializing_session_state_frames_research_results(list_key_research):
    
    for key_research in list_key_research:
        if f"df_total_PK_{key_research}" not in st.session_state:
            st.session_state[f"df_total_PK_{key_research}"] = None


def style_icon_setting_additional_parameters(key,file_name):
    if key is None and file_name is None:
       selected = option_menu(None, ["Настройка дополнительных параметров"], 
                  icons=['menu-button'], 
                  menu_icon="cast", default_index=0, orientation="vertical",
                  styles={
                    "container": {"padding": "0!important", "background-color": "#1f3b57"},
                    "icon": {"color": "#cbe4de", "font-size": "16px"}, 
                    "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#92c4e6","color": "#ffffff"},
                    "nav-link-selected": {"background-color": "#73b5f2"},
                  })
    else:
      selected = option_menu(None, [f"Настройка дополнительных параметров для «{file_name}»"], 
                 icons=['menu-button'], 
                 menu_icon="cast", default_index=0, orientation="vertical",
                 styles={
                   "container": {"padding": "0!important", "background-color": "#1f3b57"},
                   "icon": {"color": "#cbe4de", "font-size": "12px"}, 
                   "nav-link": {"font-size": "12px", "text-align": "left", "margin":"0px", "--hover-color": "#92c4e6","color": "#ffffff"},
                   "nav-link-selected": {"background-color": "#73b5f2"},
                 }, key = f"Настройка дополнительных параметров для {key} «{file_name}»")
    
    return selected

def style_icon_report():
    selected = option_menu(None, ["Cформированный отчeт"], 
               icons=['file-earmark-arrow-down-fill'], 
               menu_icon="cast", default_index=0, orientation="vertical",
               styles={
                   "container": {"padding": "0!important", "background-color": "#1f3b57"},
                   "icon": {"color": "#cbe4de", "font-size": "16px"}, 
                   "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#92c4e6","color": "#ffffff"},
                   "nav-link-selected": {"background-color": "#73b5f2"},
               })
    return selected

def initializing_session_lists_tables_graphics(option,list_heading_word,list_table_word,list_graphics_word,list_heading_graphics_word):
    ###инициализация состояния
    if f"list_heading_word_{option}" not in st.session_state: 
        st.session_state[f"list_heading_word_{option}"] = list_heading_word
        st.session_state[f"list_table_word_{option}"] = list_table_word
        st.session_state[f"list_graphics_word_{option}"] = list_graphics_word
        st.session_state[f"list_heading_graphics_word_{option}"] = list_heading_graphics_word

# Функция для сохранения DataFrame в формате Excel
def to_excel_results(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=True)
    output.seek(0)  # Возвращаем курсор в начало файла
    return output

# Обертка для скачивания файла в формате Excel с поддержкой ключа
def download_excel_button(df, label, key, file_name):
    excel_data = to_excel_results(df)
    st.download_button(
        label="Скачать",
        data=excel_data,
        file_name=file_name,
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        key=key,# Добавлен параметр key
        help = f"{label}"
    )

#округление до определенного значения значищих цифр
def round_to_significant_figures(num, sig_figs):
    # Проверка на строку "-"
    if num == "-":
        return num
    # Проверка на NaN
    elif isinstance(num, float) and math.isnan(num):
        return "-"
    # Проверка на нулевое значение
    elif num == 0:
        return 0
    # Округление для остальных значений
    else:
        # Округление числа до нужного количества значащих цифр
        rounded_num = round(num, sig_figs - int(math.floor(math.log10(abs(num))) + 1))
        
        # Если результат целое число, возвращаем его как int
        if rounded_num.is_integer():
            return int(rounded_num)
        else:
            return rounded_num


#сохранение загружаемых файлов 
def save_uploadedfile(uploadedfile):
    with open(os.path.join("Папка для сохранения файлов",uploadedfile.name),"wb") as f:
       f.write(uploadedfile.getbuffer())
     

#сохранение редактируемых файлов df_edit
def save_editfile(df_edit,uploadedfile_name):
    writer=pd.ExcelWriter(os.path.join("Папка для сохранения файлов",uploadedfile_name))
    df_edit.to_excel(writer,index=False)
    writer.save()


#превращает df в excel файл-пример
def to_excel(df_example_file):
       output = BytesIO()
       writer = pd.ExcelWriter(output, engine='xlsxwriter')
       df_example_file.to_excel(writer, index=False, sheet_name='Sheet1')
       workbook = writer.book
       worksheet = writer.sheets['Sheet1']
       format1 = workbook.add_format({'num_format': '0.00'}) 
       worksheet.set_column('A:A', None, format1)
       writer.save()  
       processed_data = output.getvalue()
       return processed_data
 

###возможность редактирования фрейма исходных данных
def edit_frame(df,uploadedfile_name):
       new_df = df
       list_columns_str = []
       for i in new_df.columns.tolist():
           i_new = str(i)
           list_columns_str.append(i_new)
       new_df.columns = list_columns_str

       edited_df = st.data_editor(new_df, key = ("edited_df" + uploadedfile_name))
       save_editfile(edited_df,uploadedfile_name)

       df_change = edited_df
       
       list_change_values = df_change.columns.tolist()
       list_change_values.remove("Номер")

       list_columns_number = []
       for i in list_change_values:
           i_new = float(i)
           list_columns_number.append(i_new)

       list_columns_number.insert(0,"Номер")

       df_change.columns = list_columns_number
       
       df = df_change
       return df


###создание Word-отчета
## функция создания отчета таблиц
def create_table(list_heading_word, list_table_word):
    zip_heading_table = list(zip(list_heading_word, list_table_word))

    doc = Document()

    # Устанавливаем горизонтальную ориентацию страницы
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Настройка стиля документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    for heading, df in zip_heading_table:
        # Добавление заголовка для каждой таблицы
        paragraph = doc.add_paragraph(heading, style='Heading 1')
        run = paragraph.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Преобразование колонок в DataFrame и добавление индексов
        name_columns = pd.DataFrame(df.columns.tolist()).T
        name_columns.columns = df.columns.tolist()
        df_columns = pd.concat([name_columns, df]).reset_index(drop=True)

        total_name_index = df.index.name or "Index"
        list_index_names = df.index.tolist()
        list_index_names.insert(0, total_name_index)
        series_index_names = pd.Series(list_index_names, name=total_name_index)
        df_series_index_names = series_index_names.to_frame()
        df_columns_indexes = pd.concat([df_series_index_names, df_columns], axis=1)

        # Создание таблицы
        t = doc.add_table(rows=df_columns_indexes.shape[0], cols=df_columns_indexes.shape[1])
        t.style = 'Table Grid'

        # Автоматическая настройка ширины колонок
        max_col_widths = [max([len(str(df_columns_indexes.iat[i, j])) for i in range(df_columns_indexes.shape[0])]) for j in range(df_columns_indexes.shape[1])]
        total_width = 26.0  # Доступная ширина в см
        col_widths = [min(w * 0.2, total_width / len(max_col_widths)) for w in max_col_widths]

        for j, width in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(width)

        # Заполнение таблицы данными
        for i, row_data in df_columns_indexes.iterrows():
            for j, value in enumerate(row_data):
                cell = t.cell(i, j)
                cell.text = str(value)
                # Настройка стиля текста
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Times New Roman'

        # Центрирование текста
        for row in t.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1  # Центрирование по вертикали

    # Сохранение документа в память
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    # Кнопка для скачивания документа
    st.download_button(
        label="Сохранить таблицы 📃",
        data=bio.getvalue(),
        file_name="Таблицы.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

#визуализация и выгрузка в excel
def visualize_table(list_heading_word,list_table_word,option):
    zip_heading_table = list(zip(list_heading_word,list_table_word)) ###еще раз объявляем, иначе не видит zip-объект
    #####визуализация
    for heading, df in zip_heading_table:
        
        if heading == "Таблица биодоступности": 
          width = 500
        else:
          width = None

        with st.container(border=True,key= f"container_PK{heading}",height=500):
             
             st.subheader(heading)

             # Словарь с форматированием для конкретных колонок
             format_rules = {
                 
             }

             # Функция округления в стиле Phoenix (3–4 значащие цифры)
             def phoenix_format(value):
                 try:
                     num = float(value)
                     if num == 0:
                         return "0"
                     elif abs(num) < 1:
                         return f"{num:.4g}"  # Маленькие числа → 3-4 значащие цифры
                     elif abs(num) < 1000:
                         return f"{num:.4g}"  # Средние числа → 4 значащие цифры
                     else:
                         return f"{num:,.0f}"  # Большие числа → Без научной нотации, с разделителями
                 except (ValueError, TypeError):
                     return value  # Оставляем строки без изменений
                 
             # Функция, применяющая нужное форматирование к каждой ячейке
             def safe_format(value, col):
                 if col in format_rules:
                     fmt = format_rules[col]
                     try:
                         return fmt.format(float(value)) if isinstance(value, (int, float)) or str(value).replace('.', '', 1).isdigit() else value
                     except ValueError:
                         return value
                 else:
                     return phoenix_format(value)  # Применяем Phoenix-форматирование

             # Применяем функцию к каждой ячейке в колонках
             # Создаём словарь форматирования для Pandas Styler
             format_dict = {col: lambda x: safe_format(x, col) for col in df.columns}

             # Инициализируем состояние, если оно ещё не задано
             if f"selected_columns{heading}_{option}" not in st.session_state:
                 columns = [str(col) for col in list(df.columns)]
                 st.session_state[f"selected_columns{heading}_{option}"] = columns   # По умолчанию все колонки
                 
             if f"selected_rows{heading}_{option}" not in st.session_state:
                 rows = [str(row) for row in list(df.index)]
                 st.session_state[f"selected_rows{heading}_{option}"] = rows  # По умолчанию все колонки    
             
             # Отображаем DataFrame с форматированием
             selection = st.dataframe(df.style.format(format_dict),on_select = "rerun",selection_mode=["multi-row", "multi-column"],width=width)
             # Проверяем, были ли выбраны колонки
             if selection:
                 selected_row_indices = selection["selection"]["rows"]  # Получаем номера выбранных строк
                 selected_rows = df.index[selected_row_indices]  # Получаем пользовательские индексы
                 selected_rows = [str(row) for row in selected_rows]  # Приведение выбранных колонок к строкам

                 selected_columns = selection["selection"]["columns"]
                 selected_columns = [str(col) for col in selected_columns]  # Приведение выбранных колонок к строкам

                 # Обновляем состояние только если выбор изменился
                 if selected_columns != st.session_state[f"selected_columns{heading}_{option}"] and selected_columns != []:
                     st.session_state[f"selected_columns{heading}_{option}"] = selected_columns

                 # Обновляем состояние только если выбор изменился
                 if selected_rows != st.session_state[f"selected_rows{heading}_{option}"] and selected_rows != []:
                     st.session_state[f"selected_rows{heading}_{option}"] = selected_rows

             st.subheader("Выбранные данные:")
             df.index = df.index.astype(str)  # Приведение к строковому типу
             df.columns = df.columns.astype(str)  # Приведение к строковому типу
             
             try:
               choice_columns = df.loc[st.session_state[f"selected_rows{heading}_{option}"], st.session_state[f"selected_columns{heading}_{option}"]]
             except KeyError as e:
               columns = [str(col) for col in list(df.columns)]
               st.session_state[f"selected_columns{heading}_{option}"] = columns
               rows = [str(row) for row in list(df.index)]
               st.session_state[f"selected_rows{heading}_{option}"] = rows
               st.rerun()

             st.dataframe(choice_columns.style.format(format_dict),width=width)
             col1,col2 = st.columns([0.2,0.8])
             with col1:
                  if st.button("Очистить выбор",key = f"Clear_selection_{heading}_{option}"):
                     st.session_state[f"selected_columns{heading}_{option}"] = []
                     st.session_state[f"selected_rows{heading}_{option}"] = []
             with col2:
                  # Используем кастомные виджеты с уникальными ключами для выгрузки Excel
                  download_excel_button(choice_columns, f"Cкачать файл {heading}", heading,f"{heading}.xlsx")


## функция создания отчета графиков
def create_graphic(list_graphics_word,list_heading_graphics_word):
    ### документ Word
    zip_graphics_heading = list(zip(list_graphics_word,list_heading_graphics_word))
    doc = Document()

    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    buf = BytesIO() #костыль для того, чтобы не вылазила ошибка
    for fig, heading in zip_graphics_heading:
        buf = BytesIO()
        fig.savefig(buf, format="jpg", dpi=300, bbox_inches='tight')
        fp = tempfile.NamedTemporaryFile() 
        with open(f"{fp.name}.jpg",'wb') as ff:
             ff.write(buf.getvalue()) 
        doc.add_picture(buf)
        doc.add_paragraph(heading)
    
    doc.save(buf)
    if doc:
        st.download_button(
            label="Сохранить графики 📈",
            data=buf.getvalue(),
            file_name="Графики.docx",
            mime="docx",
            key = "graphics"
        )